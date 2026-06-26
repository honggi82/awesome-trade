import csv
import hashlib
import html
import json
import math
import re
import shutil
import sys
import time
from collections import Counter, defaultdict
from datetime import date
from pathlib import Path

import requests
from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DATA_DIR = ROOT / "data"
DOCS_DIR = ROOT / "docs"
PAPER_DIR = ROOT / "paper"
CACHE_DIR = DATA_DIR / "cache"

START_YEAR = 2000
END_YEAR = 2026
YEARS = list(range(START_YEAR, END_YEAR + 1))
YEAR_RANGE_TEXT = f"{START_YEAR}-{END_YEAR}"
YEAR_FILE_STEM = f"{START_YEAR}_{END_YEAR}"

PAPERS_JSON = f"papers_{YEAR_FILE_STEM}.json"
PAPERS_CSV = f"papers_{YEAR_FILE_STEM}.csv"
CANDIDATES_JSON = f"candidates_top1000_{YEAR_FILE_STEM}.json"
CANDIDATES_CSV = f"candidates_top1000_{YEAR_FILE_STEM}.csv"
TAXONOMY_CSV = f"papers_taxonomy_{YEAR_FILE_STEM}.csv"
PERIOD_ANALYSIS_JSON = f"period_analysis_{YEAR_FILE_STEM}.json"
PAPER_LOCALIZATIONS_JSON = f"paper_localizations_{YEAR_FILE_STEM}.json"
GITHUB_LINKS_JSON = f"github_links_{YEAR_FILE_STEM}.json"
LINK_AUDIT_JSON = f"link_audit_{YEAR_FILE_STEM}.json"

CANDIDATES_PER_YEAR = 1000
TARGET_PER_YEAR = 100
TARGET_TOTAL = TARGET_PER_YEAR * len(YEARS)
S2_BULK_URL = "https://api.semanticscholar.org/graph/v1/paper/search/bulk"
REQUEST_DELAY = 0.05

S2_FIELDS = ",".join(
    [
        "paperId",
        "title",
        "year",
        "authors",
        "venue",
        "publicationVenue",
        "publicationDate",
        "citationCount",
        "influentialCitationCount",
        "abstract",
        "url",
        "externalIds",
        "openAccessPdf",
        "s2FieldsOfStudy",
        "publicationTypes",
    ]
)

QUERIES = [
    "artificial intelligence",
    "machine learning",
    "deep learning",
    "neural networks",
    "support vector machines",
    "data mining",
    "pattern recognition",
    "bayesian networks",
    "foundation models",
    "large language models",
    "natural language processing",
    "computer vision",
    "reinforcement learning",
    "generative AI diffusion models",
    "graph neural networks",
    "multimodal learning",
    "AI safety fairness explainability robustness",
    "AI for science healthcare robotics",
]

AI_RELEVANCE_PATTERNS = [
    r"\bartificial intelligence\b",
    r"\bAI\b",
    r"\bmachine learning\b",
    r"\bdeep learning\b",
    r"\bneural network",
    r"\bneural networks\b",
    r"\bsupport vector machine",
    r"\bSVM\b",
    r"\bdata mining\b",
    r"\bpattern recognition\b",
    r"\bbayesian network",
    r"\bclassification\b",
    r"\bclustering\b",
    r"\btransformer",
    r"\blanguage model",
    r"\blarge language model",
    r"\bLLM\b",
    r"\bfoundation model",
    r"\bcomputer vision\b",
    r"\bimage recognition\b",
    r"\bimage classification\b",
    r"\bobject detection\b",
    r"\bsemantic segmentation\b",
    r"\bnatural language processing\b",
    r"\bNLP\b",
    r"\breinforcement learning\b",
    r"\bpolicy gradient\b",
    r"\bgenerative model",
    r"\bdiffusion model",
    r"\bgenerative adversarial",
    r"\bGAN\b",
    r"\bgraph neural",
    r"\bself-supervised\b",
    r"\bcontrastive learning\b",
    r"\brepresentation learning\b",
    r"\bfederated learning\b",
    r"\bexplainable AI\b",
    r"\bXAI\b",
    r"\badversarial\b",
    r"\brobustness\b",
    r"\bfairness\b",
    r"\bprivacy\b",
    r"\bspeech recognition\b",
    r"\brecommendation",
    r"\bfew-shot\b",
    r"\bzero-shot\b",
    r"\bprompt",
    r"\bRAG\b",
    r"\bretrieval-augmented\b",
]

BAD_TITLE_PATTERNS = [
    r"^group$",
    r"^editorial\b",
    r"^preface\b",
    r"^front matter\b",
    r"^proceedings of\b",
    r"^international journal of\b",
    r"^advances in neural information processing systems \d+$",
    r"^the sciences of the artificial$",
    r"^statistical learning theory$",
    r"^generative adversarial networks$",
    r"^foundations of machine learning$",
    r"^language models$",
    r"^random forests$",
    r"^deep learning$",
    r"^machine learning$",
    r"^pattern recognition and machine learning$",
    r"^artificial intelligence a modern approach$",
]

IMPORTANT_VENUES = [
    "nature",
    "science",
    "cell",
    "neurips",
    "nips",
    "icml",
    "iclr",
    "aaai",
    "ijcai",
    "cvpr",
    "iccv",
    "eccv",
    "acl",
    "emnlp",
    "naacl",
    "kdd",
    "sigir",
    "www",
    "the web conference",
    "ieee transactions",
    "jmlr",
    "transactions on machine learning research",
    "pattern recognition",
    "artificial intelligence",
    "machine learning",
    "nature machine intelligence",
    "nature methods",
    "nature biomedical engineering",
]

CATEGORIES = [
    (
        "Foundation Models and Large Language Models",
        [
            "large language model",
            "language model",
            "foundation model",
            "gpt",
            "bert",
            "pretrain",
            "prompt",
            "retrieval-augmented",
            "rag",
            "scaling law",
            "instruction",
        ],
    ),
    (
        "Generative Models and Synthetic Media",
        [
            "diffusion",
            "generative",
            "gan",
            "adversarial network",
            "image synthesis",
            "text-to-image",
            "vae",
            "score-based",
            "synthetic",
        ],
    ),
    (
        "Vision and Multimodal Learning",
        [
            "computer vision",
            "image",
            "video",
            "vision transformer",
            "object detection",
            "segmentation",
            "clip",
            "multimodal",
            "vision-language",
            "visual",
            "point cloud",
        ],
    ),
    (
        "Natural Language Processing and Knowledge",
        [
            "natural language",
            "nlp",
            "translation",
            "question answering",
            "dialogue",
            "summarization",
            "information retrieval",
            "knowledge graph",
            "text",
            "speech recognition",
        ],
    ),
    (
        "Reinforcement Learning and Agents",
        [
            "reinforcement",
            "policy",
            "agent",
            "markov",
            "robot",
            "control",
            "planning",
            "imitation",
            "reward",
            "human feedback",
        ],
    ),
    (
        "Representation, Self-Supervised, and Transfer Learning",
        [
            "self-supervised",
            "contrastive",
            "representation",
            "domain adaptation",
            "few-shot",
            "zero-shot",
            "transfer learning",
            "metric learning",
            "distillation",
            "pre-training",
        ],
    ),
    (
        "Trustworthy, Explainable, and Responsible AI",
        [
            "explainable",
            "interpretability",
            "interpretable",
            "fairness",
            "bias",
            "robust",
            "adversarial",
            "privacy",
            "safety",
            "uncertainty",
            "calibration",
        ],
    ),
    (
        "Graph Learning, Recommendation, and Core Methods",
        [
            "graph neural",
            "recommendation",
            "recommender",
            "optimization",
            "hyperparameter",
            "neural architecture",
            "bayesian",
            "random forest",
            "ensemble",
            "kernel",
            "decision tree",
        ],
    ),
    (
        "AI for Science, Healthcare, and Robotics",
        [
            "protein",
            "molecule",
            "drug",
            "genomic",
            "healthcare",
            "medical",
            "clinical",
            "biomedical",
            "science",
            "robotics",
            "autonomous driving",
            "biology",
        ],
    ),
]

KEYWORD_CONVENTION = [
    (
        "machine-learning",
        "Statistical learning, neural networks, SVMs, trees/boosting, data mining, pattern recognition, clustering, or core ML methods.",
        "475569",
    ),
    (
        "foundation-models",
        "Large language models, foundation models, scaling, prompting, alignment, or retrieval-augmented systems.",
        "2563eb",
    ),
    (
        "generative-ai",
        "Generative adversarial, diffusion, synthetic media, text-to-image, or other model-based generation work.",
        "a855f7",
    ),
    (
        "multimodal",
        "Vision-language, audio-language, video-language, or cross-modal representation learning.",
        "0891b2",
    ),
    (
        "nlp",
        "Natural language processing, language modeling, retrieval, dialogue, summarization, or speech-language work.",
        "f59e0b",
    ),
    (
        "vision",
        "Computer vision, image/video understanding, object detection, segmentation, or visual recognition.",
        "0f766e",
    ),
    (
        "reinforcement-learning",
        "Reinforcement learning, agents, planning, control, robotics, reward modeling, or human feedback.",
        "dc2626",
    ),
    (
        "trustworthy-ai",
        "Explainability, robustness, safety, fairness, uncertainty, privacy, bias, or responsible AI.",
        "be123c",
    ),
    (
        "graph-learning",
        "Graph neural networks, recommender systems, knowledge graphs, graph benchmarks, or graph-based core AI methods.",
        "4f46e5",
    ),
    (
        "ai4science",
        "AI for science, healthcare, biology, molecules, proteins, robotics, autonomous systems, or clinical domains.",
        "16a34a",
    ),
]
KEYWORD_COLORS = {keyword: color for keyword, _, color in KEYWORD_CONVENTION}

LANGUAGES = {
    "en": "English",
    "ko": "한국어",
    "zh": "中文",
    "ja": "日本語",
}

UI_LABELS = {
    "en": {
        "papers": "papers",
        "categories": "categories",
        "overview": "Category Overview",
        "limitations": "Limitations",
        "analysis": "Selected-period analysis",
        "totalSelected": "Total selected papers",
        "categoryCount": "Categories",
        "keyIdea": "Key idea",
        "strengths": "Strengths",
        "paperLimitations": "Limitations",
    },
    "ko": {
        "papers": "편",
        "categories": "개 분류",
        "overview": "분류 개요",
        "limitations": "한계",
        "analysis": "선택 기간 분석",
        "totalSelected": "선정 논문",
        "categoryCount": "분류",
        "keyIdea": "핵심 아이디어",
        "strengths": "장점",
        "paperLimitations": "한계",
    },
    "zh": {
        "papers": "篇论文",
        "categories": "个分类",
        "overview": "分类概览",
        "limitations": "局限性",
        "analysis": "所选期间分析",
        "totalSelected": "入选论文总数",
        "categoryCount": "分类数",
        "keyIdea": "核心思想",
        "strengths": "优势",
        "paperLimitations": "局限性",
    },
    "ja": {
        "papers": "本",
        "categories": "分類",
        "overview": "カテゴリ概要",
        "limitations": "限界",
        "analysis": "選択期間の分析",
        "totalSelected": "選定論文総数",
        "categoryCount": "カテゴリ数",
        "keyIdea": "主要アイデア",
        "strengths": "強み",
        "paperLimitations": "限界",
    },
}

TAXONOMY_TRENDS = {
    "Foundation Models and Large Language Models": [
        "The main trend is a shift from task-specific NLP systems toward general-purpose foundation models that transfer across tasks through prompting, retrieval, and instruction tuning.",
        "Scaling laws, retrieval augmentation, alignment, and data governance are now central design axes rather than afterthoughts.",
        "Citation-ranked work tends to emphasize architectures, pretraining corpora, benchmark behavior, and broad capability evaluation.",
    ],
    "Generative Models and Synthetic Media": [
        "Generative AI research has moved from GAN-centered image synthesis toward diffusion, score-based modeling, and controllable multimodal generation.",
        "The area increasingly connects generation quality with data curation, safety, copyright, controllability, and evaluation reliability.",
        "Highly cited papers often introduce reusable model families, training objectives, or evaluation protocols that become infrastructure for later systems.",
    ],
    "Vision and Multimodal Learning": [
        "Vision research is increasingly organized around transformer backbones, self-supervised pretraining, segmentation/detection foundation models, and vision-language alignment.",
        "The strongest papers often combine large-scale data, reusable architectures, and transfer to multiple downstream tasks.",
        "Multimodal work is pushing vision beyond single-task recognition toward retrieval, grounding, robotics, and generative interfaces.",
    ],
    "Natural Language Processing and Knowledge": [
        "NLP is moving from supervised task pipelines toward pretrained language models, retrieval-augmented methods, and knowledge-intensive reasoning benchmarks.",
        "Search, retrieval, summarization, dialogue, and domain adaptation are increasingly evaluated as integrated knowledge workflows.",
        "Citation-ranked NLP papers often become shared components, datasets, or baseline methods for later foundation model research.",
    ],
    "Reinforcement Learning and Agents": [
        "Reinforcement learning is converging with human feedback, offline datasets, robotics, planning, and agentic tool-use settings.",
        "A major thread is reducing sample inefficiency while improving robustness under distribution shift and sparse rewards.",
        "The field is increasingly judged by transfer, safety, and real-world interaction rather than benchmark score alone.",
    ],
    "Representation, Self-Supervised, and Transfer Learning": [
        "Self-supervised and contrastive methods are reducing dependence on labeled data while improving transfer across tasks and domains.",
        "Reusable representations, distillation, domain adaptation, and few-shot learning form the connective tissue between specialized AI subfields.",
        "Citation impact is often driven by methods that become default pretraining or adaptation recipes.",
    ],
    "Trustworthy, Explainable, and Responsible AI": [
        "Trustworthy AI work is broadening from post-hoc explanations to robustness, fairness, privacy, calibration, uncertainty, and safety-aware evaluation.",
        "Highly cited papers provide taxonomies, metrics, attacks, or toolkits that make system behavior easier to inspect and compare.",
        "The area is increasingly coupled to foundation model deployment, high-stakes domains, and governance requirements.",
    ],
    "Graph Learning, Recommendation, and Core Methods": [
        "Core AI methods include graph neural networks, recommender systems, optimization, neural architecture search, Bayesian methods, and efficient training recipes.",
        "Many papers in this category become reusable algorithmic infrastructure for applied AI systems.",
        "The dominant trend is stronger inductive bias for non-Euclidean data, sparse interaction data, and efficient hyperparameter or architecture search.",
    ],
    "AI for Science, Healthcare, and Robotics": [
        "Applied AI for science and healthcare is shifting from proof-of-concept prediction toward validated workflows for biology, medicine, molecules, and robotics.",
        "The most visible work couples domain data with deep learning architectures that can transfer into laboratory, clinical, or embodied settings.",
        "Evaluation increasingly needs external validation, prospective testing, reproducible datasets, and domain-specific safety constraints.",
    ],
    "General AI Methods and Systems": [
        "General AI methods consolidate architectures, benchmarks, surveys, datasets, and system-level observations that cut across subfields.",
        "This category often captures high-citation survey or infrastructure work that shapes how later papers define progress.",
        "Citation-ranked views can be especially useful here, but they should be read as a map of influence rather than a complete quality assessment.",
    ],
}

TAXONOMY_LIMITATIONS = {
    "Foundation Models and Large Language Models": [
        "Capability gains are difficult to separate from data scale, benchmark leakage, and evaluation prompt sensitivity.",
        "Alignment and safety claims often need stronger real-world and multilingual validation.",
        "Compute-intensive training can limit reproducibility and concentrate follow-up work around a small number of institutions.",
    ],
    "Generative Models and Synthetic Media": [
        "Image and media quality metrics may not capture factuality, controllability, provenance, or downstream harms.",
        "Training data provenance and copyright constraints can be under-specified in highly cited generation work.",
        "Robust evaluation across cultures, modalities, and adversarial uses remains difficult.",
    ],
    "Vision and Multimodal Learning": [
        "Large-scale benchmark success can overstate robustness under distribution shift, rare classes, and real deployment constraints.",
        "Multimodal alignment may inherit biases and spurious correlations from web-scale data.",
        "High-performing systems often require data and compute resources that are hard for smaller labs to reproduce.",
    ],
    "Natural Language Processing and Knowledge": [
        "Benchmark scores can hide brittle reasoning, retrieval failures, hallucination, and domain transfer issues.",
        "Language coverage is often uneven, with English and high-resource domains overrepresented.",
        "Evaluation can be sensitive to annotation protocols, prompt wording, and changing model APIs.",
    ],
    "Reinforcement Learning and Agents": [
        "Sample efficiency, reward misspecification, simulator bias, and safety under exploration remain persistent barriers.",
        "Benchmark performance may not transfer to physical robots, human-facing tools, or open-ended environments.",
        "Agentic systems need stronger evidence on reliability, recovery from errors, and long-horizon oversight.",
    ],
    "Representation, Self-Supervised, and Transfer Learning": [
        "Transfer claims depend heavily on downstream task choice, data overlap, and evaluation protocol.",
        "Contrastive and self-supervised methods can learn spurious shortcuts when augmentations or negatives are poorly matched.",
        "Representation quality is hard to compare when model size, data scale, and training recipes differ.",
    ],
    "Trustworthy, Explainable, and Responsible AI": [
        "Explanations can be persuasive without being faithful to model internals or decision processes.",
        "Fairness, robustness, privacy, and safety metrics can conflict and require domain-specific tradeoffs.",
        "Responsible AI results often need stronger deployment evidence beyond benchmark or synthetic settings.",
    ],
    "Graph Learning, Recommendation, and Core Methods": [
        "Graph and recommender benchmarks can contain temporal leakage, popularity bias, or unrealistic train/test splits.",
        "Algorithmic gains may be sensitive to hyperparameter budgets and implementation details.",
        "Core methods need careful ablations before broad claims about generality or efficiency are accepted.",
    ],
    "AI for Science, Healthcare, and Robotics": [
        "External validation, prospective testing, and domain expert review are often more important than retrospective benchmark scores.",
        "Clinical, biological, or robotic deployment can fail when data collection protocols differ from training assumptions.",
        "Safety, interpretability, uncertainty, and regulatory evidence remain essential for translation.",
    ],
    "General AI Methods and Systems": [
        "Survey and infrastructure papers can dominate citations while empirical evidence remains distributed across subfields.",
        "Broad claims need careful mapping to specific tasks, datasets, and operational constraints.",
        "Metadata-driven ranking cannot replace expert reading of full papers and experimental details.",
    ],
}

KOREAN_CATEGORY_NAMES = {
    "Foundation Models and Large Language Models": "파운데이션 모델 및 대규모 언어모델",
    "Generative Models and Synthetic Media": "생성 모델 및 합성 미디어",
    "Vision and Multimodal Learning": "비전 및 멀티모달 학습",
    "Natural Language Processing and Knowledge": "자연어처리 및 지식",
    "Reinforcement Learning and Agents": "강화학습 및 에이전트",
    "Representation, Self-Supervised, and Transfer Learning": "표현학습, 자기지도학습 및 전이학습",
    "Trustworthy, Explainable, and Responsible AI": "신뢰가능하고 설명가능한 책임 AI",
    "Graph Learning, Recommendation, and Core Methods": "그래프 학습, 추천 및 핵심 방법론",
    "AI for Science, Healthcare, and Robotics": "과학, 헬스케어 및 로보틱스를 위한 AI",
    "General AI Methods and Systems": "일반 AI 방법론 및 시스템",
}


def norm_text(value):
    return re.sub(r"\s+", " ", value or "").strip()


def safe_slug(value):
    value = re.sub(r"[^a-zA-Z0-9]+", "-", value.lower()).strip("-")
    return value[:90] or "paper"


def normalize_title(value):
    return re.sub(r"[^a-z0-9]+", " ", (value or "").lower()).strip()


def paper_key(paper):
    ext = paper.get("externalIds") or {}
    for key in ("DOI", "ArXiv", "PubMed", "CorpusId"):
        value = ext.get(key)
        if value:
            return f"{key.lower()}:{str(value).lower()}"
    if paper.get("paperId"):
        return f"s2:{paper['paperId']}"
    return f"title:{normalize_title(paper.get('title'))}"


def cache_name(year, query):
    digest = hashlib.sha1(f"{year}:{query}".encode("utf-8")).hexdigest()[:12]
    return CACHE_DIR / f"s2_{year}_{safe_slug(query)}_{digest}.json"


def fetch_year_query(year, query):
    CACHE_DIR.mkdir(parents=True, exist_ok=True)
    cache_file = cache_name(year, query)
    if cache_file.exists():
        return json.loads(cache_file.read_text(encoding="utf-8"))

    params = {
        "query": query,
        "year": str(year),
        "fields": S2_FIELDS,
        "sort": "citationCount:desc",
        "limit": str(CANDIDATES_PER_YEAR),
    }
    last_error = None
    for attempt in range(4):
        try:
            response = requests.get(S2_BULK_URL, params=params, timeout=90)
            if response.status_code == 429:
                time.sleep(15 + attempt * 10)
                continue
            response.raise_for_status()
            payload = response.json()
            rows = payload.get("data") or []
            cache_file.write_text(json.dumps(rows, ensure_ascii=False), encoding="utf-8")
            return rows
        except Exception as exc:
            last_error = exc
            time.sleep(5 + attempt * 5)
    raise RuntimeError(f"Semantic Scholar query failed for {year} {query!r}: {last_error}")


def venue_name(paper):
    publication_venue = paper.get("publicationVenue") or {}
    return norm_text(publication_venue.get("name") or paper.get("venue") or "")


def authors_text(paper, max_authors=8):
    authors = paper.get("authors") or []
    names = [norm_text(a.get("name")) for a in authors if norm_text(a.get("name"))]
    if len(names) > max_authors:
        return ", ".join(names[:max_authors]) + ", et al."
    return ", ".join(names)


def primary_link(paper):
    ext = paper.get("externalIds") or {}
    doi = ext.get("DOI")
    if doi:
        doi = str(doi)
        return doi if doi.startswith("http") else f"https://doi.org/{doi}"
    if paper.get("url"):
        return paper["url"]
    pdf = paper.get("openAccessPdf") or {}
    return pdf.get("url") or ""


def title_is_bad(title):
    title_l = normalize_title(title)
    if len(title_l) < 8:
        return True
    return any(re.search(pattern, title_l, re.I) for pattern in BAD_TITLE_PATTERNS)


def relevance_counts(paper):
    title = paper.get("title") or ""
    abstract = paper.get("abstract") or ""
    venue = venue_name(paper)
    text = f"{title} {abstract} {venue}"
    title_patterns = globals().get("TITLE_RELEVANCE_PATTERNS", AI_RELEVANCE_PATTERNS)
    title_hits = sum(1 for pattern in title_patterns if re.search(pattern, title, re.I))
    text_hits = sum(1 for pattern in AI_RELEVANCE_PATTERNS if re.search(pattern, text, re.I))
    return title_hits, text_hits


def is_relevant(paper):
    title = norm_text(paper.get("title"))
    if not title or title_is_bad(title):
        return False
    publication_types = [str(x).lower() for x in (paper.get("publicationTypes") or [])]
    if any(kind in publication_types for kind in ("editorial", "news", "lettersandcomments")):
        return False
    title_hits, text_hits = relevance_counts(paper)
    fields = " ".join(
        norm_text(item.get("category"))
        for item in (paper.get("s2FieldsOfStudy") or [])
        if item.get("category")
    ).lower()
    if title_hits:
        return True
    if text_hits >= 2:
        return True
    return text_hits >= 1 and "computer science" in fields and "review" in title.lower()


def importance_score(paper):
    title = paper.get("title") or ""
    abstract = paper.get("abstract") or ""
    venue = venue_name(paper)
    text = f"{title} {abstract} {venue}".lower()
    citations = int(paper.get("citationCount") or 0)
    influential = int(paper.get("influentialCitationCount") or 0)
    title_hits, text_hits = relevance_counts(paper)
    query_hits = len(paper.get("sourceQueries") or [])
    score = math.log1p(citations) * 22.0 + math.log1p(influential) * 14.0
    reasons = [f"citations={citations}", f"influential={influential}"]
    if any(v in venue.lower() for v in IMPORTANT_VENUES):
        score += 10
        reasons.append("recognized venue")
    if re.search(r"\b(review|survey|systematic review|benchmark|dataset)\b", text):
        score += 5
        reasons.append("survey/benchmark signal")
    if re.search(r"\b(open-source|code|dataset|benchmark|corpus)\b", text):
        score += 3
        reasons.append("resource signal")
    if re.search(r"\b(safety|fairness|privacy|robust|explainable|interpretability|clinical|healthcare)\b", text):
        score += 3
        reasons.append("translation/trustworthiness signal")
    if paper.get("openAccessPdf"):
        score += 1
        reasons.append("open PDF metadata")
    if title_hits:
        score += min(8, title_hits * 2)
        reasons.append(f"title AI matches={title_hits}")
    if text_hits:
        score += min(8, text_hits)
        reasons.append(f"AI term matches={text_hits}")
    if query_hits > 1:
        score += min(6, query_hits)
        reasons.append(f"query hits={query_hits}")
    return round(score, 3), "; ".join(reasons)


def text_for(paper):
    return f"{paper.get('title', '')} {paper.get('abstract', '')} {paper.get('venue', '')} {paper.get('methodTags', '')}".lower()


def category_for(paper):
    text = text_for(paper)
    if re.search(r"\b(large language models?|language models?|llm|foundation model|gpt|chatgpt|prompting?|instruction[- ]?tuning|scaling laws?|retrieval-augmented generation|rag)\b", text, re.I):
        return "Foundation Models and Large Language Models"
    if re.search(r"\b(protein|molecule|drug|genomic|healthcare|medical|clinical|biomedical|biology|autonomous driving)\b", text, re.I):
        return "AI for Science, Healthcare, and Robotics"
    if re.search(r"\b(graph neural|graph convolution|open graph benchmark|knowledge graph|recommender|recommendation)\b", text, re.I):
        return "Graph Learning, Recommendation, and Core Methods"
    best = ("General AI Methods and Systems", 0)
    for category, terms in CATEGORIES:
        score = sum(1 for term in terms if term in text)
        if score > best[1]:
            best = (category, score)
    return best[0]


def method_tags(paper):
    text = text_for(paper)
    rules = [
        ("survey/review", r"\b(review|survey|overview)\b"),
        ("benchmark/dataset", r"\b(benchmark|dataset|corpus|leaderboard|open graph)\b"),
        ("transformer", r"\b(transformer|attention|bert|gpt|vit)\b"),
        ("self-supervised", r"\b(self-supervised|contrastive|masked|pretrain|pre-training)\b"),
        ("generative", r"\b(diffusion|generative|gan|vae|score-based|text-to-image)\b"),
        ("graph-learning", r"\b(graph neural|gnn|knowledge graph)\b"),
        ("retrieval", r"\b(retrieval|rag|dense passage|information retrieval)\b"),
        ("trustworthy-ai", r"\b(explainable|interpretability|fairness|robust|privacy|safety|uncertainty)\b"),
        ("ai4science", r"\b(protein|molecule|drug|clinical|medical|healthcare|robot|biology)\b"),
        ("reinforcement-learning", r"\b(reinforcement|policy|agent|reward|human feedback)\b"),
    ]
    tags = [name for name, pattern in rules if re.search(pattern, text, re.I)]
    return tags[:6] or ["metadata-ranked"]


def keyword_tags(paper, category=None):
    text = text_for(paper)
    tags = set()
    if re.search(r"\b(machine learning|statistical learning|neural network|support vector|svm|decision tree|random forest|xgboost|boosting|classification|clustering|data mining|pattern recognition|bayesian network)\b", text, re.I):
        tags.add("machine-learning")
    if re.search(r"\b(language model|large language|llm|foundation model|gpt|bert|prompt|rag|retrieval-augmented|scaling law)\b", text, re.I):
        tags.add("foundation-models")
    if re.search(r"\b(diffusion|generative|gan|vae|synthetic|text-to-image|image synthesis)\b", text, re.I):
        tags.add("generative-ai")
    if re.search(r"\b(multimodal|multi-modal|vision-language|clip|cross-modal|video-language)\b", text, re.I):
        tags.add("multimodal")
    if re.search(r"\b(natural language|nlp|text|translation|question answering|dialogue|summarization|speech recognition|information retrieval|language model)\b", text, re.I):
        tags.add("nlp")
    if re.search(r"\b(computer vision|image|video|object detection|segmentation|visual|vision transformer|point cloud)\b", text, re.I):
        tags.add("vision")
    if re.search(r"\b(reinforcement|policy|agent|robot|control|planning|reward|human feedback|imitation)\b", text, re.I):
        tags.add("reinforcement-learning")
    if re.search(r"\b(explainable|interpretability|fairness|bias|robust|adversarial|privacy|safety|uncertainty|calibration|responsible)\b", text, re.I):
        tags.add("trustworthy-ai")
    if re.search(r"\b(graph neural|graph convolution|open graph benchmark|knowledge graph|recommender|recommendation|graph machine learning)\b", text, re.I):
        tags.add("graph-learning")
    if re.search(r"\b(protein|molecule|drug|genomic|healthcare|medical|clinical|biomedical|science|robotics|autonomous driving|biology)\b", text, re.I):
        tags.add("ai4science")

    category = category or category_for(paper)
    defaults = {
        "Foundation Models and Large Language Models": "foundation-models",
        "Generative Models and Synthetic Media": "generative-ai",
        "Vision and Multimodal Learning": "vision",
        "Natural Language Processing and Knowledge": "nlp",
        "Reinforcement Learning and Agents": "reinforcement-learning",
        "Trustworthy, Explainable, and Responsible AI": "trustworthy-ai",
        "Graph Learning, Recommendation, and Core Methods": "graph-learning",
        "AI for Science, Healthcare, and Robotics": "ai4science",
        "General AI Methods and Systems": "machine-learning",
    }
    if not tags and category in defaults:
        tags.add(defaults[category])
    ordered = [keyword for keyword, _, _ in KEYWORD_CONVENTION if keyword in tags]
    return ordered or ["machine-learning"]


def first_sentence(value, fallback):
    text = norm_text(value)
    if not text:
        return fallback
    match = re.search(r"(.{40,280}?[.!?])\s", text + " ")
    return match.group(1) if match else text[:280]


def research_limitations(paper, method_tag_list, keyword_list):
    category = paper.get("category") or category_for(paper)
    items = list(TAXONOMY_LIMITATIONS.get(category, next(iter(TAXONOMY_LIMITATIONS.values()))))
    tags = set(method_tag_list + keyword_list)
    if "survey/review" in tags:
        items.append("Review-level synthesis cannot resolve inconsistent study quality, benchmark leakage, or reproducibility gaps.")
    if "benchmark/dataset" in tags:
        items.append("Benchmark value depends on sustained maintenance, clear licensing, and representative task design.")
    if "trustworthy-ai" in tags:
        items.append("Trustworthiness metrics require domain-specific thresholds and may trade off against accuracy or utility.")
    if "ai4science" in tags:
        items.append("Domain translation needs external validation under laboratory, clinical, or embodied deployment conditions.")
    if "foundation-models" in tags:
        items.append("Large-model conclusions can be sensitive to data scale, prompt design, and hidden training-set overlap.")
    return items[:3]


def localized_paper_fields(paper, language):
    if language == "en":
        return {
            "keyIdea": paper["keyIdea"],
            "strengths": paper["strengths"],
            "limitations": paper["limitations"],
        }
    category_name = KOREAN_CATEGORY_NAMES.get(paper["category"], paper["category"])
    tags = paper.get("methodTags", "").replace("; ", ", ") or "AI 방법론"
    return {
        "keyIdea": f"'{paper['title']}'은(는) {category_name} 분야에서 {tags} 흐름을 보여주는 citation-ranked AI 연구입니다.",
        "strengths": paper["strengths"]
        .replace("high citation signal", "높은 인용 신호")
        .replace("influential citation signal", "영향력 있는 인용 신호")
        .replace("recognized venue", "주요 학술지/학회 신호")
        .replace("open-access PDF metadata", "오픈액세스 PDF 메타데이터")
        .replace("selected by citation count from the audited AI candidate pool", "감사 가능한 AI 후보군에서 인용수 기준으로 선정"),
        "limitations": " ".join(TAXONOMY_LIMITATIONS.get(paper["category"], next(iter(TAXONOMY_LIMITATIONS.values())))[:2]),
    }


def enrich_paper(paper):
    base = dict(paper)
    category = category_for(base)
    methods = method_tags(base)
    keywords = keyword_tags(base, category)
    citations = int(base.get("citationCount") or 0)
    influential = int(base.get("influentialCitationCount") or 0)
    strengths = []
    if citations >= 100:
        strengths.append(f"high citation signal ({citations:,})")
    if influential >= 10:
        strengths.append(f"influential citation signal ({influential:,})")
    if "recognized venue" in base.get("importanceReasons", ""):
        strengths.append("recognized venue")
    if base.get("openAccessPdf"):
        strengths.append("open-access PDF metadata")
    if not strengths:
        strengths.append("selected by citation count from the audited AI candidate pool")
    base["category"] = category
    base["methodTags"] = "; ".join(methods)
    base["keywordTags"] = "; ".join(keywords)
    base["keyIdea"] = first_sentence(
        base.get("abstract", ""),
        f"Positions {base.get('title', 'this paper')} within {category}.",
    )
    base["strengths"] = "; ".join(strengths[:4])
    base["limitations"] = "; ".join(research_limitations(base, methods, keywords))
    return base


def normalize_paper(paper, year, candidate_rank):
    ext = paper.get("externalIds") or {}
    score, reasons = importance_score(paper)
    pdf = paper.get("openAccessPdf") or {}
    row = {
        "paperId": paper.get("paperId") or "",
        "title": norm_text(paper.get("title")),
        "authors": authors_text(paper),
        "year": int(paper.get("year") or year),
        "venue": venue_name(paper),
        "publicationDate": paper.get("publicationDate") or "",
        "citationCount": int(paper.get("citationCount") or 0),
        "influentialCitationCount": int(paper.get("influentialCitationCount") or 0),
        "abstract": norm_text(paper.get("abstract")),
        "url": primary_link(paper),
        "semanticScholarUrl": paper.get("url") or "",
        "openAccessPdf": pdf.get("url") or "",
        "doi": ext.get("DOI") or "",
        "arxiv": ext.get("ArXiv") or "",
        "pubmed": ext.get("PubMed") or "",
        "corpusId": ext.get("CorpusId") or "",
        "fieldsOfStudy": "; ".join(
            sorted(
                {
                    norm_text(item.get("category"))
                    for item in (paper.get("s2FieldsOfStudy") or [])
                    if norm_text(item.get("category"))
                }
            )
        ),
        "publicationTypes": "; ".join(paper.get("publicationTypes") or []),
        "sourceQueries": "; ".join(sorted(paper.get("sourceQueries") or [])),
        "queryHitCount": len(paper.get("sourceQueries") or []),
        "candidateRank": candidate_rank,
        "importanceScore": score,
        "importanceReasons": reasons,
    }
    return enrich_paper(row)


def collect_papers():
    candidates_by_year = {}
    for year in YEARS:
        merged = {}
        for query in QUERIES:
            print(f"[collect] {year} :: {query}", flush=True)
            try:
                papers = fetch_year_query(year, query)
            except Exception as exc:
                print(f"[warn] {year} {query}: {exc}", flush=True)
                continue
            for paper in papers:
                if paper.get("year") != year:
                    continue
                if not is_relevant(paper):
                    continue
                key = paper_key(paper)
                if key not in merged:
                    paper["sourceQueries"] = set()
                    merged[key] = paper
                merged[key]["sourceQueries"].add(query)
                if int(paper.get("citationCount") or 0) > int(merged[key].get("citationCount") or 0):
                    merged[key].update(paper)
            time.sleep(REQUEST_DELAY)

        ranked = sorted(
            merged.values(),
            key=lambda p: (
                int(p.get("citationCount") or 0),
                int(p.get("influentialCitationCount") or 0),
                importance_score(p)[0],
                p.get("title") or "",
            ),
            reverse=True,
        )
        candidate_pool = ranked[:CANDIDATES_PER_YEAR]
        normalized = [normalize_paper(p, year, i + 1) for i, p in enumerate(candidate_pool)]
        candidates_by_year[year] = normalized
        print(
            f"[collect] {year}: retained {len(normalized):,}/{CANDIDATES_PER_YEAR:,} candidates from {len(ranked):,} relevant records",
            flush=True,
        )

    selected, selected_by_year = select_papers_by_year(candidates_by_year)
    all_candidates = [p for rows in candidates_by_year.values() for p in rows]
    print(f"[collect] selected {len(selected):,} papers ({TARGET_PER_YEAR:,} per year where available)", flush=True)
    return selected, selected_by_year, candidates_by_year


def select_papers_by_year(candidates_by_year):
    selected_by_year = {year: [] for year in YEARS}
    selected = []
    for year in YEARS:
        rows = sorted(
            [paper for paper in candidates_by_year.get(year, []) if is_relevant(paper)],
            key=lambda p: int(p.get("candidateRank") or 999999),
        )[:TARGET_PER_YEAR]
        for year_rank, paper in enumerate(rows, 1):
            paper["yearRank"] = year_rank
            selected_by_year[year].append(paper)
            selected.append(paper)
    selected.sort(
        key=lambda p: (
            p["citationCount"],
            p["influentialCitationCount"],
            p["importanceScore"],
            p["title"],
        ),
        reverse=True,
    )
    for rank, paper in enumerate(selected, 1):
        paper["rank"] = rank
    for rows in selected_by_year.values():
        rows.sort(key=lambda p: p.get("yearRank") or 999999)
    return selected, selected_by_year


def title_key(value):
    text = norm_text(value).lower()
    text = re.sub(r"[^a-z0-9]+", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def clean_github_url(value):
    text = norm_text(value)
    match = re.search(r"https?://github\.com/[^\s<>\"')]+", text, flags=re.I)
    if not match:
        return ""
    path = match.group(0).split("github.com/", 1)[1].rstrip(".,;")
    return f"https://github.com/{path}"


def is_official_github_entry(value):
    return isinstance(value, dict) and bool(value.get("githubOfficial")) and bool(value.get("mentionedInPaper"))


def official_github_url_from_paper_text(row):
    official_context = re.compile(
        r"\b(code|codes|source|implementation|repo|repository|software|toolbox|package|model|models|weights|data|dataset|project|available|released|open-source|publicly)\b",
        flags=re.I,
    )
    for field in ("abstract", "abstractSnippet"):
        text = norm_text(row.get(field))
        for match in re.finditer(r"https?://github\.com/[^\s<>\"')]+", text, flags=re.I):
            window = text[max(0, match.start() - 160): match.end() + 80]
            if official_context.search(window):
                return clean_github_url(match.group(0))
    return ""


def load_github_links():
    path = DATA_DIR / GITHUB_LINKS_JSON
    if not path.exists():
        return {}
    payload = json.loads(path.read_text(encoding="utf-8"))
    links = {}
    for key, value in payload.get("links", {}).items():
        if not is_official_github_entry(value):
            continue
        github_url = clean_github_url(value.get("githubUrl") if isinstance(value, dict) else value)
        if github_url:
            links[key] = {**value, "githubUrl": github_url} if isinstance(value, dict) else {"githubUrl": github_url}
    return links


def sync_github_keyword_tag(row):
    raw = str(row.get("keywordTags") or "")
    tags = [tag.strip() for tag in raw.split(";") if tag.strip() and tag.strip() != "github"]
    if row.get("githubUrl"):
        tags.append("github")
    row["keywordTags"] = "; ".join(tags)


def apply_github_links(rows):
    links = load_github_links()
    for row in rows:
        match = links.get(title_key(row.get("title")))
        if match and is_official_github_entry(match):
            row["githubUrl"] = clean_github_url(match.get("githubUrl", ""))
        else:
            row["githubUrl"] = ""
        sync_github_keyword_tag(row)
    return rows


def reuse_existing_candidates():
    rows = []
    path = DATA_DIR / CANDIDATES_JSON
    if path.exists():
        payload = json.loads(path.read_text(encoding="utf-8"))
        for row in payload.get("candidates") or []:
            rows.append(normalize_candidate_row(row))
    if not rows:
        for year in YEARS:
            year_path = DATA_DIR / f"candidates_top1000_{year}.json"
            if not year_path.exists():
                raise FileNotFoundError(f"Missing existing candidate pool: {year_path}")
            payload = json.loads(year_path.read_text(encoding="utf-8"))
            for row in payload.get("candidates") or []:
                rows.append(normalize_candidate_row(row))
    candidates_by_year = {}
    for year in YEARS:
        year_rows = [p for p in rows if p["year"] == year]
        year_rows.sort(key=lambda p: int(p.get("candidateRank") or 999999))
        candidates_by_year[year] = year_rows[:CANDIDATES_PER_YEAR]
    all_candidates = [p for values in candidates_by_year.values() for p in values]
    selected, selected_by_year = select_papers_by_year(candidates_by_year)
    print(f"[reuse] regenerated {len(selected):,} selected papers ({TARGET_PER_YEAR:,} per year) from {len(all_candidates):,} existing candidates", flush=True)
    return selected, selected_by_year, candidates_by_year


def normalize_candidate_row(row):
    normalized = dict(row)
    normalized["year"] = int(normalized.get("year") or 0)
    normalized["citationCount"] = int(normalized.get("citationCount") or 0)
    normalized["influentialCitationCount"] = int(normalized.get("influentialCitationCount") or 0)
    normalized["importanceScore"] = float(normalized.get("importanceScore") or 0)
    return enrich_paper(normalized)


def csv_ready(row):
    out = {}
    for key, value in row.items():
        if isinstance(value, (list, dict, set)):
            out[key] = json.dumps(value, ensure_ascii=False)
        else:
            out[key] = value
    return out


def write_csv(path, rows, fields):
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fields, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(csv_ready(row) for row in rows)


def write_data(selected, selected_by_year, candidates_by_year):
    DATA_DIR.mkdir(exist_ok=True)
    apply_github_links(selected)
    for rows in selected_by_year.values():
        apply_github_links(rows)
    flat_candidates = [p for rows in candidates_by_year.values() for p in rows]
    metadata = {
        "topic": "AI research",
        "source": "Semantic Scholar Academic Graph bulk search",
        "generated": date.today().isoformat(),
        "years": YEARS,
        "candidate_pool_per_year": CANDIDATES_PER_YEAR,
        "selected_per_year": TARGET_PER_YEAR,
        "selected_total": TARGET_TOTAL,
        "selection_scope": "top 100 papers per publication year by citation count",
        "ranking": "citationCount desc, influentialCitationCount desc",
        "queries": QUERIES,
        "openalex_concept": "https://openalex.org/C154945302",
        "openalex_subfield": "https://openalex.org/subfields/1702",
    }
    (DATA_DIR / PAPERS_JSON).write_text(
        json.dumps({"metadata": metadata, "papers": selected, "byYear": selected_by_year}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    (DATA_DIR / CANDIDATES_JSON).write_text(
        json.dumps(
            {
                "metadata": metadata,
                "candidate_count": len(flat_candidates),
                "combined_csv": CANDIDATES_CSV,
                "year_files": {str(year): f"candidates_top1000_{year}.json" for year in YEARS},
                "note": "Candidate records are stored in the combined CSV and per-year JSON files to keep GitHub file sizes below the 100 MB single-file limit.",
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )
    fields = [
        "rank",
        "yearRank",
        "candidateRank",
        "year",
        "title",
        "authors",
        "venue",
        "publicationDate",
        "citationCount",
        "influentialCitationCount",
        "importanceScore",
        "category",
        "methodTags",
        "keywordTags",
        "keyIdea",
        "strengths",
        "limitations",
        "url",
        "semanticScholarUrl",
        "openAccessPdf",
        "githubUrl",
        "doi",
        "arxiv",
        "pubmed",
        "corpusId",
        "fieldsOfStudy",
        "publicationTypes",
        "sourceQueries",
        "queryHitCount",
        "importanceReasons",
        "abstract",
    ]
    candidate_fields = [field for field in fields if field not in ("rank", "yearRank")]
    write_csv(DATA_DIR / PAPERS_CSV, selected, fields)
    write_csv(DATA_DIR / CANDIDATES_CSV, flat_candidates, candidate_fields)
    write_paper_localizations(DATA_DIR / PAPER_LOCALIZATIONS_JSON, selected)
    for year, rows in candidates_by_year.items():
        stem = f"candidates_top1000_{year}"
        (DATA_DIR / f"{stem}.json").write_text(
            json.dumps({"metadata": metadata, "year": year, "candidates": rows}, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        write_csv(DATA_DIR / f"{stem}.csv", rows, candidate_fields)
    for year, rows in selected_by_year.items():
        write_csv(DATA_DIR / f"papers_{year}.csv", rows, fields)
    return flat_candidates


def category_groups(rows):
    groups = defaultdict(list)
    for paper in rows:
        groups[paper["category"]].append(paper)
    for papers in groups.values():
        papers.sort(key=lambda p: (p["citationCount"], p["influentialCitationCount"], p["importanceScore"]), reverse=True)
    return groups


def category_stats(rows):
    return Counter(p["category"] for p in rows)


def year_stats(rows):
    stats = {}
    for year in YEARS:
        year_rows = [p for p in rows if p["year"] == year]
        if not year_rows:
            continue
        stats[year] = {
            "count": len(year_rows),
            "citations": sum(p["citationCount"] for p in year_rows),
            "top": max(year_rows, key=lambda p: (p["citationCount"], p["influentialCitationCount"])),
        }
    return stats


def period_key(start, end):
    return f"{start}-{end}"


def period_label(start, end):
    if start == START_YEAR and end == END_YEAR:
        return f"All years ({YEAR_RANGE_TEXT})"
    return str(start) if start == end else f"{start}-{end}"


def all_period_ranges():
    return [(start, end) for start in YEARS for end in range(start, END_YEAR + 1)]


def period_select_ranges():
    full = (START_YEAR, END_YEAR)
    return [full] + [pair for pair in all_period_ranges() if pair != full]


def research_overview_html():
    return """
    <section class="research-brief" id="researchBrief" aria-labelledby="research-timeline-title">
      <h2 id="research-timeline-title">Research Timeline</h2>
      <div class="timeline-copy">
        <p>2000-2026년 AI 연구 코퍼스는 통계적 머신러닝과 커널·그래프·추천 방법론에서 출발해 deep learning, vision, representation learning, transformer 기반 NLP, foundation model, generative AI, multimodal learning, reinforcement learning과 agents, trustworthy AI, AI for Science로 확장된 흐름을 보여준다. 총 2,700편의 선별 논문과 856만 회 이상의 인용 신호는 AI가 개별 태스크 알고리즘의 집합에서 범용 representation과 대규모 사전학습 인프라 중심의 연구 생태계로 이동했음을 드러낸다.</p>
        <p>큰 taxonomy 축은 General AI Methods and Systems, Foundation Models and Large Language Models, Vision and Multimodal Learning이며, 이 세 축이 나머지 생성 모델, NLP/knowledge, RL/agents, graph learning, trustworthy AI, AI4Science를 연결한다. 2010년대 중반 이후 ResNet, attention/transformer, BERT류 사전학습, diffusion과 multimodal model이 인용 질량을 끌어올렸고, 2020년대에는 모델 규모보다 평가, alignment, 데이터 출처, 효율성, 배포 책임성이 더 중요한 연구 질문으로 부상했다.</p>
      </div>
      <h2>Research Insights</h2>
      <div class="research-insights">
        <article class="insight-box">
          <div class="insight-label">Foundation Layer</div>
          <h3>파운데이션 모델이 공통 연구 층이 된다</h3>
          <p>LLM과 대규모 사전학습은 NLP 내부의 한 흐름을 넘어 retrieval, code, multimodal, robotics, science 응용이 공유하는 기본 인프라가 되었다.</p>
          <p class="insight-implication">시사점: 새 연구는 모델 자체뿐 아니라 데이터, instruction, evaluation, deployment interface를 함께 설계해야 한다.</p>
        </article>
        <article class="insight-box">
          <div class="insight-label">Vision And Multimodal</div>
          <h3>비전과 멀티모달은 AI 영향력의 핵심 증폭기다</h3>
          <p>ImageNet, ResNet, vision transformer, CLIP 계열 흐름은 시각 인식에서 언어·영상·로봇·의료 데이터까지 연결되는 표현학습 기반을 만들었다.</p>
          <p class="insight-implication">시사점: 단일 modality benchmark보다 cross-modal grounding과 실제 데이터 변화에 대한 강건성이 중요하다.</p>
        </article>
        <article class="insight-box">
          <div class="insight-label">Generative AI</div>
          <h3>생성 모델은 합성 미디어에서 설계 도구로 이동한다</h3>
          <p>GAN, VAE, diffusion, text-to-image 흐름은 이미지 생성에 머물지 않고 데이터 증강, 시뮬레이션, 과학·의료·콘텐츠 제작 워크플로로 확장된다.</p>
          <p class="insight-implication">시사점: 생성 품질과 함께 provenance, controllability, misuse risk를 같이 평가해야 한다.</p>
        </article>
        <article class="insight-box">
          <div class="insight-label">Agents And RL</div>
          <h3>RL과 agents는 벤치마크에서 도구 사용으로 이동한다</h3>
          <p>강화학습, planning, control, human feedback은 게임·로봇 문제를 넘어 LLM agent, tool use, embodied decision-making과 연결된다.</p>
          <p class="insight-implication">시사점: agent 연구는 reward 설계보다 장기 안정성, 관찰 가능성, 실패 복구, human oversight가 핵심이다.</p>
        </article>
        <article class="insight-box">
          <div class="insight-label">Trustworthy AI</div>
          <h3>신뢰성은 배포의 병목으로 남는다</h3>
          <p>explainability, robustness, fairness, privacy, safety 연구는 성능 중심 논문 지형에 비해 규모는 작지만 실제 적용의 승인 조건을 결정한다.</p>
          <p class="insight-implication">시사점: 높은 benchmark score는 calibration, uncertainty, auditability 없이는 충분한 배포 근거가 되지 않는다.</p>
        </article>
        <article class="insight-box">
          <div class="insight-label">AI4Science</div>
          <h3>AI4Science는 모델의 외부 타당성을 시험한다</h3>
          <p>단백질, 분자, 의료, 로보틱스, 자율주행 연구는 AI 모델이 텍스트와 이미지 밖의 물리·생물학적 제약을 얼마나 잘 다루는지 검증한다.</p>
          <p class="insight-implication">시사점: 향후 영향력은 leaderboard보다 실험 검증, 도메인 제약, 실패 비용을 반영하는 평가에서 갈린다.</p>
        </article>
      </div>
    </section>
"""


def research_copy():
    return {
        "en": """
      <h2 id="research-timeline-title">Research Timeline</h2>
      <div class="timeline-copy">
        <p>The 2000-2026 AI corpus traces a shift from statistical machine learning, kernels, graph and recommendation methods toward deep learning, vision, representation learning, transformer-based NLP, foundation models, generative AI, multimodal learning, reinforcement learning and agents, trustworthy AI, and AI for Science. The 2,700 selected papers and more than 8.56 million citations show AI moving from task-specific algorithms toward general representations and large-scale pretraining infrastructure.</p>
        <p>The largest taxonomy axes are General AI Methods and Systems, Foundation Models and Large Language Models, and Vision and Multimodal Learning. These connect generative models, NLP and knowledge, RL and agents, graph learning, trustworthy AI, and AI4Science. Since the mid-2010s, ResNet, attention/transformers, BERT-style pretraining, diffusion, and multimodal models have driven citation mass; in the 2020s, evaluation, alignment, data provenance, efficiency, and responsible deployment have become central questions.</p>
      </div>
      <h2>Research Insights</h2>
      <div class="research-insights">
        <article class="insight-box"><div class="insight-label">Foundation Layer</div><h3>Foundation models become a shared research layer</h3><p>LLMs and large-scale pretraining now support retrieval, code, multimodal, robotics, and science applications beyond NLP alone.</p><p class="insight-implication">Implication: research must design models together with data, instructions, evaluation, and deployment interfaces.</p></article>
        <article class="insight-box"><div class="insight-label">Vision And Multimodal</div><h3>Vision and multimodal learning amplify AI impact</h3><p>ImageNet, ResNet, vision transformers, and CLIP-like models link visual recognition to language, video, robotics, and medical data.</p><p class="insight-implication">Implication: cross-modal grounding and robustness to real data shifts matter more than single-modality benchmarks.</p></article>
        <article class="insight-box"><div class="insight-label">Generative AI</div><h3>Generative models move from media synthesis to design tools</h3><p>GANs, VAEs, diffusion, and text-to-image systems extend into augmentation, simulation, science, medicine, and content workflows.</p><p class="insight-implication">Implication: provenance, controllability, and misuse risk should be evaluated with output quality.</p></article>
        <article class="insight-box"><div class="insight-label">Agents And RL</div><h3>RL and agents move from benchmarks to tool use</h3><p>Reinforcement learning, planning, control, and human feedback now connect games and robotics to LLM agents and embodied decision-making.</p><p class="insight-implication">Implication: long-term stability, observability, recovery from failure, and human oversight are key.</p></article>
        <article class="insight-box"><div class="insight-label">Trustworthy AI</div><h3>Trustworthiness remains the deployment bottleneck</h3><p>Explainability, robustness, fairness, privacy, and safety are smaller in count but decisive for real-world approval.</p><p class="insight-implication">Implication: benchmark scores are not enough without calibration, uncertainty, and auditability.</p></article>
        <article class="insight-box"><div class="insight-label">AI4Science</div><h3>AI4Science tests external validity</h3><p>Protein, molecule, medical, robotics, and autonomous-systems work tests whether AI handles physical and biological constraints beyond text and images.</p><p class="insight-implication">Implication: future impact depends on experimental validation, domain constraints, and failure costs.</p></article>
      </div>
""",
        "ko": """
      <h2 id="research-timeline-title">연구 타임라인</h2>
      <div class="timeline-copy">
        <p>2000-2026년 AI 연구 코퍼스는 통계적 머신러닝과 커널·그래프·추천 방법론에서 출발해 deep learning, vision, representation learning, transformer 기반 NLP, foundation model, generative AI, multimodal learning, reinforcement learning과 agents, trustworthy AI, AI for Science로 확장된 흐름을 보여준다. 총 2,700편의 선별 논문과 856만 회 이상의 인용 신호는 AI가 개별 태스크 알고리즘의 집합에서 범용 representation과 대규모 사전학습 인프라 중심의 연구 생태계로 이동했음을 드러낸다.</p>
        <p>큰 taxonomy 축은 General AI Methods and Systems, Foundation Models and Large Language Models, Vision and Multimodal Learning이며, 이 세 축이 나머지 생성 모델, NLP/knowledge, RL/agents, graph learning, trustworthy AI, AI4Science를 연결한다. 2010년대 중반 이후 ResNet, attention/transformer, BERT류 사전학습, diffusion과 multimodal model이 인용 질량을 끌어올렸고, 2020년대에는 모델 규모보다 평가, alignment, 데이터 출처, 효율성, 배포 책임성이 더 중요한 연구 질문으로 부상했다.</p>
      </div>
      <h2>연구 인사이트</h2>
      <div class="research-insights">
        <article class="insight-box"><div class="insight-label">Foundation Layer</div><h3>파운데이션 모델이 공통 연구 층이 된다</h3><p>LLM과 대규모 사전학습은 NLP 내부의 한 흐름을 넘어 retrieval, code, multimodal, robotics, science 응용이 공유하는 기본 인프라가 되었다.</p><p class="insight-implication">시사점: 새 연구는 모델 자체뿐 아니라 데이터, instruction, evaluation, deployment interface를 함께 설계해야 한다.</p></article>
        <article class="insight-box"><div class="insight-label">Vision And Multimodal</div><h3>비전과 멀티모달은 AI 영향력의 핵심 증폭기다</h3><p>ImageNet, ResNet, vision transformer, CLIP 계열 흐름은 시각 인식에서 언어·영상·로봇·의료 데이터까지 연결되는 표현학습 기반을 만들었다.</p><p class="insight-implication">시사점: 단일 modality benchmark보다 cross-modal grounding과 실제 데이터 변화에 대한 강건성이 중요하다.</p></article>
        <article class="insight-box"><div class="insight-label">Generative AI</div><h3>생성 모델은 합성 미디어에서 설계 도구로 이동한다</h3><p>GAN, VAE, diffusion, text-to-image 흐름은 이미지 생성에 머물지 않고 데이터 증강, 시뮬레이션, 과학·의료·콘텐츠 제작 워크플로로 확장된다.</p><p class="insight-implication">시사점: 생성 품질과 함께 provenance, controllability, misuse risk를 같이 평가해야 한다.</p></article>
        <article class="insight-box"><div class="insight-label">Agents And RL</div><h3>RL과 agents는 벤치마크에서 도구 사용으로 이동한다</h3><p>강화학습, planning, control, human feedback은 게임·로봇 문제를 넘어 LLM agent, tool use, embodied decision-making과 연결된다.</p><p class="insight-implication">시사점: agent 연구는 reward 설계보다 장기 안정성, 관찰 가능성, 실패 복구, human oversight가 핵심이다.</p></article>
        <article class="insight-box"><div class="insight-label">Trustworthy AI</div><h3>신뢰성은 배포의 병목으로 남는다</h3><p>explainability, robustness, fairness, privacy, safety 연구는 성능 중심 논문 지형에 비해 규모는 작지만 실제 적용의 승인 조건을 결정한다.</p><p class="insight-implication">시사점: 높은 benchmark score는 calibration, uncertainty, auditability 없이는 충분한 배포 근거가 되지 않는다.</p></article>
        <article class="insight-box"><div class="insight-label">AI4Science</div><h3>AI4Science는 모델의 외부 타당성을 시험한다</h3><p>단백질, 분자, 의료, 로보틱스, 자율주행 연구는 AI 모델이 텍스트와 이미지 밖의 물리·생물학적 제약을 얼마나 잘 다루는지 검증한다.</p><p class="insight-implication">시사점: 향후 영향력은 leaderboard보다 실험 검증, 도메인 제약, 실패 비용을 반영하는 평가에서 갈린다.</p></article>
      </div>
""",
    }


def top_metadata_values(rows, key, limit=3):
    counts = Counter()
    for row in rows:
        for value in str(row.get(key) or "").split("; "):
            if value:
                counts[value] += 1
    return [value for value, _ in counts.most_common(limit)]


def category_display_name(category, language):
    if language == "ko":
        return KOREAN_CATEGORY_NAMES.get(category, category)
    return category


def language_period_analysis(language, category, rows, start, end):
    count = len(rows)
    citations = sum(p["citationCount"] for p in rows)
    top = max(rows, key=lambda p: (p["citationCount"], p["influentialCitationCount"], p["title"]))
    active_year, active_count = Counter(p["year"] for p in rows).most_common(1)[0]
    tags = ", ".join(top_metadata_values(rows, "methodTags")) or "metadata-ranked"
    venues = ", ".join(top_metadata_values(rows, "venue")) or "mixed venues"
    name = category_display_name(category, language)
    if language == "ko":
        return {
            "categoryName": name,
            "overview": [
                f"{start}-{end} 기간의 {name} 분류에는 선정 논문 {count:,}편과 인용 {citations:,}회가 포함됩니다. 가장 활발한 연도는 {active_year}년({active_count:,}편)이며, 대표 상위 논문은 '{top['title']}'({top['citationCount']:,}회 인용)입니다.",
                f"이 기간의 주요 방법 태그는 {tags}이고, 자주 보이는 venue는 {venues}입니다.",
            ],
            "limitations": [
                TAXONOMY_LIMITATIONS.get(category, next(iter(TAXONOMY_LIMITATIONS.values())))[0],
                "기간별 citation ranking은 최근 논문을 구조적으로 불리하게 만들 수 있으므로 최신성은 별도 전문가 검토가 필요합니다.",
            ],
        }
    return {
        "categoryName": name,
        "overview": [
            f"In {start}-{end}, {name} contains {count:,} selected papers and {citations:,} citations. The most active year is {active_year} ({active_count:,} papers), and the leading citation-ranked paper is \"{top['title']}\" ({top['citationCount']:,} citations).",
            f"Frequent method tags include {tags}, with venue concentration around {venues}.",
        ],
        "limitations": [
            TAXONOMY_LIMITATIONS.get(category, next(iter(TAXONOMY_LIMITATIONS.values())))[0],
            "Citation-ranked period views structurally disadvantage very recent papers, so novelty needs separate expert review.",
        ],
    }


def overall_period_summary(rows, start, end):
    category_counts = Counter(p["category"] for p in rows)
    category_citations = defaultdict(int)
    keyword_counts = Counter()
    year_counts = Counter()
    year_citations = defaultdict(int)
    for paper in rows:
        category_citations[paper["category"]] += paper["citationCount"]
        year_counts[paper["year"]] += 1
        year_citations[paper["year"]] += paper["citationCount"]
        for keyword in str(paper.get("keywordTags") or "").split("; "):
            if keyword:
                keyword_counts[keyword] += 1
    peak_year, peak_count = year_counts.most_common(1)[0] if year_counts else (None, 0)
    peak_citation_year = max(year_citations, key=year_citations.get) if year_citations else None
    top = max(rows, key=lambda p: (p["citationCount"], p["influentialCitationCount"], p["title"])) if rows else None
    return {
        "startYear": start,
        "endYear": end,
        "rangeLabel": str(start) if start == end else f"{start}-{end}",
        "totalPapers": len(rows),
        "activeYears": len(year_counts),
        "citationCount": sum(p["citationCount"] for p in rows),
        "categoryCount": len(category_counts),
        "topCategories": [
            {"name": category, "count": count, "citations": category_citations[category]}
            for category, count in category_counts.most_common(6)
        ],
        "topKeywords": [
            {"name": keyword, "count": count}
            for keyword, count in keyword_counts.most_common(6)
        ],
        "peakYear": peak_year,
        "peakYearCount": peak_count,
        "peakCitationYear": peak_citation_year,
        "peakCitationCount": year_citations.get(peak_citation_year, 0) if peak_citation_year else 0,
        "topPaper": {
            "title": top["title"],
            "year": top["year"],
            "category": top["category"],
            "url": top["url"],
            "citations": top["citationCount"],
        } if top else None,
    }


def overall_research_templates():
    return {
        "en": {
            "timelineTitle": "Research Timeline",
            "summary": [
                "For {range}, the AI corpus contains {papers} selected papers across {activeYears} active years, with {citations} citations. The strongest taxonomy signals are {topCategories}, and the most active year is {peakYear} ({peakYearCount} papers).",
                "The leading citation-ranked paper is \"{topPaper}\" ({topPaperYear}, {topPaperCitations} citations) in {topPaperCategory}. Keywords such as {topKeywords} show how this period connects machine learning methods, foundation models, multimodal systems, generative AI, agents, trustworthy AI, and AI for Science.",
            ],
            "insightsTitle": "Research Insights",
            "insights": [
                {"label": "Foundation Layer", "title": "General representations become shared infrastructure", "body": "{topCategories} dominate {range}, showing which layers of AI became reusable research infrastructure.", "implication": "Implication: compare periods by data, model scale, evaluation, and deployment interface, not by algorithm names alone."},
                {"label": "Citation Mass", "title": "Citation peaks reveal infrastructure moments", "body": "The selected range carries {citations} citations, with citation mass peaking around {peakCitationYear}.", "implication": "Implication: high-impact AI work often packages reusable architectures, objectives, datasets, or benchmarks."},
                {"label": "Multimodal Shift", "title": "AI impact grows when modalities connect", "body": "Frequent tags such as {topKeywords} indicate whether the period is centered on language, vision, generation, graphs, agents, or scientific domains.", "implication": "Implication: robust cross-modal grounding and external validation matter more than isolated benchmark gains."},
                {"label": "Trust And Deployment", "title": "Trustworthiness remains the adoption bottleneck", "body": "Even when the leading paper is in {topPaperCategory}, responsible deployment questions shape how the period's methods are used outside benchmarks.", "implication": "Implication: calibration, uncertainty, provenance, privacy, and auditability should be tracked with performance."},
                {"label": "Open Gaps", "title": "Citation-ranked AI maps need recency correction", "body": "Recent model families, open-source systems, safety work, and negative results can be underweighted before citation signals mature.", "implication": "Implication: use this view as a stable map, then layer in recent expert review for fast-moving subfields."},
            ],
        },
        "ko": {
            "timelineTitle": "연구 타임라인",
            "summary": [
                "{range} 기간의 AI 코퍼스는 활성 연도 {activeYears}년에 걸쳐 선별 논문 {papers}편과 인용 {citations}회를 포함합니다. 가장 강한 taxonomy 신호는 {topCategories}이며, 논문 수가 가장 많은 해는 {peakYear}년({peakYearCount}편)입니다.",
                "인용 기준 최상위 논문은 {topPaperCategory} 분류의 \"{topPaper}\"({topPaperYear}, {topPaperCitations}회 인용)입니다. {topKeywords} 같은 키워드는 이 기간이 machine learning 방법론, foundation model, multimodal system, generative AI, agent, trustworthy AI, AI for Science를 어떻게 연결하는지 보여줍니다.",
            ],
            "insightsTitle": "Research Insights",
            "insights": [
                {"label": "Foundation Layer", "title": "범용 representation이 공유 인프라가 됩니다", "body": "{topCategories}가 {range}를 주도하며, AI의 어떤 층이 재사용 가능한 연구 인프라가 되었는지 보여줍니다.", "implication": "시사점: 기간 비교는 알고리즘 이름만이 아니라 data, model scale, evaluation, deployment interface를 함께 봐야 합니다."},
                {"label": "Citation Mass", "title": "인용 피크는 인프라 전환점을 드러냅니다", "body": "선택 기간은 총 {citations}회 인용을 가지며, 인용 질량은 {peakCitationYear}년 부근에서 정점입니다.", "implication": "시사점: 영향력 큰 AI 연구는 재사용 가능한 architecture, objective, dataset, benchmark를 함께 제공하는 경우가 많습니다."},
                {"label": "Multimodal Shift", "title": "AI 영향력은 modality가 연결될 때 커집니다", "body": "{topKeywords} 같은 빈도 높은 태그는 해당 기간이 language, vision, generation, graph, agent, scientific domain 중 어디에 중심을 두는지 나타냅니다.", "implication": "시사점: 단일 benchmark 향상보다 cross-modal grounding과 외부 검증이 중요합니다."},
                {"label": "Trust And Deployment", "title": "신뢰성은 여전히 도입의 병목입니다", "body": "최상위 논문이 {topPaperCategory}에 속하더라도, responsible deployment 질문은 benchmark 밖에서 방법이 쓰이는 방식을 결정합니다.", "implication": "시사점: 성능과 함께 calibration, uncertainty, provenance, privacy, auditability를 추적해야 합니다."},
                {"label": "Open Gaps", "title": "인용 기반 AI 지도에는 최신성 보정이 필요합니다", "body": "최신 모델 계열, open-source system, safety 연구, 부정 결과는 인용 신호가 성숙하기 전까지 과소평가될 수 있습니다.", "implication": "시사점: 이 뷰는 안정적 지도층으로 쓰고, 빠르게 움직이는 하위 분야는 최신 전문가 리뷰를 덧붙여야 합니다."},
            ],
        },
        "zh": {
            "timelineTitle": "研究时间线",
            "summary": [
                "在 {range} 期间，AI 语料包含 {papers} 篇入选论文，覆盖 {activeYears} 个活跃年份，总引用为 {citations} 次。最强的 taxonomy 信号是 {topCategories}，论文数量峰值出现在 {peakYear} 年（{peakYearCount} 篇）。",
                "按引用排序的领先论文是 {topPaperCategory} 中的《{topPaper}》（{topPaperYear}，{topPaperCitations} 次引用）。{topKeywords} 等关键词显示该时期如何连接机器学习方法、基础模型、多模态系统、生成式 AI、智能体、可信 AI 和 AI for Science。",
            ],
            "insightsTitle": "Research Insights",
            "insights": [
                {"label": "Foundation Layer", "title": "通用表示成为共享基础设施", "body": "{topCategories} 主导 {range}，显示 AI 的哪些层变成可复用研究基础设施。", "implication": "启示：比较时期时应同时看数据、模型规模、评估和部署接口，而不只是算法名称。"},
                {"label": "Citation Mass", "title": "引用峰值揭示基础设施时刻", "body": "所选时期共有 {citations} 次引用，引用质量在 {peakCitationYear} 年附近达到峰值。", "implication": "启示：高影响 AI 工作往往提供可复用架构、目标函数、数据集或基准。"},
                {"label": "Multimodal Shift", "title": "模态连接时 AI 影响力扩大", "body": "{topKeywords} 等高频标签显示该时期是围绕语言、视觉、生成、图、智能体还是科学领域组织的。", "implication": "启示：稳健的跨模态 grounding 和外部验证比孤立 benchmark 增益更重要。"},
                {"label": "Trust And Deployment", "title": "可信性仍是采用瓶颈", "body": "即使领先论文属于 {topPaperCategory}，负责任部署问题也会塑造这些方法在 benchmark 之外的使用方式。", "implication": "启示：应与性能一起追踪校准、不确定性、来源、隐私和可审计性。"},
                {"label": "Open Gaps", "title": "引用排序 AI 地图需要新近性校正", "body": "最新模型家族、开源系统、安全研究和负结果在引用信号成熟前可能被低估。", "implication": "启示：把此视图作为稳定地图，再为快速变化的子领域叠加最新专家评审。"},
            ],
        },
        "ja": {
            "timelineTitle": "研究タイムライン",
            "summary": [
                "{range} の AI コーパスには、{activeYears} の対象年にわたる選定論文 {papers} 本、引用 {citations} 件が含まれます。最も強い taxonomy 信号は {topCategories} で、論文数のピークは {peakYear} 年（{peakYearCount} 本）です。",
                "引用順で最上位の論文は {topPaperCategory} の「{topPaper}」（{topPaperYear}、{topPaperCitations} 件引用）です。{topKeywords} などのキーワードは、この期間が machine learning 方法、foundation model、multimodal system、generative AI、agent、trustworthy AI、AI for Science をどう結びつけているかを示します。",
            ],
            "insightsTitle": "Research Insights",
            "insights": [
                {"label": "Foundation Layer", "title": "汎用表現が共有インフラになります", "body": "{topCategories} が {range} を主導し、AI のどの層が再利用可能な研究インフラになったかを示します。", "implication": "示唆：期間比較ではアルゴリズム名だけでなく、data、model scale、evaluation、deployment interface を見る必要があります。"},
                {"label": "Citation Mass", "title": "引用ピークはインフラ化の瞬間を示します", "body": "選択期間の引用は {citations} 件で、引用質量は {peakCitationYear} 年付近でピークになります。", "implication": "示唆：高インパクトな AI 研究は再利用可能な architecture、objective、dataset、benchmark を伴うことが多いです。"},
                {"label": "Multimodal Shift", "title": "モダリティがつながると AI の影響は広がります", "body": "{topKeywords} などの高頻度タグは、その期間が language、vision、generation、graph、agent、scientific domain のどこを中心にしているかを示します。", "implication": "示唆：孤立した benchmark 改善より、cross-modal grounding と外部検証が重要です。"},
                {"label": "Trust And Deployment", "title": "信頼性は導入のボトルネックです", "body": "最上位論文が {topPaperCategory} に属していても、responsible deployment の問いは benchmark 外での使われ方を左右します。", "implication": "示唆：性能とともに calibration、uncertainty、provenance、privacy、auditability を追跡すべきです。"},
                {"label": "Open Gaps", "title": "引用ベースの AI 地図には最新性補正が必要です", "body": "新しいモデル系列、open-source system、安全性研究、否定的結果は、引用信号が成熟するまで過小評価されることがあります。", "implication": "示唆：このビューを安定した地図として使い、動きの速い下位分野には最新の専門家レビューを重ねてください。"},
            ],
        },
    }


def build_period_analysis(selected):
    ranges = [
        {"key": period_key(start, end), "label": period_label(start, end), "from": start, "to": end}
        for start, end in period_select_ranges()
    ]
    analysis = {}
    for start, end in all_period_ranges():
        rows = [p for p in selected if start <= p["year"] <= end]
        groups = category_groups(rows)
        entry = {}
        entry["__overall__"] = overall_period_summary(rows, start, end)
        for category, papers in groups.items():
            entry[safe_slug(category)] = {
                language: language_period_analysis(language, category, papers, start, end)
                for language in LANGUAGES
            }
        analysis[period_key(start, end)] = entry
    return {
        "generated": date.today().isoformat(),
        "yearRange": YEAR_RANGE_TEXT,
        "languages": LANGUAGES,
        "uiLabels": UI_LABELS,
        "ranges": ranges,
        "analysis": analysis,
    }


def write_period_analysis(selected):
    payload = build_period_analysis(selected)
    target = DATA_DIR / PERIOD_ANALYSIS_JSON
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text(json.dumps(payload, ensure_ascii=False, separators=(",", ":")), encoding="utf-8")


def write_taxonomy_dataset(selected):
    rows = []
    for category, papers in sorted(category_groups(selected).items()):
        overview = " ".join(TAXONOMY_TRENDS.get(category, []))
        limitations = " ".join(TAXONOMY_LIMITATIONS.get(category, []))
        for idx, paper in enumerate(papers, 1):
            row = dict(paper)
            row["taxonomyRank"] = idx
            row["categoryOverview"] = overview
            row["categoryLimitations"] = limitations
            rows.append(row)
    fields = [
        "category",
        "taxonomyRank",
        "rank",
        "yearRank",
        "candidateRank",
        "year",
        "title",
        "authors",
        "venue",
        "publicationDate",
        "citationCount",
        "influentialCitationCount",
        "importanceScore",
        "categoryOverview",
        "categoryLimitations",
        "methodTags",
        "keywordTags",
        "keyIdea",
        "strengths",
        "limitations",
        "url",
        "semanticScholarUrl",
        "openAccessPdf",
        "githubUrl",
        "doi",
        "arxiv",
        "pubmed",
        "fieldsOfStudy",
        "sourceQueries",
        "importanceReasons",
    ]
    write_csv(DATA_DIR / TAXONOMY_CSV, rows, fields)


def shields_keyword_badge(keyword):
    color = KEYWORD_COLORS.get(keyword, "64748b")
    badge = keyword.replace("-", "--")
    return f"https://img.shields.io/badge/keyword-{badge}-{color}"


def readme_keyword_badges(keywords):
    return " ".join(
        f'<img alt="{html.escape(keyword)}" src="{shields_keyword_badge(keyword)}">'
        for keyword in keywords
    )


def readme_keyword_convention_lines():
    lines = []
    for keyword, description, _ in KEYWORD_CONVENTION:
        lines.append(
            f"- ![{keyword}]({shields_keyword_badge(keyword)}) **{keyword}**: {description}"
        )
    return lines


def md_link(label, url):
    return f"[{label}]({url})" if url else label


def readme_taxonomy_table(rows, max_rows=10):
    out = [
        '<table width="100%">',
        "<colgroup>",
        '<col width="5%">',
        '<col width="22%">',
        '<col width="12%">',
        '<col width="12%">',
        '<col width="25%">',
        '<col width="12%">',
        '<col width="12%">',
        "</colgroup>",
        "<thead><tr>",
        '<th align="right">Rank</th><th>Paper</th><th>Meta</th><th>Keywords</th><th>Key idea</th><th>Strengths</th><th>Limitations</th>',
        "</tr></thead><tbody>",
    ]
    for paper in rows[:max_rows]:
        keywords = [x for x in paper["keywordTags"].split("; ") if x]
        out.extend(
            [
                "<tr>",
                f'<td align="right">{paper["rank"]}</td>',
                f'<td>{md_link(html.escape(paper["title"]), html.escape(paper["url"]))}<br><sub>{html.escape(paper["authors"] or "Unknown authors")}</sub></td>',
                f'<td>{paper["year"]}<br>{html.escape(paper["venue"] or "Unknown venue")}<br>{paper["citationCount"]:,} citations</td>',
                f"<td>{readme_keyword_badges(keywords)}</td>",
                f"<td>{html.escape(paper['keyIdea'])}</td>",
                f"<td>{html.escape(paper['strengths'])}</td>",
                f"<td>{html.escape(paper['limitations'])}</td>",
                "</tr>",
            ]
        )
    if len(rows) > max_rows:
        out.append(
            f'<tr><td colspan="7"><em>{len(rows) - max_rows:,} additional selected papers in this category are available in the dataset and website.</em></td></tr>'
        )
    out.extend(["</tbody></table>"])
    return "\n".join(out)


def write_readme(selected, candidates):
    stats = category_stats(selected)
    years = year_stats(selected)
    total_cites = sum(p["citationCount"] for p in selected)
    candidate_counts = Counter(p["year"] for p in candidates)
    groups = category_groups(selected)
    lines = [
        "# Awesome Test",
        "",
        "[![Awesome](https://awesome.re/badge-flat.svg)](https://awesome.re)",
        "",
        f"A taxonomy-first, citation-ranked map of AI research from {START_YEAR} through {END_YEAR}.",
        "",
        '<p align="center">',
        '  <a href="https://honggi82.github.io/awesome-test/">',
        '    <img src="https://img.shields.io/badge/Open_Interactive_Website-honggi82.github.io%2Fawesome--test-0f766e?style=for-the-badge" alt="Open Interactive Website">',
        "  </a>",
        "</p>",
        "",
        f"Generated on {date.today().isoformat()} from free public Semantic Scholar metadata. This edition investigates up to {CANDIDATES_PER_YEAR:,} AI-related candidate papers per year for {YEAR_RANGE_TEXT}, keeps an audited candidate pool of {len(candidates):,} records, selects the top {TARGET_PER_YEAR:,} papers from each year by citation count ({TARGET_TOTAL:,} papers total), and reorganizes them by AI research taxonomy.",
        "",
        "## Project Links",
        "",
        "- Open Interactive Website: https://honggi82.github.io/awesome-test/",
        f"- Selected dataset: `data/{PAPERS_CSV}`",
        f"- Taxonomy dataset with paper-level ideas, strengths, and limitations: `data/{TAXONOMY_CSV}`",
        f"- Precomputed period and language analysis: `data/{PERIOD_ANALYSIS_JSON}`",
        f"- Candidate Pool: `data/{CANDIDATES_CSV}`",
        "- English review draft: `paper/review_en.html`, `paper/review_en.docx`",
        "- Korean review draft: `paper/review_ko.html`",
        "",
        "## Keywords Convention",
        "",
        "These badges define the AI keyword tags used to read and extend this collection.",
        "",
        *readme_keyword_convention_lines(),
        "",
        "## Taxonomy Overview",
        "",
        f"- **Total selected papers**: {len(selected):,} papers",
        f"- **Candidate pool audited**: {len(candidates):,} papers ({', '.join(f'{year}: {candidate_counts[year]:,}' for year in YEARS)})",
        f"- **Citation count in selected set**: {total_cites:,}",
    ]
    for category, count in stats.most_common():
        lines.append(f"- **{category}**: {count:,} papers")
    lines.extend(["", "## Taxonomy Collections", ""])
    for category, count in stats.most_common():
        rows = groups[category]
        citations = sum(p["citationCount"] for p in rows)
        covered_years = sorted({p["year"] for p in rows})
        lines.extend(
            [
                f"### {category}",
                "",
                f"- Papers selected: **{count:,}**",
                f"- Years covered: **{covered_years[0]}-{covered_years[-1]}**",
                f"- Citation count in selected set: **{citations:,}**",
                "- Category Overview (main research trends):",
            ]
        )
        lines.extend(f"  - {item}" for item in TAXONOMY_TRENDS[category])
        lines.append("- Limitations:")
        lines.extend(f"  - {item}" for item in TAXONOMY_LIMITATIONS[category])
        lines.extend(
            [
                "",
                f"<details>",
                f"<summary><strong>Show representative papers for {category}</strong></summary>",
                "",
                readme_taxonomy_table(rows),
                "",
                "</details>",
                "",
            ]
        )
    lines.extend(["## Yearly Coverage", ""])
    lines.append("| Year | Candidate papers audited | Selected top-100 papers | Citations in selected set | Top selected paper |")
    lines.append("|---:|---:|---:|---:|---|")
    for year in YEARS:
        stat = years.get(year)
        top = md_link(html.escape(stat["top"]["title"]), html.escape(stat["top"]["url"])) if stat else "-"
        lines.append(
            f"| {year} | {candidate_counts[year]:,} | {stat['count'] if stat else 0:,} | {stat['citations'] if stat else 0:,} | {top} |"
        )
    lines.extend(
        [
            "",
            "## Methodology",
            "",
            f"The collection uses Semantic Scholar Academic Graph bulk search. Queries cover broad AI, machine learning, deep learning, foundation models, language, vision, reinforcement learning, generative models, graph learning, multimodal learning, trustworthy AI, and AI-for-science themes. For each year from {START_YEAR} through {END_YEAR}, results are filtered to the publication year, screened with explicit AI relevance expressions in title/abstract/venue metadata, deduplicated by DOI, arXiv, PubMed, CorpusId, paperId, then title, and reduced to at most {CANDIDATES_PER_YEAR:,} candidates by citation count. The final awesome list selects the top {TARGET_PER_YEAR:,} papers within each publication year by citation count; influential citation count and a deterministic metadata importance score are retained as tie-breakers and audit signals.",
            "",
            "The taxonomy, key ideas, strengths, limitations, method tags, and keyword tags are generated deterministically from public metadata and rule-based domain conventions. No paid API, paid LLM, paid translation, or paid compute was used.",
            "",
            "## Caveats",
            "",
            "- This is a metadata-driven citation map, not a full systematic review of every PDF.",
            f"- Citation counts favor older papers; {END_YEAR} should be interpreted as a partial and still-moving year.",
            "- Semantic Scholar metadata can omit venues, abstracts, PDFs, or influential citation counts for some records.",
            "- Citation ranking measures influence and visibility; it does not directly measure methodological quality, safety, or reproducibility.",
            "",
            "## Acknowledgements",
            "",
            "This repository and interactive site were created with appreciation for [jehyunlee/paper-curation](https://github.com/jehyunlee/paper-curation). Its paper-curation workflow and repository organization informed the approach used here for a taxonomy-first, citation-ranked research map.",
            "",
            "## License",
            "",
            "CC-BY-4.0 for text and metadata curation; upstream paper metadata belongs to original sources.",
        ]
    )
    (ROOT / "README.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


def chart_color(index):
    colors = ["#2563eb", "#0f766e", "#a855f7", "#f59e0b", "#dc2626", "#0891b2", "#be123c", "#16a34a", "#64748b"]
    return colors[index % len(colors)]


def write_svg(path, content):
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")


def svg_text(value):
    return html.escape(str(value), quote=True)


def render_category_chart(rows, start, end):
    data = category_stats(rows).most_common()
    width = 1000
    row_h = 44
    height = max(260, 96 + row_h * max(1, len(data)))
    max_value = max([value for _, value in data] or [1])
    bars = []
    if not data:
        bars.append('<text x="500" y="150" text-anchor="middle" fill="#64748b" font-size="24">No selected papers</text>')
    for i, (category, value) in enumerate(data):
        y = 72 + i * row_h
        bar_w = int((value / max_value) * 520)
        bars.append(f'<text x="32" y="{y + 24}" fill="#172033" font-size="15">{svg_text(category[:52])}</text>')
        bars.append(f'<rect x="420" y="{y + 4}" width="{bar_w}" height="26" rx="5" fill="{chart_color(i)}"/>')
        bars.append(f'<text x="{430 + bar_w}" y="{y + 24}" fill="#172033" font-size="15" font-weight="700">{value}</text>')
    return f"""<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" viewBox="0 0 {width} {height}" role="img" aria-label="Category distribution {start}-{end}">
<rect width="100%" height="100%" fill="#ffffff"/>
<text x="32" y="38" fill="#0f172a" font-size="24" font-weight="800">AI Paper Taxonomy, {period_label(start, end)}</text>
<text x="32" y="60" fill="#64748b" font-size="14">Selected papers by category</text>
{''.join(bars)}
</svg>"""


def render_citation_chart(rows, start, end):
    width = 1000
    height = 420
    stats = year_stats(rows)
    years = list(range(start, end + 1))
    values = [stats.get(year, {}).get("citations", 0) for year in years]
    max_value = max(values or [1]) or 1
    plot_x = 80
    plot_y = 70
    plot_w = 850
    plot_h = 270
    gap = 14
    bar_w = max(24, int((plot_w - gap * (len(years) - 1)) / max(1, len(years))))
    bars = []
    for i, (year, value) in enumerate(zip(years, values)):
        h = int((value / max_value) * plot_h) if value else 0
        x = plot_x + i * (bar_w + gap)
        y = plot_y + plot_h - h
        bars.append(f'<rect x="{x}" y="{y}" width="{bar_w}" height="{h}" rx="5" fill="#7c3aed"/>')
        bars.append(f'<text x="{x + bar_w / 2:.1f}" y="{plot_y + plot_h + 28}" text-anchor="middle" fill="#475569" font-size="14">{year}</text>')
        if value:
            bars.append(f'<text x="{x + bar_w / 2:.1f}" y="{max(58, y - 8)}" text-anchor="middle" fill="#172033" font-size="13" font-weight="700">{value:,}</text>')
    return f"""<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" viewBox="0 0 {width} {height}" role="img" aria-label="Yearly citation mass {start}-{end}">
<rect width="100%" height="100%" fill="#ffffff"/>
<text x="32" y="38" fill="#0f172a" font-size="24" font-weight="800">Yearly Citation Mass, {period_label(start, end)}</text>
<text x="32" y="60" fill="#64748b" font-size="14">Citation counts among visible selected papers</text>
<line x1="{plot_x}" y1="{plot_y + plot_h}" x2="{plot_x + plot_w}" y2="{plot_y + plot_h}" stroke="#cbd5e1" stroke-width="2"/>
{''.join(bars)}
</svg>"""


def write_charts(selected):
    period_dir = DOCS_DIR / "assets" / "periods"
    if period_dir.exists():
        for stale in period_dir.glob("*.svg"):
            stale.unlink()
    for start, end in all_period_ranges():
        rows = [p for p in selected if start <= p["year"] <= end]
        write_svg(period_dir / f"category_distribution_{start}_{end}.svg", render_category_chart(rows, start, end))
        write_svg(period_dir / f"yearly_citations_{start}_{end}.svg", render_citation_chart(rows, start, end))
    assets = DOCS_DIR / "assets"
    shutil.copyfile(period_dir / f"category_distribution_{START_YEAR}_{END_YEAR}.svg", assets / "category_distribution.svg")
    shutil.copyfile(period_dir / f"yearly_citations_{START_YEAR}_{END_YEAR}.svg", assets / "yearly_citations.svg")


def taxonomy_icon_svg(category):
    colors = {
        "Foundation Models and Large Language Models": ("#2563eb", "#93c5fd"),
        "Generative Models and Synthetic Media": ("#a855f7", "#f0abfc"),
        "Vision and Multimodal Learning": ("#0f766e", "#67e8f9"),
        "Natural Language Processing and Knowledge": ("#f59e0b", "#fde68a"),
        "Reinforcement Learning and Agents": ("#dc2626", "#fca5a5"),
        "Representation, Self-Supervised, and Transfer Learning": ("#0891b2", "#7dd3fc"),
        "Trustworthy, Explainable, and Responsible AI": ("#be123c", "#f9a8d4"),
        "Graph Learning, Recommendation, and Core Methods": ("#4f46e5", "#c4b5fd"),
        "AI for Science, Healthcare, and Robotics": ("#16a34a", "#86efac"),
        "General AI Methods and Systems": ("#334155", "#cbd5e1"),
    }
    accent, soft = colors.get(category, colors["General AI Methods and Systems"])
    label = category.split(" and ")[0][:28]
    return f"""<svg xmlns="http://www.w3.org/2000/svg" width="320" height="200" viewBox="0 0 320 200" role="img" aria-label="{svg_text(category)}">
<rect x="0" y="0" width="320" height="200" rx="18" fill="#f8fafc"/>
<circle cx="82" cy="78" r="34" fill="{soft}" stroke="{accent}" stroke-width="5"/>
<circle cx="164" cy="58" r="24" fill="#fff" stroke="{accent}" stroke-width="5"/>
<circle cx="222" cy="120" r="36" fill="{soft}" stroke="{accent}" stroke-width="5"/>
<path d="M112 78 C132 64 138 60 142 58" stroke="{accent}" stroke-width="6" fill="none" stroke-linecap="round"/>
<path d="M182 72 C200 88 206 100 212 108" stroke="{accent}" stroke-width="6" fill="none" stroke-linecap="round"/>
<path d="M92 112 C118 148 176 158 214 140" stroke="{accent}" stroke-width="4" fill="none" stroke-linecap="round" stroke-dasharray="10 8"/>
<text x="160" y="178" text-anchor="middle" fill="#172033" font-size="18" font-weight="800">{svg_text(label)}</text>
</svg>"""


def write_taxonomy_icons(selected):
    icon_dir = DOCS_DIR / "assets" / "taxonomy"
    icon_dir.mkdir(parents=True, exist_ok=True)
    missing = []
    for category in category_stats(selected):
        target = icon_dir / f"{safe_slug(category)}.png"
        if not target.exists():
            missing.append(target.as_posix())
    if missing:
        raise FileNotFoundError("Missing taxonomy PNG assets: " + ", ".join(missing))


def html_attrs(paper):
    attrs = {
        "data-paper-id": paper["rank"],
        "data-year": paper["year"],
        "data-citations": paper["citationCount"],
        "data-keywords": " ".join(paper["keywordTags"].split("; ")),
        "data-title": paper["title"],
        "data-venue": paper["venue"],
    }
    return " ".join(f'{key}="{html.escape(str(value), quote=True)}"' for key, value in attrs.items())


def paper_localization_payload(rows):
    return {
        "generated": date.today().isoformat(),
        "languages": LANGUAGES,
        "papers": {
            str(paper["rank"]): {
                language: localized_paper_fields(paper, language)
                for language in LANGUAGES
            }
            for paper in rows
        },
    }


def write_paper_localizations(target, rows):
    target.write_text(
        json.dumps(paper_localization_payload(rows), ensure_ascii=False, separators=(",", ":")),
        encoding="utf-8",
    )


def keyword_chips_html(keyword_string):
    tags = [tag for tag in keyword_string.split("; ") if tag]
    return "".join(
        f'<span class="keyword-chip paper-keyword" style="--chip-color:#{KEYWORD_COLORS.get(tag, "64748b")}">{html.escape(tag)}</span>'
        for tag in tags
    )


def paper_card(paper, taxonomy_rank):
    semantic = f'<a href="{html.escape(paper["semanticScholarUrl"])}">Semantic Scholar</a>' if paper.get("semanticScholarUrl") else ""
    pdf = f'<a href="{html.escape(paper["openAccessPdf"])}">PDF</a>' if paper.get("openAccessPdf") else ""
    doi = f'<a href="{html.escape(paper["url"])}">Paper</a>' if paper.get("url") else ""
    github = f'<a href="{html.escape(paper["githubUrl"])}">GitHub</a>' if paper.get("githubUrl") else ""
    links = " ".join(x for x in (doi, semantic, pdf, github) if x)
    return f"""
      <article class="paper-card" {html_attrs(paper)}>
        <div class="paper-rank">#{paper['rank']}</div>
        <div>
          <h3><a href="{html.escape(paper['url'] or paper['semanticScholarUrl'])}">{html.escape(paper['title'])}</a></h3>
          <p class="authors">{html.escape(paper['authors'] or 'Unknown authors')}</p>
          <div class="meta">
            <span>{paper['year']}</span>
            <span>{html.escape(paper['venue'] or 'Unknown venue')}</span>
            <span>{paper['citationCount']:,} citations</span>
            <span>{paper['influentialCitationCount']:,} influential</span>
            <span>score {paper['importanceScore']}</span>
            <span>year #{paper.get('yearRank', paper.get('candidateRank', ''))}</span>
            <span>taxonomy #{taxonomy_rank}</span>
          </div>
          <div class="paper-keywords" aria-label="Keywords">{keyword_chips_html(paper['keywordTags'])}</div>
          <div class="assessment">
            <p><strong class="paper-key-idea-label">Key idea:</strong> <span class="paper-key-idea-text">{html.escape(paper['keyIdea'])}</span></p>
            <p><strong class="paper-strengths-label">Strengths:</strong> <span class="paper-strengths-text">{html.escape(paper['strengths'])}</span></p>
            <p><strong class="paper-limitations-label">Limitations:</strong> <span class="paper-limitations-text">{html.escape(paper['limitations'])}</span></p>
          </div>
          <p class="links">{links}</p>
        </div>
      </article>
"""


def taxonomy_section(category, rows):
    slug = safe_slug(category)
    citations = sum(p["citationCount"] for p in rows)
    years = sorted({p["year"] for p in rows})
    top = rows[0]
    cards = "".join(paper_card(paper, idx) for idx, paper in enumerate(rows, 1))
    trends = "".join(f"<li>{html.escape(item)}</li>" for item in TAXONOMY_TRENDS[category])
    limitations = "".join(f"<li>{html.escape(item)}</li>" for item in TAXONOMY_LIMITATIONS[category])
    return f"""
    <section class="taxonomy-section" data-category="{slug}">
      <details>
        <summary>
          <img class="summary-thumb" src="assets/taxonomy/{slug}.png" alt="">
          <span class="summary-title">{html.escape(category)}</span>
          <span class="category-count">{len(rows):,} papers</span>
          <span class="category-years">{years[0]}-{years[-1]}</span>
          <span class="category-citations">{citations:,} citations</span>
        </summary>
        <div class="section-intro">
          <figure class="section-visual"><img src="assets/taxonomy/{slug}.png" alt="{html.escape(category)} illustration"></figure>
          <p><strong class="top-paper-label">Top paper:</strong> <span class="top-paper">{html.escape(top['title'])}</span></p>
          <div class="insight-grid">
            <div class="insight-box">
              <strong class="overview-heading">Category Overview</strong>
              <ul class="category-overview-list">{trends}</ul>
            </div>
            <div class="insight-box limitation-box">
              <strong class="limitation-heading">Limitations</strong>
              <ul class="category-limitations-list">{limitations}</ul>
            </div>
          </div>
        </div>
        <div class="paper-list">{cards}</div>
      </details>
    </section>
"""


def all_taxonomy_section(rows):
    years = sorted({p["year"] for p in rows})
    citations = sum(p["citationCount"] for p in rows)
    top = max(rows, key=lambda p: p["citationCount"])
    return f"""
    <section class="taxonomy-section all-taxonomy-section" data-category="all-taxonomies">
      <details>
        <summary>
          <span class="all-taxonomy-thumb" aria-hidden="true">All</span>
          <span class="summary-title">All Taxonomies</span>
          <span class="category-count">{len(rows):,} papers</span>
          <span class="category-years">{years[0]}-{years[-1]}</span>
          <span class="category-citations">{citations:,} citations</span>
        </summary>
        <div class="section-intro">
          <p><strong class="top-paper-label">Top paper:</strong> <span class="top-paper">{html.escape(top['title'])}</span></p>
        </div>
        <div class="paper-list all-taxonomy-list"></div>
      </details>
    </section>
"""


def site_keyword_convention_html():
    return "\n".join(
        f"<button class='keyword-item' type='button' data-keyword='{html.escape(keyword)}' aria-pressed='false'><span class='keyword-chip' style='--chip-color:#{color}'>{html.escape(keyword)}</span><span>{html.escape(description)}</span></button>"
        for keyword, description, color in KEYWORD_CONVENTION
    )


def write_site(selected):
    apply_github_links(selected)
    DOCS_DIR.mkdir(exist_ok=True)
    (DOCS_DIR / "data").mkdir(exist_ok=True)
    (DOCS_DIR / "paper").mkdir(exist_ok=True)
    groups = category_groups(selected)
    stats = category_stats(selected)
    total_cites = sum(p["citationCount"] for p in selected)
    start_year_options = "\n".join(f'<option value="{year}"{" selected" if year == START_YEAR else ""}>{year}</option>' for year in YEARS)
    end_year_options = "\n".join(f'<option value="{year}"{" selected" if year == END_YEAR else ""}>{year}</option>' for year in YEARS)
    period_options = "\n".join(
        f'<option value="{period_key(start, end)}" data-from="{start}" data-to="{end}"{" selected" if start == START_YEAR and end == END_YEAR else ""}>{period_label(start, end)}</option>'
        for start, end in period_select_ranges()
    )
    language_options = "\n".join(
        f'<option value="{code}"{" selected" if code == "en" else ""}>{html.escape(label)}</option>'
        for code, label in LANGUAGES.items()
    )
    sections = "\n".join(
        [all_taxonomy_section(selected).strip()]
        + [taxonomy_section(category, groups[category]).strip() for category, _ in stats.most_common()]
    )
    keyword_html = site_keyword_convention_html()
    research_overview = research_overview_html().strip()
    research_copy_payload = json.dumps(research_copy(), ensure_ascii=False)
    overall_research_templates_payload = json.dumps(overall_research_templates(), ensure_ascii=False)
    static_site_copy_payload = json.dumps(site_static_copy(), ensure_ascii=False)
    period_analysis_url = f"https://raw.githubusercontent.com/{PROJECT_OWNER}/{PROJECT_REPO}/main/data/{PERIOD_ANALYSIS_JSON}"
    candidate_pool_url = f"https://raw.githubusercontent.com/{PROJECT_OWNER}/{PROJECT_REPO}/main/data/{CANDIDATES_CSV}"
    year_script = f"""
  <script>
    (() => {{
      const periodSelect = document.getElementById("periodPreset");
      const languageSelect = document.getElementById("languageSelect");
      const startSelect = document.getElementById("startYear");
      const endSelect = document.getElementById("endYear");
      const resetButton = document.getElementById("resetYears");
      const rangeStatus = document.getElementById("rangeStatus");
      const statPapers = document.getElementById("statPapers");
      const statYears = document.getElementById("statYears");
      const statCitations = document.getElementById("statCitations");
      const statCategories = document.getElementById("statCategories");
      const taxonomyTotalSummary = document.getElementById("taxonomyTotalSummary");
      const keywordFilterStatus = document.getElementById("keywordFilterStatus");
      const categoryChart = document.getElementById("categoryDistributionChart");
      const citationChart = document.getElementById("yearlyCitationsChart");
      const categoryChartCaption = document.getElementById("categoryChartCaption");
      const citationChartCaption = document.getElementById("citationChartCaption");
      const defaultStart = startSelect.value;
      const defaultEnd = endSelect.value;
      const defaultLanguage = languageSelect.value;
      const researchCopy = {research_copy_payload};
      const overallResearchTemplates = {overall_research_templates_payload};
      const staticCopy = {static_site_copy_payload};
      const periodOptions = Array.from(periodSelect.options);
      const validYears = Array.from(startSelect.options).map(option => option.value);
      const keywordGrid = document.querySelector(".keyword-grid");
      const keywordButtons = Array.from(document.querySelectorAll(".keyword-item[data-keyword]"));
      const allTaxonomiesSection = document.querySelector(".all-taxonomy-section");
      const allTaxonomiesDetails = allTaxonomiesSection?.querySelector("details");
      const allTaxonomiesList = allTaxonomiesSection?.querySelector(".all-taxonomy-list");
      let allTaxonomiesCards = [];
      let precomputed = null;
      let paperLocalizations = null;

      function formatNumber(value) {{ return Number(value).toLocaleString("en-US"); }}
      function rangeKey(start, end) {{ return `${{start}}-${{end}}`; }}
      function rangeLabel(start, end) {{ return start === end ? String(start) : `${{start}}-${{end}}`; }}
      function chartPath(kind, start, end) {{ return `assets/periods/${{kind}}_${{start}}_${{end}}.svg`; }}
      function selectedKeywords() {{
        return keywordButtons.filter(button => button.getAttribute("aria-pressed") === "true").map(button => button.dataset.keyword);
      }}
      function setKeywordPressed(button, pressed) {{
        button.setAttribute("aria-pressed", pressed ? "true" : "false");
        button.classList.toggle("is-selected", pressed);
      }}
      function keywordMatches(card, selected) {{
        if (!selected.length) return true;
        const cardKeywords = (card.dataset.keywords || "").split(" ").filter(Boolean);
        return selected.some(keyword => cardKeywords.includes(keyword));
      }}
      function yearRangeText(years) {{
        if (!years.length) return "No years";
        const sorted = [...new Set(years)].sort((a, b) => a - b);
        return sorted[0] === sorted[sorted.length - 1] ? String(sorted[0]) : `${{sorted[0]}}-${{sorted[sorted.length - 1]}}`;
      }}
      function selectedRangeValue(start, end) {{
        const match = periodOptions.find(option => option.dataset.from === String(start) && option.dataset.to === String(end));
        return match ? match.value : "custom";
      }}
      function updatePeriodSelect(start, end) {{ periodSelect.value = selectedRangeValue(start, end); }}
      function labels() {{
        const fallback = {{
          papers: "papers", categories: "categories", overview: "Category Overview", limitations: "Limitations",
          totalSelected: "Total selected papers", categoryCount: "Categories",
          keyIdea: "Key idea", strengths: "Strengths", paperLimitations: "Limitations",
          citations: "citations", selectedKeyword: "Selected keyword", matchingPapers: "Matching papers",
          allKeywords: "all", categoryChart: "Category distribution", citationChart: "Yearly citation mass",
          topPaper: "Top paper", allTaxonomies: "All Taxonomies"
        }};
        return {{...fallback, ...(precomputed?.uiLabels?.en || {{}}), ...(precomputed?.uiLabels?.[languageSelect.value] || {{}})}};
      }}
      function languageStaticCopy() {{
        return {{...(staticCopy.en || {{}}), ...(staticCopy[languageSelect.value] || {{}})}};
      }}
      function displayName(value) {{
        const current = languageStaticCopy();
        return current.categoryNames?.[value] || value;
      }}
      function updateStaticLocalization(copy) {{
        const current = languageStaticCopy();
        document.documentElement.lang = languageSelect.value;
        document.querySelectorAll("[data-i18n]").forEach(node => {{
          const key = node.dataset.i18n;
          if (current[key]) node.textContent = current[key];
        }});
        document.querySelectorAll(".keyword-description").forEach(node => {{
          const suffix = languageSelect.value.charAt(0).toUpperCase() + languageSelect.value.slice(1);
          const value = node.dataset[`description${{suffix}}`] || node.dataset.descriptionEn;
          if (value) node.textContent = value;
        }});
        document.querySelectorAll(".top-paper-label").forEach(node => {{ node.textContent = `${{copy.topPaper}}:`; }});
        const allTitle = allTaxonomiesSection?.querySelector(".summary-title");
        if (allTitle) allTitle.textContent = copy.allTaxonomies;
      }}
      function escapeHtml(value) {{
        const escapeMap = {{ "&": "&amp;", "<": "&lt;", ">": "&gt;", '"': "&quot;", "'": "&#39;" }};
        return String(value ?? "").replace(/[&<>"']/g, ch => escapeMap[ch]);
      }}
      function names(items, key = "name") {{
        return (items || []).slice(0, 3).map(item => displayName(item[key])).filter(Boolean).join(", ") || "n/a";
      }}
      function researchTemplateData(metric) {{
        const topCategories = metric.topCategories || [];
        const topKeywords = metric.topKeywords || [];
        const topCategory = topCategories[0] || {{}};
        const topPaper = metric.topPaper || {{}};
        return {{
          range: metric.rangeLabel || `${{metric.startYear}}-${{metric.endYear}}`,
          papers: formatNumber(metric.totalPapers),
          activeYears: formatNumber(metric.activeYears),
          citations: formatNumber(metric.citationCount),
          topCategories: names(topCategories),
          topCategory: displayName(topCategory.name || "n/a"),
          topCategoryCount: formatNumber(topCategory.count || 0),
          topKeywords: names(topKeywords),
          peakYear: metric.peakYear || "n/a",
          peakYearCount: formatNumber(metric.peakYearCount || 0),
          peakCitationYear: metric.peakCitationYear || "n/a",
          topPaper: topPaper.title || "n/a",
          topPaperYear: topPaper.year || "n/a",
          topPaperCategory: displayName(topPaper.category || "n/a"),
          topPaperCitations: formatNumber(topPaper.citations || 0)
        }};
      }}
      function applyTemplate(template, data) {{
        let output = template || "";
        Object.keys(data).forEach(key => {{
          output = output.split("{{" + key + "}}").join(escapeHtml(data[key]));
        }});
        return output;
      }}
      function renderOverallResearch(metric) {{
        const copy = overallResearchTemplates[languageSelect.value] || overallResearchTemplates.en;
        const data = researchTemplateData(metric);
        const summaryHtml = (copy.summary || []).map(text => `<p>${{applyTemplate(text, data)}}</p>`).join("");
        const insightHtml = (copy.insights || []).map(item => `
          <article class="insight-box">
            <div class="insight-label">${{escapeHtml(item.label)}}</div>
            <h3>${{applyTemplate(item.title, data)}}</h3>
            <p>${{applyTemplate(item.body, data)}}</p>
            <p class="insight-implication">${{applyTemplate(item.implication, data)}}</p>
          </article>`).join("");
        return `
          <h2 id="research-timeline-title">${{escapeHtml(copy.timelineTitle)}}</h2>
          <div class="timeline-copy">${{summaryHtml}}</div>
          <h2>${{escapeHtml(copy.insightsTitle)}}</h2>
          <div class="research-insights">${{insightHtml}}</div>`;
      }}
      function updateResearchCopy(start, end) {{
        const brief = document.getElementById("researchBrief");
        if (!brief) return;
        const metric = precomputed?.analysis?.[rangeKey(start, end)]?.__overall__;
        brief.innerHTML = metric ? renderOverallResearch(metric) : (researchCopy[languageSelect.value] || researchCopy.en);
      }}
      function setList(target, items) {{
        if (!target || !items) return;
        target.innerHTML = "";
        items.forEach(item => {{
          const li = document.createElement("li");
          li.textContent = item;
          target.appendChild(li);
        }});
      }}
      function updateCharts(start, end) {{
        const label = rangeLabel(start, end);
        if (categoryChart) {{
          categoryChart.src = chartPath("category_distribution", start, end);
          categoryChart.alt = `Category distribution chart for ${{label}}`;
        }}
        if (citationChart) {{
          citationChart.src = chartPath("yearly_citations", start, end);
          citationChart.alt = `Yearly citation chart for ${{label}}`;
        }}
        const copy = labels();
        if (categoryChartCaption) categoryChartCaption.textContent = `${{copy.categoryChart}} (${{label}})`;
        if (citationChartCaption) citationChartCaption.textContent = `${{copy.citationChart}} (${{label}})`;
      }}
      function localizedCardText(card, field, language) {{
        const paper = paperLocalizations?.papers?.[card.dataset.paperId];
        const selected = paper?.[language] || paper?.en;
        const jsonField = field === "paperLimitations" ? "limitations" : field;
        return selected?.[jsonField] || "";
      }}
      function applyPaperLocalization(card, copy) {{
        const language = languageSelect.value;
        [
          ["keyIdea", ".paper-key-idea-label", ".paper-key-idea-text", copy.keyIdea],
          ["strengths", ".paper-strengths-label", ".paper-strengths-text", copy.strengths],
          ["paperLimitations", ".paper-limitations-label", ".paper-limitations-text", copy.paperLimitations]
        ].forEach(([field, labelSelector, textSelector, label]) => {{
          const labelNode = card.querySelector(labelSelector);
          const textNode = card.querySelector(textSelector);
          const localized = localizedCardText(card, field, language);
          if (labelNode) labelNode.textContent = `${{label}}:`;
          if (textNode && localized) textNode.textContent = localized;
        }});
      }}
      function applyPrecomputedAnalysis(section, start, end) {{
        const entry = precomputed?.analysis?.[rangeKey(start, end)]?.[section.dataset.category];
        const analysis = entry?.[languageSelect.value] || entry?.en;
        if (!analysis) return;
        const copy = labels();
        const title = section.querySelector(".summary-title");
        if (title) title.textContent = analysis.categoryName;
        const overviewHeading = section.querySelector(".overview-heading");
        const limitationHeading = section.querySelector(".limitation-heading");
        if (overviewHeading) overviewHeading.textContent = copy.overview;
        if (limitationHeading) limitationHeading.textContent = copy.limitations;
        setList(section.querySelector(".category-overview-list"), analysis.overview);
        setList(section.querySelector(".category-limitations-list"), analysis.limitations);
      }}
      function syncUrl(start, end) {{
        const url = new URL(window.location.href);
        const language = languageSelect.value;
        if (language === defaultLanguage) url.searchParams.delete("lang"); else url.searchParams.set("lang", language);
        if (String(start) === defaultStart && String(end) === defaultEnd) {{
          url.searchParams.delete("period"); url.searchParams.delete("from"); url.searchParams.delete("to");
        }} else {{
          const value = selectedRangeValue(start, end);
          if (value !== "custom") {{ url.searchParams.set("period", value); url.searchParams.delete("from"); url.searchParams.delete("to"); }}
          else {{ url.searchParams.delete("period"); url.searchParams.set("from", start); url.searchParams.set("to", end); }}
        }}
        const keywords = selectedKeywords();
        if (keywords.length) url.searchParams.set("keywords", keywords.join(",")); else url.searchParams.delete("keywords");
        window.history.replaceState(null, "", url);
      }}
      function setFromUrl() {{
        const params = new URLSearchParams(window.location.search);
        const lang = params.get("lang");
        if (lang && Array.from(languageSelect.options).some(option => option.value === lang)) languageSelect.value = lang;
        const requestedKeywords = (params.get("keywords") || "").split(",").filter(Boolean);
        const keywordValues = keywordButtons.map(button => button.dataset.keyword);
        const activeKeyword = requestedKeywords.find(keyword => keywordValues.includes(keyword)) || "";
        keywordButtons.forEach(button => setKeywordPressed(button, button.dataset.keyword === activeKeyword));
        const period = params.get("period");
        if (period) {{
          const option = periodOptions.find(item => item.value === period && item.dataset.from && item.dataset.to);
          if (option) {{ periodSelect.value = period; startSelect.value = option.dataset.from; endSelect.value = option.dataset.to; return; }}
        }}
        const from = params.get("from");
        const to = params.get("to");
        if (validYears.includes(from)) startSelect.value = from;
        if (validYears.includes(to)) endSelect.value = to;
        updatePeriodSelect(startSelect.value, endSelect.value);
      }}
      function updateKeywordFilterStatus(selected, totalPapers, copy) {{
        if (!keywordFilterStatus) return;
        const keyword = selected.length ? selected[0] : "all";
        const keywordLabel = keyword === "all" ? copy.allKeywords : keyword;
        keywordFilterStatus.textContent = `${{copy.selectedKeyword}}: ${{keywordLabel}} | ${{copy.matchingPapers}}: ${{formatNumber(totalPapers)}} ${{copy.papers}}`;
      }}
      function renderAllTaxonomiesCards() {{
        if (!allTaxonomiesList || !allTaxonomiesDetails?.open) return;
        allTaxonomiesList.innerHTML = "";
        const fragment = document.createDocumentFragment();
        const sortedCards = [...allTaxonomiesCards].sort((a, b) => Number(b.dataset.citations || 0) - Number(a.dataset.citations || 0));
        sortedCards.forEach(card => fragment.appendChild(card.cloneNode(true)));
        allTaxonomiesList.appendChild(fragment);
      }}
      function updateAllTaxonomiesSection(cards, totalPapers, years, totalCitations, copy) {{
        if (!allTaxonomiesSection) return;
        allTaxonomiesCards = cards;
        allTaxonomiesSection.hidden = totalPapers === 0;
        allTaxonomiesSection.querySelector(".category-count").textContent = `${{formatNumber(totalPapers)}} ${{copy.papers}}`;
        allTaxonomiesSection.querySelector(".category-years").textContent = yearRangeText(years);
        allTaxonomiesSection.querySelector(".category-citations").textContent = `${{formatNumber(totalCitations)}} ${{copy.citations}}`;
        const topPaperTarget = allTaxonomiesSection.querySelector(".top-paper");
        const topPaper = [...cards].sort((a, b) => Number(b.dataset.citations || 0) - Number(a.dataset.citations || 0))[0]?.querySelector("h3");
        if (topPaper && topPaperTarget) topPaperTarget.textContent = topPaper.textContent.trim();
        if (allTaxonomiesDetails?.open) renderAllTaxonomiesCards();
        else if (allTaxonomiesList) allTaxonomiesList.innerHTML = "";
      }}
      function applyYearFilter(sync = true) {{
        let start = Number(startSelect.value);
        let end = Number(endSelect.value);
        if (start > end) {{ const previous = start; start = end; end = previous; startSelect.value = String(start); endSelect.value = String(end); }}
        const copy = labels();
        updateStaticLocalization(copy);
        updateResearchCopy(start, end);
        const activeKeywords = selectedKeywords();
        let totalPapers = 0;
        let totalCitations = 0;
        let activeCategories = 0;
        const activeYears = [];
        const visibleCards = [];
        document.querySelectorAll(".taxonomy-section:not(.all-taxonomy-section)").forEach(section => {{
          let sectionCount = 0;
          let sectionCitations = 0;
          const sectionYears = [];
          section.querySelectorAll(".paper-card").forEach(card => {{
            const year = Number(card.dataset.year);
            const citations = Number(card.dataset.citations || 0);
            const visible = year >= start && year <= end && keywordMatches(card, activeKeywords);
            applyPaperLocalization(card, copy);
            card.hidden = !visible;
            if (visible) {{ sectionCount += 1; sectionCitations += citations; sectionYears.push(year); activeYears.push(year); visibleCards.push(card); }}
          }});
          const hasPapers = sectionCount > 0;
          section.hidden = !hasPapers;
          if (!hasPapers) return;
          activeCategories += 1;
          totalPapers += sectionCount;
          totalCitations += sectionCitations;
          section.querySelector(".category-count").textContent = `${{formatNumber(sectionCount)}} ${{copy.papers}}`;
          section.querySelector(".category-years").textContent = yearRangeText(sectionYears);
          section.querySelector(".category-citations").textContent = `${{formatNumber(sectionCitations)}} ${{copy.citations}}`;
          const topPaper = section.querySelector(".paper-card:not([hidden]) h3");
          const topPaperTarget = section.querySelector(".top-paper");
          if (topPaper && topPaperTarget) topPaperTarget.textContent = topPaper.textContent.trim();
          applyPrecomputedAnalysis(section, start, end);
        }});
        updateAllTaxonomiesSection(visibleCards, totalPapers, activeYears, totalCitations, copy);
        statPapers.textContent = formatNumber(totalPapers);
        statYears.textContent = formatNumber(new Set(activeYears).size);
        statCitations.textContent = formatNumber(totalCitations);
        statCategories.textContent = formatNumber(activeCategories);
        updatePeriodSelect(start, end);
        updateCharts(start, end);
        taxonomyTotalSummary.innerHTML = `<strong>${{copy.totalSelected}}:</strong> ${{formatNumber(totalPapers)}} ${{copy.papers}}; <strong>${{copy.categoryCount}}:</strong> ${{formatNumber(activeCategories)}} ${{copy.categories}}.`;
        updateKeywordFilterStatus(activeKeywords, totalPapers, copy);
        const keywordText = activeKeywords.length ? ` | ${{activeKeywords[0]}}` : "";
        rangeStatus.textContent = `${{start}}-${{end}} | ${{formatNumber(totalPapers)}} ${{copy.papers}} | ${{formatNumber(activeCategories)}} ${{copy.categories}}${{keywordText}}`;
        if (sync) syncUrl(start, end);
      }}
      setFromUrl();
      applyYearFilter(false);
      Promise.allSettled([
        fetch("{period_analysis_url}").then(response => response.json()),
        fetch("data/{PAPER_LOCALIZATIONS_JSON}").then(response => response.json())
      ]).then(results => {{
        if (results[0].status === "fulfilled") precomputed = results[0].value;
        if (results[1].status === "fulfilled") paperLocalizations = results[1].value;
        applyYearFilter(false);
        if (results[0].status === "rejected") rangeStatus.textContent = "Precomputed analysis could not be loaded.";
      }});
      periodSelect.addEventListener("change", () => {{
        const option = periodSelect.selectedOptions[0];
        if (option && option.dataset.from && option.dataset.to) {{ startSelect.value = option.dataset.from; endSelect.value = option.dataset.to; }}
        applyYearFilter(true);
      }});
      languageSelect.addEventListener("change", () => applyYearFilter(true));
      startSelect.addEventListener("change", () => applyYearFilter(true));
      endSelect.addEventListener("change", () => applyYearFilter(true));
      if (keywordGrid) {{
        keywordGrid.addEventListener("click", event => {{
          const button = event.target.closest(".keyword-item[data-keyword]");
          if (!button || !keywordGrid.contains(button)) return;
          event.preventDefault();
          const pressed = button.getAttribute("aria-pressed") === "true";
          keywordButtons.forEach(item => setKeywordPressed(item, !pressed && item === button));
          applyYearFilter(true);
        }});
      }}
      allTaxonomiesDetails?.addEventListener("toggle", () => {{
        if (allTaxonomiesDetails.open) renderAllTaxonomiesCards();
        else if (allTaxonomiesList) allTaxonomiesList.innerHTML = "";
      }});
      resetButton.addEventListener("click", () => {{
        startSelect.value = defaultStart;
        endSelect.value = defaultEnd;
        periodSelect.value = `${{defaultStart}}-${{defaultEnd}}`;
        keywordButtons.forEach(button => setKeywordPressed(button, false));
        applyYearFilter(true);
      }});
    }})();
  </script>
"""
    html_doc = f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <link rel="icon" href="data:,">
  <title>Awesome Test</title>
  <style>
    :root {{ color-scheme: light; --ink:#172033; --muted:#5b6678; --line:#d9dee8; --accent:#2563eb; --accent2:#0f766e; --bg:#f7f9fc; --panel:#ffffff; }}
    body {{ margin:0; font-family: Inter, Segoe UI, Arial, sans-serif; color:var(--ink); background:var(--bg); }}
    header {{ padding:54px 7vw 34px; background:linear-gradient(120deg,#eff6ff,#ecfdf5); border-bottom:1px solid var(--line); }}
    h1 {{ font-size:48px; margin:0 0 12px; letter-spacing:0; }}
    h2 {{ margin-top:36px; }}
    p {{ line-height:1.65; color:var(--muted); }}
    main {{ padding:28px 7vw 72px; }}
    nav a {{ display:inline-block; margin:0 12px 10px 0; color:#0f5f97; font-weight:700; }}
    a {{ color:#0f5f97; text-decoration:none; }}
    a:hover {{ text-decoration:underline; }}
    .stats {{ display:grid; grid-template-columns:repeat(auto-fit,minmax(190px,1fr)); gap:12px; margin:24px 0; }}
    .stat {{ background:white; border:1px solid var(--line); border-radius:8px; padding:16px; }}
    .stat strong {{ display:block; font-size:28px; color:var(--accent); }}
    .stat span {{ display:block; margin-top:8px; color:var(--muted); }}
    .research-brief {{ margin:28px 0; padding:24px 0; border-top:1px solid var(--line); border-bottom:1px solid var(--line); }}
    .timeline-copy {{ max-width:1080px; }}
    .research-insights {{ display:grid; grid-template-columns:repeat(auto-fit,minmax(260px,1fr)); gap:12px; margin-top:12px; }}
    .research-insights .insight-label {{ color:var(--accent); font-size:12px; font-weight:800; text-transform:uppercase; letter-spacing:0; }}
    .research-insights .insight-box h3 {{ margin:6px 0 8px; font-size:17px; }}
    .research-insights .insight-box p {{ margin:8px 0 0; }}
    .insight-implication {{ color:var(--ink); font-weight:700; }}
    .filters {{ display:flex; flex-wrap:wrap; align-items:end; gap:12px; margin:24px 0; padding:14px; background:white; border:1px solid var(--line); border-radius:8px; }}
    .filter-field {{ display:grid; gap:6px; }}
    .wide-field {{ min-width:min(100%,280px); }}
    .filter-field label {{ font-weight:700; font-size:13px; color:#344255; }}
    select {{ min-width:104px; height:38px; border:1px solid var(--line); border-radius:8px; background:white; color:var(--ink); padding:0 10px; font:inherit; }}
    #periodPreset {{ min-width:280px; }}
    button {{ min-height:38px; border:1px solid #bfc8d8; border-radius:8px; background:#f6f8fb; color:var(--ink); padding:0 14px; font-weight:700; cursor:pointer; }}
    button:hover {{ background:#eef2f7; }}
    #rangeStatus {{ color:var(--muted); font-weight:700; min-height:38px; display:inline-flex; align-items:center; }}
    .keyword-section {{ margin:28px 0; }}
    .keyword-grid {{ display:grid; grid-template-columns:repeat(auto-fit,minmax(240px,1fr)); gap:10px; }}
    .keyword-item {{ display:grid; gap:8px; align-items:flex-start; height:auto; min-height:96px; padding:12px; background:white; border:1px solid var(--line); border-radius:8px; color:var(--muted); line-height:1.45; text-align:left; font:inherit; cursor:pointer; }}
    .keyword-item[aria-pressed="true"], .keyword-item.is-selected {{ border-color:var(--accent); box-shadow:0 0 0 2px rgba(37,99,235,0.16); color:var(--ink); }}
    .keyword-chip {{ justify-self:start; min-width:0; text-align:center; background:var(--chip-color); color:white; border-radius:999px; padding:4px 9px; font-size:13px; font-weight:800; }}
    .keyword-item > span:not(.keyword-chip) {{ display:block; width:100%; }}
    .keyword-filter-status {{ margin:10px 0 0; font-weight:700; color:var(--accent2); }}
    .figures {{ display:grid; grid-template-columns:repeat(auto-fit,minmax(280px,1fr)); gap:16px; margin:24px 0; }}
    .chart-figure {{ margin:0; }}
    .chart-figure figcaption {{ margin-top:8px; color:var(--muted); font-size:13px; font-weight:700; }}
    .figures img {{ width:100%; aspect-ratio:16 / 9; object-fit:contain; background:white; border:1px solid var(--line); border-radius:8px; display:block; }}
    .taxonomy-section {{ margin-top:16px; }}
    .taxonomy-section[hidden], .paper-card[hidden] {{ display:none !important; }}
    details {{ background:var(--panel); border:1px solid var(--line); border-radius:8px; overflow:hidden; }}
    summary {{ cursor:pointer; display:grid; grid-template-columns:64px minmax(260px,1fr) repeat(3,minmax(110px,auto)); gap:12px; align-items:center; padding:14px 18px; font-weight:700; }}
    .summary-thumb, .all-taxonomy-thumb {{ width:56px; height:40px; object-fit:cover; border:1px solid var(--line); border-radius:6px; background:#f8fafc; }}
    .all-taxonomy-thumb {{ display:inline-flex; align-items:center; justify-content:center; color:var(--accent); font-weight:800; }}
    .summary-title {{ color:var(--accent); }}
    .section-intro {{ padding:0 18px 14px; border-top:1px solid var(--line); }}
    .section-visual {{ margin:14px 0 4px; }}
    .section-visual img {{ width:min(320px,100%); aspect-ratio:16 / 11; object-fit:contain; border:1px solid var(--line); border-radius:8px; background:#f8fafc; display:block; }}
    .insight-grid {{ display:grid; grid-template-columns:repeat(auto-fit,minmax(260px,1fr)); gap:12px; margin-top:12px; }}
    .insight-box {{ padding:12px 14px; background:#f4f8ff; border:1px solid #cfe0ff; border-radius:8px; }}
    .limitation-box {{ background:#fff8f1; border-color:#ead7c1; }}
    .insight-box ul {{ margin:8px 0 0; padding-left:20px; color:var(--muted); line-height:1.55; }}
    .paper-list {{ display:grid; gap:12px; padding:16px; background:#f9fbfd; }}
    .paper-card {{ display:grid; grid-template-columns:56px 1fr; gap:14px; padding:16px; background:white; border:1px solid var(--line); border-radius:8px; }}
    .paper-rank {{ font-weight:800; color:var(--accent2); }}
    .paper-card h3 {{ margin:0 0 6px; font-size:18px; line-height:1.35; }}
    .authors {{ margin:0 0 8px; }}
    .meta {{ display:flex; flex-wrap:wrap; gap:8px; margin:8px 0 10px; }}
    .meta span {{ display:inline-block; background:#eef2f7; border:1px solid #dce3ee; border-radius:999px; padding:5px 9px; color:#344255; font-size:13px; }}
    .paper-keywords {{ display:flex; flex-wrap:wrap; gap:6px; margin:0 0 10px; }}
    .paper-keyword {{ min-width:0; font-size:12px; padding:3px 8px; }}
    .assessment {{ display:grid; grid-template-columns:repeat(auto-fit,minmax(260px,1fr)); gap:8px; }}
    .links a {{ margin-right:12px; font-weight:700; }}
    @media (max-width:760px) {{
      h1 {{ font-size:36px; }}
      summary {{ grid-template-columns:1fr; }}
      .paper-card {{ grid-template-columns:1fr; }}
    }}
  </style>
</head>
<body>
  <header>
    <h1>Awesome Test</h1>
    <p data-i18n="hero">A taxonomy-first, citation-ranked map of AI research from {START_YEAR} through {END_YEAR}. Each year investigates up to {CANDIDATES_PER_YEAR:,} candidate papers; the final collection selects the top {TARGET_PER_YEAR:,} papers from each year by citation count ({TARGET_TOTAL:,} papers total).</p>
    <nav>
      <a href="https://github.com/honggi82/awesome-test" data-i18n="navReadme">README</a>
      <a href="data/{PAPERS_CSV}" data-i18n="navDataset">CSV Dataset</a>
      <a href="data/{TAXONOMY_CSV}" data-i18n="navTaxonomyCsv">Taxonomy CSV</a>
      <a href="{period_analysis_url}" data-i18n="navPeriodJson">Period Analysis JSON</a>
      <a href="#keywords-convention" data-i18n="navKeywords">Keywords Convention</a>
      <a href="{candidate_pool_url}" data-i18n="navCandidatePool">Candidate Pool</a>
      <a href="paper/review_en.html" data-i18n="navReview">Review Paper</a>
      <a href="paper/review_ko.html" data-i18n="navKoreanReview">Korean Review</a>
    </nav>
  </header>
  <main>
    <div class="stats">
      <div class="stat"><strong id="statPapers">{len(selected)}</strong><span data-i18n="statSelected">selected papers</span></div>
      <div class="stat"><strong id="statYears">{len(year_stats(selected))}</strong><span data-i18n="statYears">years represented</span></div>
      <div class="stat"><strong id="statCitations">{total_cites:,}</strong><span data-i18n="statCitations">citation count total</span></div>
      <div class="stat"><strong id="statCategories">{len(stats)}</strong><span data-i18n="statCategories">topic categories</span></div>
    </div>
    <form class="filters" id="yearFilter">
      <div class="filter-field wide-field">
        <label for="periodPreset" data-i18n="filterPeriod">Period</label>
        <select id="periodPreset" name="period">{period_options}</select>
      </div>
      <div class="filter-field">
        <label for="languageSelect" data-i18n="filterLanguage">Language</label>
        <select id="languageSelect" name="lang">{language_options}</select>
      </div>
      <div class="filter-field">
        <label for="startYear" data-i18n="filterStart">Start year</label>
        <select id="startYear" name="from">{start_year_options}</select>
      </div>
      <div class="filter-field">
        <label for="endYear" data-i18n="filterEnd">End year</label>
        <select id="endYear" name="to">{end_year_options}</select>
      </div>
      <button type="button" id="resetYears" data-i18n="filterReset">Reset</button>
      <span id="rangeStatus"></span>
    </form>
    {research_overview}
    <section class="keyword-section" id="keywords-convention">
      <h2 data-i18n="keywordsTitle">Keywords Convention</h2>
      <p data-i18n="keywordsIntro">These clickable keyword tags define the AI-specific convention used to scan, filter, and extend this collection.</p>
      <div class="keyword-grid">{keyword_html}</div>
      <p class="keyword-filter-status" id="keywordFilterStatus">Selected keyword: all | Matching papers: {len(selected):,} papers</p>
    </section>
    <h2 data-i18n="taxonomyTitle">Taxonomy</h2>
    <p id="taxonomyTotalSummary"><strong>Total selected papers:</strong> {len(selected):,} papers; <strong>Categories:</strong> {len(stats)} categories.</p>
    <p data-i18n="taxonomyIntro">Each taxonomy section lists papers with publication year, venue, citation count, influential citations, score, keywords, key idea, strengths, research-focused limitations, and paper links. Sections are collapsed by default.</p>
    <div class="figures">
      <figure class="chart-figure">
        <img id="categoryDistributionChart" src="assets/periods/category_distribution_{START_YEAR}_{END_YEAR}.svg" alt="Category distribution chart for {YEAR_RANGE_TEXT}">
        <figcaption id="categoryChartCaption">Category distribution ({YEAR_RANGE_TEXT})</figcaption>
      </figure>
      <figure class="chart-figure">
        <img id="yearlyCitationsChart" src="assets/periods/yearly_citations_{START_YEAR}_{END_YEAR}.svg" alt="Yearly citation chart for {YEAR_RANGE_TEXT}">
        <figcaption id="citationChartCaption">Yearly citation mass ({YEAR_RANGE_TEXT})</figcaption>
      </figure>
    </div>
    {sections}
  </main>
{year_script}
</body>
</html>
"""
    (DOCS_DIR / "index.html").write_text(html_doc, encoding="utf-8")
    (DOCS_DIR / ".nojekyll").write_text("", encoding="utf-8")


def reference_line(paper):
    venue = paper["venue"] or "Unknown venue"
    return f"{paper['authors'] or 'Unknown authors'}. ({paper['year']}). {paper['title']}. {venue}. {paper['url'] or paper['semanticScholarUrl']}"


def review_sections(selected, korean=False):
    stats = year_stats(selected)
    cats = category_stats(selected)
    total_cites = sum(p["citationCount"] for p in selected)
    top_cited = sorted(selected, key=lambda p: p["citationCount"], reverse=True)[:12]
    top_scored = sorted(selected, key=lambda p: p["importanceScore"], reverse=True)[:12]
    peak_year = max(stats, key=lambda y: stats[y]["citations"])
    leading_cat, leading_count = cats.most_common(1)[0]
    if korean:
        title = f"{YEAR_RANGE_TEXT} AI 연구 동향: 공개 메타데이터 기반 citation-ranked 리뷰"
        abstract = (
            f"이 리뷰 초안은 {START_YEAR}년부터 {END_YEAR}년까지 AI 연구를 연도별 최대 {CANDIDATES_PER_YEAR:,}편의 후보 논문으로 조사하고, "
            f"각 연도에서 citation이 높은 {TARGET_PER_YEAR:,}편씩 총 {TARGET_TOTAL:,}편을 선정해 taxonomy-first 방식으로 분석한다. "
            "선정과 분류는 Semantic Scholar 공개 메타데이터, 명시적 AI 관련성 필터, DOI/arXiv/PubMed/CorpusId/paperId 중복 제거, 인용수 정렬을 사용했다."
        )
        methods = (
            "각 연도에 대해 AI, machine learning, deep learning, foundation model, LLM, NLP, computer vision, reinforcement learning, generative AI, graph learning, multimodal learning, trustworthy AI, AI for science 관련 질의를 보냈다. "
            f"제목/초록/venue에서 AI 관련 표현이 확인되는 record만 유지하고, 연도별 최대 {CANDIDATES_PER_YEAR:,}편을 citation count 기준 후보군으로 저장했다. 최종 목록은 각 연도별 citation count 상위 {TARGET_PER_YEAR:,}편이다."
        )
        findings = [
            f"선정 논문 {len(selected):,}편은 총 {total_cites:,}회의 인용을 포함하며, citation mass가 가장 큰 연도는 {peak_year}년이다.",
            f"가장 큰 분류는 {KOREAN_CATEGORY_NAMES.get(leading_cat, leading_cat)}({leading_count}편)이다.",
            f"{YEAR_RANGE_TEXT} 구간은 전통적 machine learning, data mining, pattern recognition부터 vision transformer, LLM, RAG, self-supervised learning, diffusion/generative AI, trustworthy AI, AI for science까지 이어지는 흐름을 보인다.",
            f"{END_YEAR}년 논문은 아직 인용 누적 시간이 짧으므로 최신성과 영향력은 분리해서 읽어야 한다.",
        ]
        caveat = "이 문서는 PDF 전문 기반 systematic review가 아니라 공개 메타데이터 기반의 citation map이다. 인용수는 영향력을 보여주지만 방법론적 품질, 안전성, 재현성을 직접 보장하지 않는다."
        conclusion = "AI 연구는 대규모 모델과 데이터 중심 방법론을 축으로 확장하면서, 동시에 안전성, 설명가능성, domain translation, 재현성 검증을 더 강하게 요구받고 있다."
    else:
        title = f"AI Research from {START_YEAR} to {END_YEAR}: A Metadata-Driven Citation Map"
        abstract = (
            f"This draft review maps AI research from {START_YEAR} through {END_YEAR}, investigating up to {CANDIDATES_PER_YEAR:,} candidate papers per year "
            f"from free public Semantic Scholar metadata and selecting the top {TARGET_PER_YEAR:,} papers from each year by citation count ({TARGET_TOTAL:,} papers total). "
            "The resulting collection is organized by research taxonomy and enriched with deterministic key ideas, strengths, limitations, and AI-specific keyword tags."
        )
        methods = (
            f"For each year, broad AI-oriented queries were sent to Semantic Scholar Academic Graph bulk search. Records were retained when title, abstract, or venue metadata matched explicit AI relevance expressions, deduplicated by DOI, arXiv, PubMed, CorpusId, paperId, then title, and reduced to at most {CANDIDATES_PER_YEAR:,} candidates per year by citation count. The final collection selects the top {TARGET_PER_YEAR:,} papers within each publication year by citation count, with influential citation count and metadata importance score retained as tie-breakers and audit signals."
        )
        findings = [
            f"The {len(selected):,} selected papers account for {total_cites:,} citations in the selected set, with the largest citation mass in {peak_year}.",
            f"The largest category is {leading_cat} ({leading_count} papers), reflecting the influence of large-scale model and representation work.",
            "The period connects vision transformers, LLMs, retrieval augmentation, self-supervised learning, diffusion/generative AI, trustworthy AI, and AI-for-science workflows.",
            f"Papers from {END_YEAR} are structurally citation-disadvantaged because the year is partial and citation accumulation is still immature.",
        ]
        caveat = "This is a metadata-driven citation map rather than a full systematic review of every PDF. Citation count is a useful influence signal, but it is not a direct measure of methodological quality, safety, or reproducibility."
        conclusion = "AI research is expanding around large-scale models and data-centric methods while simultaneously increasing pressure for safety, interpretability, domain validation, and reproducible evaluation."
    category_lines = [f"{cat}: {count}" for cat, count in cats.most_common()]
    year_lines = [
        f"{year}: {stats[year]['count']} selected papers, {stats[year]['citations']:,} citations, top selected paper: {stats[year]['top']['title']}"
        for year in YEARS
        if year in stats
    ]
    return {
        "title": title,
        "abstract": abstract,
        "methods": methods,
        "findings": findings,
        "category_lines": category_lines,
        "year_lines": year_lines,
        "top_cited": top_cited,
        "top_scored": top_scored,
        "caveat": caveat,
        "conclusion": conclusion,
    }


def html_ranked_table(rows, metric):
    label = "Citations" if metric == "citations" else "Importance"
    out = [f"<table><thead><tr><th>Year</th><th>Rank</th><th>Paper</th><th>{label}</th><th>Category</th></tr></thead><tbody>"]
    for paper in rows:
        value = paper["citationCount"] if metric == "citations" else paper["importanceScore"]
        link = f'<a href="{html.escape(paper["url"] or paper["semanticScholarUrl"])}">{html.escape(paper["title"])}</a>'
        out.append(f"<tr><td>{paper['year']}</td><td>{paper['rank']}</td><td>{link}</td><td>{value}</td><td>{html.escape(paper['category'])}</td></tr>")
    out.append("</tbody></table>")
    return "\n".join(out)


def write_review_html(selected, korean=False):
    content = review_sections(selected, korean=korean)
    lang = "ko" if korean else "en"
    name = "review_ko.html" if korean else "review_en.html"
    html_doc = f"""<!doctype html>
<html lang="{lang}">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{html.escape(content['title'])}</title>
  <style>
    body {{ font-family: Georgia, 'Noto Serif KR', serif; max-width: 920px; margin: 40px auto; padding: 0 22px; line-height: 1.72; color:#172033; }}
    h1 {{ line-height:1.15; }}
    h2 {{ margin-top:34px; }}
    .abstract {{ background:#f5f7fb; border-left:4px solid #2563eb; padding:14px 18px; }}
    table {{ width:100%; border-collapse:collapse; margin:16px 0; }}
    th, td {{ border-bottom:1px solid #d9dee8; padding:8px; vertical-align:top; text-align:left; }}
    th {{ background:#f4f6fa; }}
  </style>
</head>
<body>
  <h1>{html.escape(content['title'])}</h1>
  <p><strong>Generated:</strong> {date.today().isoformat()} &middot; <strong>Dataset:</strong> {len(selected):,} selected papers</p>
  <h2>Abstract</h2>
  <p class="abstract">{html.escape(content['abstract'])}</p>
  <h2>1. Scope and Methods</h2>
  <p>{html.escape(content['methods'])}</p>
  <h2>2. Key Findings</h2>
  <ul>{''.join(f'<li>{html.escape(item)}</li>' for item in content['findings'])}</ul>
  <h2>3. Taxonomy</h2>
  <ul>{''.join(f'<li>{html.escape(item)}</li>' for item in content['category_lines'])}</ul>
  <h2>4. Year-by-Year Coverage</h2>
  <ul>{''.join(f'<li>{html.escape(item)}</li>' for item in content['year_lines'])}</ul>
  <h2>5. Top Papers by Citation Count</h2>
  {html_ranked_table(content['top_cited'], 'citations')}
  <h2>6. Top Papers by Metadata Importance Score</h2>
  {html_ranked_table(content['top_scored'], 'importance')}
  <h2>7. Limitations</h2>
  <p>{html.escape(content['caveat'])}</p>
  <h2>8. Conclusion</h2>
  <p>{html.escape(content['conclusion'])}</p>
  <h2>Selected References</h2>
  <ol>{''.join(f'<li>{html.escape(reference_line(paper))}</li>' for paper in content['top_cited'])}</ol>
</body>
</html>
"""
    (PAPER_DIR / name).write_text(html_doc, encoding="utf-8")


def write_review_docx(selected):
    content = review_sections(selected, korean=False)
    doc = Document()
    doc.add_heading(content["title"], level=0)
    doc.add_paragraph(f"Generated: {date.today().isoformat()} | Dataset: {len(selected):,} selected papers")
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(content["abstract"])
    doc.add_heading("1. Scope and Methods", level=1)
    doc.add_paragraph(content["methods"])
    doc.add_heading("2. Key Findings", level=1)
    for item in content["findings"]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("3. Taxonomy", level=1)
    for item in content["category_lines"]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("4. Year-by-Year Coverage", level=1)
    for item in content["year_lines"]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("5. Top Papers by Citation Count", level=1)
    for paper in content["top_cited"]:
        doc.add_paragraph(f"{paper['year']} #{paper['rank']}: {paper['title']} ({paper['citationCount']:,} citations)", style="List Number")
    doc.add_heading("6. Limitations", level=1)
    doc.add_paragraph(content["caveat"])
    doc.add_heading("7. Conclusion", level=1)
    doc.add_paragraph(content["conclusion"])
    doc.save(PAPER_DIR / "review_en.docx")


def write_curation_method_html(markdown_text):
    body = []
    for line in markdown_text.splitlines():
        if line.startswith("# "):
            body.append(f"<h1>{html.escape(line[2:])}</h1>")
        elif line.startswith("## "):
            body.append(f"<h2>{html.escape(line[3:])}</h2>")
        elif line.startswith("- "):
            body.append(f"<li>{html.escape(line[2:])}</li>")
        elif line.strip():
            body.append(f"<p>{html.escape(line)}</p>")
    html_doc = f"""<!doctype html>
<html lang="en">
<head><meta charset="utf-8"><meta name="viewport" content="width=device-width, initial-scale=1"><title>Curation Method</title></head>
<body>{''.join(body)}</body>
</html>
"""
    (PAPER_DIR / "curation_method.html").write_text(html_doc, encoding="utf-8")


def write_project_files(selected, candidates):
    citation = f"""cff-version: 1.2.0
title: "Awesome Test: A Metadata-Driven Citation Map of AI Research, {YEAR_RANGE_TEXT}"
message: "If you use this curated dataset or review draft, please cite this repository."
type: dataset
authors:
  - name: "Honggi"
repository-code: "https://github.com/honggi82/awesome-test"
date-released: "{date.today().isoformat()}"
license: "CC-BY-4.0"
keywords:
  - "artificial intelligence"
  - "machine learning"
  - "deep learning"
  - "foundation models"
  - "bibliometrics"
"""
    (ROOT / "CITATION.cff").write_text(citation, encoding="utf-8")
    (ROOT / "LICENSE").write_text("CC-BY-4.0 for text and metadata curation; upstream paper metadata belongs to original sources.\n", encoding="utf-8")
    (ROOT / ".gitignore").write_text("__pycache__/\n*.pyc\n.tools/\ndata/cache/\n.playwright-cli/\noutput/playwright/\n", encoding="utf-8")
    publish = r"""@echo off
setlocal
cd /d "%~dp0"

set "GH_EXE=%~dp0.tools\gh\bin\gh.exe"
if not exist "%GH_EXE%" if exist "%~dp0..\awesome-BCI\.tools\gh\bin\gh.exe" set "GH_EXE=%~dp0..\awesome-BCI\.tools\gh\bin\gh.exe"
if not exist "%GH_EXE%" set "GH_EXE=gh"

"%GH_EXE%" auth status
if errorlevel 1 (
  echo.
  echo GitHub login is required. Run:
  echo   "%GH_EXE%" auth login --hostname github.com --web --scopes repo
  exit /b 1
)

"%GH_EXE%" repo view honggi82/awesome-test >nul 2>nul
if errorlevel 1 (
  "%GH_EXE%" repo create honggi82/awesome-test --public --description "Awesome Test: metadata-driven AI paper curation, 2000-2026" --source . --remote origin --push
) else (
  git remote set-url origin https://github.com/honggi82/awesome-test.git
  git push -u origin main
)
if errorlevel 1 exit /b %errorlevel%

"%GH_EXE%" api repos/honggi82/awesome-test/pages -X POST -f "source[branch]=main" -f "source[path]=/docs" >nul 2>nul
if errorlevel 1 (
  "%GH_EXE%" api repos/honggi82/awesome-test/pages -X PUT -f "source[branch]=main" -f "source[path]=/docs" >nul 2>nul
)

echo.
echo Done: https://github.com/honggi82/awesome-test
echo Pages: https://honggi82.github.io/awesome-test/
"""
    (ROOT / "publish_to_github.bat").write_text(publish, encoding="utf-8")
    method = f"""# Curation Method

## Scope

- Topic: AI research across machine learning, deep learning, foundation models, NLP, vision, reinforcement learning, generative AI, trustworthy AI, graph learning, multimodal learning, and AI for science.
- Period: {YEAR_RANGE_TEXT}.
- Candidate target: up to {CANDIDATES_PER_YEAR:,} papers per year.
- Final selection: top {TARGET_PER_YEAR:,} papers per year by citation count from the audited yearly candidate pools ({TARGET_TOTAL:,} papers total).

## Data Source

Metadata comes from the free public Semantic Scholar Academic Graph bulk search endpoint. OpenAlex was used only as a public concept/subfield reference for the AI topic (`Artificial intelligence`, `C154945302`; Artificial Intelligence subfield `1702`).

## Ranking

Records are filtered to the requested publication year, screened for explicit AI relevance in title/abstract/venue metadata, deduplicated by DOI, arXiv, PubMed, CorpusId, paperId, and normalized title, then ranked by citation count. The selected set keeps the top {TARGET_PER_YEAR:,} papers in each year. Influential citation count and a deterministic metadata importance score are retained as audit fields.

## Enrichment

Taxonomy, key ideas, strengths, limitations, method tags, and keyword convention tags are generated with deterministic rules from public metadata. No paid API, paid LLM, paid translation, or paid compute was used.
"""
    (PAPER_DIR / "curation_method.md").write_text(method, encoding="utf-8")
    write_curation_method_html(method)


def copy_public_assets():
    for filename in (CANDIDATES_CSV, PERIOD_ANALYSIS_JSON):
        (DOCS_DIR / "data" / filename).unlink(missing_ok=True)
    for filename in (PAPERS_CSV, TAXONOMY_CSV, PAPER_LOCALIZATIONS_JSON):
        shutil.copyfile(DATA_DIR / filename, DOCS_DIR / "data" / filename)
    if (DATA_DIR / GITHUB_LINKS_JSON).exists():
        shutil.copyfile(DATA_DIR / GITHUB_LINKS_JSON, DOCS_DIR / "data" / GITHUB_LINKS_JSON)
    if (DATA_DIR / LINK_AUDIT_JSON).exists():
        shutil.copyfile(DATA_DIR / LINK_AUDIT_JSON, DOCS_DIR / "data" / LINK_AUDIT_JSON)
    shutil.copyfile(PAPER_DIR / "review_en.html", DOCS_DIR / "paper" / "review_en.html")
    shutil.copyfile(PAPER_DIR / "review_ko.html", DOCS_DIR / "paper" / "review_ko.html")
    shutil.copyfile(PAPER_DIR / "curation_method.html", DOCS_DIR / "paper" / "curation_method.html")


def main():
    for path in (DATA_DIR, DOCS_DIR, PAPER_DIR):
        path.mkdir(exist_ok=True)
    if "--reuse-candidates" in sys.argv:
        selected, selected_by_year, candidates_by_year = reuse_existing_candidates()
    else:
        selected, selected_by_year, candidates_by_year = collect_papers()
    candidates = write_data(selected, selected_by_year, candidates_by_year)
    write_taxonomy_dataset(selected)
    write_period_analysis(selected)
    write_readme(selected, candidates)
    write_taxonomy_icons(selected)
    write_charts(selected)
    write_site(selected)
    write_review_html(selected, korean=False)
    write_review_html(selected, korean=True)
    write_review_docx(selected)
    write_project_files(selected, candidates)
    copy_public_assets()
    print(f"[done] generated {len(selected):,} selected papers from {len(candidates):,} candidates", flush=True)


# ---------------------------------------------------------------------------
# Awesome Trade overrides
#
# The base generator came from the local awesome-ai/awesome-test workflow. Keep
# the tested data/site machinery, but replace the domain model here so the
# generated repository is explicitly about stock investment and AI trading.
# ---------------------------------------------------------------------------

PROJECT_OWNER = "honggi82"
PROJECT_REPO = "awesome-trade"
PROJECT_TITLE = "Awesome Trade"
PROJECT_GITHUB_URL = f"https://github.com/{PROJECT_OWNER}/{PROJECT_REPO}"
PROJECT_PAGES_URL = f"https://{PROJECT_OWNER}.github.io/{PROJECT_REPO}/"
PROJECT_TOPIC = "stock investment and AI-driven trading research"
PROJECT_DESCRIPTION = (
    f"A taxonomy-first, citation-ranked map of stock investment and AI-driven "
    f"trading research from {START_YEAR} through {END_YEAR}."
)
SKILL2_PROVENANCE_JSON = "paper_curation_skill2_provenance.json"

QUERIES = [
    "stock market prediction",
    "stock return prediction",
    "stock price prediction",
    "equity market forecasting",
    "financial market prediction",
    "financial time series forecasting",
    "machine learning stock market",
    "deep learning stock prediction",
    "neural networks stock market",
    "support vector machine stock prediction",
    "reinforcement learning trading",
    "algorithmic trading",
    "quantitative trading",
    "portfolio optimization stock market",
    "asset pricing stock returns",
    "factor investing equity",
    "technical analysis stock market",
    "investor sentiment stock returns",
    "news sentiment stock market",
    "high frequency trading stock market",
    "market microstructure trading",
    "volatility forecasting stock market",
    "risk management portfolio stock",
    "large language models finance trading",
    "stock returns",
    "equity returns",
    "financial economics stock market",
    "investment management equity",
    "securities market",
    "capital markets investment",
    "market efficiency stock returns",
    "mutual fund performance stocks",
    "hedge fund performance equity",
    "fund performance stock market",
    "options market stock returns",
    "derivatives stock market",
    "risk premium stock returns",
    "equity factor models",
    "market anomalies stock returns",
    "momentum investing stock market",
    "value investing stock market",
    "earnings announcements stock returns",
    "analyst forecasts stock returns",
    "price discovery stock market",
    "stock exchange market efficiency",
    "securities trading investment",
    "equity portfolio performance",
    "trading strategies stock market",
]

AI_RELEVANCE_PATTERNS = [
    r"\bstock market",
    r"\bstock price",
    r"\bstock prices",
    r"\bstock return",
    r"\bstock returns",
    r"\bcommon stock\b",
    r"\bstock investment\b",
    r"\bequity market",
    r"\bequity markets",
    r"\bequity return",
    r"\bequity returns",
    r"\bequities\b",
    r"\bshare price",
    r"\bfinancial market",
    r"\bfinancial markets",
    r"\bcapital market",
    r"\bcapital markets",
    r"\bsecurities market",
    r"\bsecurities trading",
    r"\bstock exchange",
    r"\bbehavioral finance\b",
    r"\binvestment management\b",
    r"\basset pricing\b",
    r"\breturn predictability\b",
    r"\brisk premium\b",
    r"\basset price\b",
    r"\basset prices\b",
    r"\bfactor investing\b",
    r"\bfactor model",
    r"\bequity factor",
    r"\bmarket anomal",
    r"\bmarket efficiency\b",
    r"\bprice discovery\b",
    r"\bmarket liquidity\b",
    r"\bliquidity risk\b",
    r"\bportfolio",
    r"\basset allocation\b",
    r"\bmutual fund",
    r"\bhedge fund",
    r"\bfund performance\b",
    r"\bETF\b",
    r"\bexchange[- ]traded\b",
    r"\btrading volume\b",
    r"\btrading activity\b",
    r"\btrading simulation\b",
    r"\bequity trading\b",
    r"\bstock trading\b",
    r"\bfinancial trading\b",
    r"\binformed trading\b",
    r"\btrading strateg",
    r"\balgorithmic trading\b",
    r"\bquantitative trading\b",
    r"\bhigh[- ]frequency trading\b",
    r"\bmarket microstructure\b",
    r"\blimit order book\b",
    r"\border flow\b",
    r"\btransaction costs?\b",
    r"\btechnical analysis\b",
    r"\bfundamental analysis\b",
    r"\bvolatility forecasting\b",
    r"\brealized volatility\b",
    r"\breturn volatility\b",
    r"\bgarch\b",
    r"\bvalue at risk\b",
    r"\bfinancial risk\b",
    r"\bfinancial time series\b",
    r"\bfinancial economics\b",
    r"\bearnings announcement",
    r"\banalyst forecast",
    r"\boptions market\b",
    r"\bderivative",
    r"\binvestor sentiment\b",
    r"\bnews sentiment\b",
    r"\balternative data\b",
    r"\bfinancial forecasting\b",
    r"\bfinancial prediction",
    r"\bfinancial econom",
    r"\bfinancial signal",
    r"\bfinancial distress\b",
    r"\bequity price",
    r"\bequity prices",
    r"\bmarket timing\b",
    r"\binvestment performance\b",
    r"\bactive investing\b",
    r"\blow[- ]latency trading\b",
    r"\bnon[- ]synchronous trading\b",
    r"\boptimal execution\b",
    r"\bbank equity\b",
    r"\breturn on equity\b",
    r"\bWall Street\b",
    r"\bstock investing\b",
    r"\bequity investing\b",
    r"\bS&P\s*500\b",
    r"\bNASDAQ\b",
    r"\bNYSE\b",
    r"\bDJIA\b",
    r"\bBlack Monday\b",
]

TITLE_RELEVANCE_PATTERNS = AI_RELEVANCE_PATTERNS + [
    r"\bstock\b",
    r"\bstocks\b",
    r"\bportfolios?\b",
]

CORE_MARKET_TITLE_PATTERNS = [
    r"\bstock\b",
    r"\bstocks\b",
    r"\bstock market\b",
    r"\bstock price",
    r"\bstock return",
    r"\bequity market",
    r"\bequity return",
    r"\bfinancial market",
    r"\bcapital market",
    r"\bfinancial time[- ]series\b",
    r"\bS&P\s*500\b",
    r"\bNASDAQ\b",
    r"\bNYSE\b",
    r"\bDJIA\b",
]

OFF_TOPIC_TITLE_PATTERNS = [
    r"\bhealth equity\b",
    r"\bdigital health\b",
    r"\bhealthcare\b",
    r"\bclinical\b",
    r"\bmedical\b",
    r"\bbiomedical\b",
    r"\bbiohacking\b",
    r"\bbioinnovation\b",
    r"\bpharmaceutical\b",
    r"\benergy trading\b",
    r"\bcarbon trading\b",
    r"\bsmart grid\b",
    r"\boil\s*&\s*gas\b",
    r"\bfood insecurity\b",
]

FINANCE_FIELD_PATTERNS = [
    r"\beconomics\b",
    r"\bbusiness\b",
    r"\bfinance\b",
]

FINANCE_VENUE_PATTERNS = [
    r"\bjournal of finance\b",
    r"\bjournal of financial economics\b",
    r"\breview of financial studies\b",
    r"\bjournal of financial and quantitative analysis\b",
    r"\bjournal of banking\b",
    r"\bjournal of portfolio management\b",
    r"\bfinancial analysts journal\b",
    r"\bjournal of empirical finance\b",
    r"\bquantitative finance\b",
    r"\bfinancial\b",
    r"\bfinance\b",
    r"\beconomics\b",
]

BAD_TITLE_PATTERNS = [
    r"^group$",
    r"^editorial\b",
    r"^preface\b",
    r"^front matter\b",
    r"^proceedings of\b",
    r"^book review\b",
    r"^call for papers\b",
    r"^erratum\b",
    r"^corrigendum\b",
    r"^press releases?$",
]


def publication_type_values(paper):
    value = paper.get("publicationTypes") or []
    if isinstance(value, str):
        return [part.strip().lower() for part in re.split(r"[;,]", value) if part.strip()]
    return [str(item).lower() for item in value]


def fields_of_study_text(paper):
    values = []
    fields = paper.get("s2FieldsOfStudy") or []
    if isinstance(fields, str):
        values.append(fields)
    else:
        for item in fields:
            if isinstance(item, dict) and item.get("category"):
                values.append(str(item.get("category")))
            elif item:
                values.append(str(item))
    if paper.get("fieldsOfStudy"):
        values.append(str(paper.get("fieldsOfStudy")))
    return " ".join(values).lower()


def finance_field_match(paper):
    fields = fields_of_study_text(paper)
    return any(re.search(pattern, fields, re.I) for pattern in FINANCE_FIELD_PATTERNS)


def finance_venue_match(paper):
    venue = venue_name(paper)
    return any(re.search(pattern, venue, re.I) for pattern in FINANCE_VENUE_PATTERNS)


def core_market_title_match(paper):
    title = norm_text(paper.get("title"))
    return any(re.search(pattern, title, re.I) for pattern in CORE_MARKET_TITLE_PATTERNS)


def title_relevance_match(paper):
    title = norm_text(paper.get("title"))
    return any(re.search(pattern, title, re.I) for pattern in TITLE_RELEVANCE_PATTERNS)


def off_topic_match(paper):
    title = norm_text(paper.get("title"))
    if any(re.search(pattern, title, re.I) for pattern in OFF_TOPIC_TITLE_PATTERNS):
        return not core_market_title_match(paper)
    fields = fields_of_study_text(paper)
    if "medicine" in fields and not title_relevance_match(paper) and not finance_venue_match(paper):
        return True
    if "agricultural and food sciences" in fields and not title_relevance_match(paper) and not finance_venue_match(paper):
        return True
    return False

IMPORTANT_VENUES = [
    "journal of finance",
    "journal of financial economics",
    "review of financial studies",
    "journal of financial and quantitative analysis",
    "management science",
    "quantitative finance",
    "journal of banking and finance",
    "journal of portfolio management",
    "financial analysts journal",
    "journal of empirical finance",
    "journal of forecasting",
    "international journal of forecasting",
    "expert systems with applications",
    "decision support systems",
    "applied soft computing",
    "knowledge-based systems",
    "neurocomputing",
    "ieee access",
    "ieee transactions",
    "acm transactions",
    "kdd",
    "icml",
    "neurips",
]

CATEGORIES = [
    (
        "Asset Pricing and Return Predictability",
        [
            "asset pricing",
            "return predictability",
            "stock return",
            "expected return",
            "factor",
            "cross-section",
            "anomaly",
            "momentum",
            "value premium",
            "equity premium",
        ],
    ),
    (
        "Portfolio Optimization and Asset Allocation",
        [
            "portfolio",
            "asset allocation",
            "mean variance",
            "efficient frontier",
            "risk parity",
            "portfolio optimization",
            "portfolio selection",
            "allocation",
            "rebalancing",
        ],
    ),
    (
        "Machine Learning for Stock Prediction",
        [
            "machine learning",
            "support vector",
            "svm",
            "random forest",
            "xgboost",
            "boosting",
            "classification",
            "regression tree",
            "data mining",
            "kernel",
        ],
    ),
    (
        "Deep Learning and Financial Time Series",
        [
            "deep learning",
            "neural network",
            "lstm",
            "gru",
            "cnn",
            "rnn",
            "transformer",
            "attention",
            "representation learning",
            "financial time series",
        ],
    ),
    (
        "Reinforcement Learning and Algorithmic Trading",
        [
            "reinforcement learning",
            "q-learning",
            "policy gradient",
            "actor critic",
            "algorithmic trading",
            "trading strategy",
            "execution",
            "market making",
            "backtesting",
        ],
    ),
    (
        "Sentiment, News, and Alternative Data",
        [
            "sentiment",
            "news",
            "twitter",
            "social media",
            "text mining",
            "natural language",
            "nlp",
            "large language model",
            "alternative data",
            "earnings call",
        ],
    ),
    (
        "Market Microstructure and High-Frequency Trading",
        [
            "market microstructure",
            "high-frequency",
            "high frequency",
            "limit order book",
            "order flow",
            "liquidity",
            "bid ask",
            "price impact",
            "transaction cost",
        ],
    ),
    (
        "Risk, Volatility, and Forecast Evaluation",
        [
            "risk",
            "volatility",
            "garch",
            "var",
            "value at risk",
            "drawdown",
            "tail risk",
            "forecast evaluation",
            "out-of-sample",
            "stress testing",
        ],
    ),
    (
        "Behavioral Finance and Investor Decision-Making",
        [
            "behavioral finance",
            "investor behavior",
            "investor sentiment",
            "attention",
            "overreaction",
            "underreaction",
            "herding",
            "disposition effect",
            "market efficiency",
        ],
    ),
    (
        "General Finance, Surveys, and Trading Systems",
        [
            "survey",
            "review",
            "overview",
            "framework",
            "system",
            "investment",
            "finance",
            "trading",
            "stock",
            "equity",
        ],
    ),
]

KEYWORD_CONVENTION = [
    ("stock-prediction", "Stock price, return, direction, trend, or market-index prediction and forecasting.", "0f766e"),
    ("machine-learning", "SVMs, trees, boosting, random forests, kernels, data mining, and classical ML for markets.", "475569"),
    ("deep-learning", "Neural networks, LSTM/GRU/CNN, attention, transformers, and representation learning for financial time series.", "2563eb"),
    ("ai-based-trade", "AI/ML-driven stock trading and investment signals, including predictive models, NLP/LLMs, and automated strategies.", "16a34a"),
    ("portfolio", "Portfolio optimization, asset allocation, portfolio selection, risk parity, and rebalancing.", "7c2d12"),
    ("reinforcement-trading", "Reinforcement learning, algorithmic trading, execution, market making, and backtested trading strategies.", "dc2626"),
    ("sentiment-altdata", "News, social media, NLP, LLMs, earnings calls, and alternative data for stock investment.", "f59e0b"),
    ("high-frequency", "Market microstructure, limit order books, order flow, liquidity, and high-frequency trading.", "0891b2"),
    ("risk-volatility", "Volatility forecasting, Value at Risk, drawdowns, tail risk, stress testing, and forecast evaluation.", "be123c"),
    ("asset-pricing", "Asset pricing, equity factors, anomalies, return predictability, and cross-sectional stock returns.", "4f46e5"),
    ("behavioral-finance", "Investor sentiment, attention, behavioral bias, market efficiency, and decision-making.", "a855f7"),
    ("github", "Papers with an official GitHub or code repository link identified in the metadata audit.", "24292f"),
]
KEYWORD_COLORS = {keyword: color for keyword, _, color in KEYWORD_CONVENTION}

LANGUAGES = {"en": "English"}
UI_LABELS = {
    "en": {
        "papers": "papers",
        "categories": "categories",
        "overview": "Category Overview",
        "limitations": "Limitations",
        "analysis": "Selected-period analysis",
        "totalSelected": "Total selected papers",
        "categoryCount": "Categories",
        "keyIdea": "Key idea",
        "strengths": "Strengths",
        "paperLimitations": "Limitations",
    }
}

TAXONOMY_TRENDS = {
    "Asset Pricing and Return Predictability": [
        "Asset-pricing work connects stock returns to factors, anomalies, valuation signals, and changing market efficiency.",
        "Citation-ranked papers often define reusable signals or empirical designs that later trading and portfolio papers inherit.",
        "The central question is whether return predictability survives costs, multiple testing, and out-of-sample evaluation.",
    ],
    "Portfolio Optimization and Asset Allocation": [
        "Portfolio research turns return forecasts into allocation decisions under uncertainty, constraints, and risk budgets.",
        "The literature moves from mean-variance foundations toward robust, Bayesian, machine-learning, and risk-aware allocation.",
        "High-impact work is valuable when it clarifies how estimation error changes realized portfolio performance.",
    ],
    "Machine Learning for Stock Prediction": [
        "Classical machine-learning papers apply kernels, trees, boosting, ensembles, and data-mining methods to stock prediction.",
        "The main trend is a shift from single-indicator technical models toward richer feature sets and stricter validation.",
        "Useful results separate predictive accuracy from economically meaningful profitability after costs.",
    ],
    "Deep Learning and Financial Time Series": [
        "Deep-learning work uses recurrent, convolutional, attention, and transformer models to represent nonlinear market dynamics.",
        "Recent papers increasingly combine price series with text, order-book, macro, or cross-asset signals.",
        "The key methodological issue is controlling overfitting in nonstationary, low-signal financial data.",
    ],
    "Reinforcement Learning and Algorithmic Trading": [
        "Reinforcement-learning and algorithmic-trading papers frame trading as sequential decision-making under costs and risk.",
        "The area connects policy learning, execution, market making, and portfolio rebalancing.",
        "Robust evaluation needs realistic costs, liquidity constraints, and stress tests beyond benchmark backtests.",
    ],
    "Sentiment, News, and Alternative Data": [
        "Text and alternative-data research studies whether news, social media, filings, and language models improve stock decisions.",
        "The field has shifted from dictionary sentiment to NLP, representation learning, and LLM-assisted financial text analysis.",
        "Strong papers align timestamped information with tradable horizons to avoid look-ahead bias.",
    ],
    "Market Microstructure and High-Frequency Trading": [
        "Microstructure and high-frequency work studies liquidity, order flow, price impact, and short-horizon prediction.",
        "The literature is shaped by transaction costs, latency, execution risk, and limit-order-book dynamics.",
        "The most useful results make trading constraints explicit rather than treating prices as frictionless signals.",
    ],
    "Risk, Volatility, and Forecast Evaluation": [
        "Risk and volatility papers estimate uncertainty, tails, drawdowns, and forecast reliability for stock investment.",
        "The area links econometrics, machine learning, and practical portfolio risk controls.",
        "Evaluation quality depends on out-of-sample tests, regime sensitivity, and transparent loss functions.",
    ],
    "Behavioral Finance and Investor Decision-Making": [
        "Behavioral-finance research explains stock-market anomalies through investor attention, sentiment, bias, and limits to arbitrage.",
        "Citation-ranked work often supplies mechanisms that complement purely statistical prediction models.",
        "The strongest papers connect behavioral signals to implementable investment constraints.",
    ],
    "General Finance, Surveys, and Trading Systems": [
        "General papers include surveys, frameworks, benchmark discussions, and systems that organize stock-investment research.",
        "This category helps readers connect finance theory, empirical design, and AI-based trading practice.",
        "Metadata-driven ranking is useful as a map of influence but does not replace full-paper expert review.",
    ],
}

TAXONOMY_LIMITATIONS = {
    "Asset Pricing and Return Predictability": [
        "Return predictability can disappear after transaction costs, crowding, and multiple-testing correction.",
        "Citation impact can favor well-known anomalies even when live performance later weakens.",
        "Metadata cannot verify whether data-snooping controls were adequate.",
    ],
    "Portfolio Optimization and Asset Allocation": [
        "Allocation gains can be dominated by estimation error, turnover, short-sale constraints, and benchmark choice.",
        "Backtests may hide capacity and liquidity limits.",
        "Metadata cannot confirm implementation details such as costs, leverage, or rebalancing rules.",
    ],
    "Machine Learning for Stock Prediction": [
        "Predictive accuracy may not translate into risk-adjusted returns after costs.",
        "Feature leakage, survivorship bias, and nonstationarity are persistent concerns.",
        "Metadata-driven screening cannot audit the exact train/test split.",
    ],
    "Deep Learning and Financial Time Series": [
        "Deep models can overfit noisy, nonstationary, and low signal-to-noise market data.",
        "Compute-heavy models may be hard to reproduce without code and data access.",
        "Reported gains need robustness checks across regimes and markets.",
    ],
    "Reinforcement Learning and Algorithmic Trading": [
        "Simulated rewards can diverge from executable trading under slippage, latency, and market impact.",
        "Policies can be brittle during regime shifts or stress periods.",
        "Backtest-only evidence needs live or paper-trading validation.",
    ],
    "Sentiment, News, and Alternative Data": [
        "Timestamp alignment and look-ahead bias are critical for text and alternative-data studies.",
        "Sentiment signals can decay quickly as information becomes crowded.",
        "Data licensing and coverage bias can limit reproducibility.",
    ],
    "Market Microstructure and High-Frequency Trading": [
        "High-frequency results depend heavily on market access, latency, fees, and order-book reconstruction quality.",
        "Public metadata rarely reveals execution assumptions.",
        "Short-horizon predictability can vanish when capacity and competition are considered.",
    ],
    "Risk, Volatility, and Forecast Evaluation": [
        "Risk forecasts can fail under regime breaks and extreme events.",
        "Choice of loss function strongly affects model rankings.",
        "Citation counts do not prove suitability for a specific mandate or risk budget.",
    ],
    "Behavioral Finance and Investor Decision-Making": [
        "Behavioral mechanisms can be market-specific and time-varying.",
        "Sentiment proxies may capture several confounded channels.",
        "Trading implementation requires constraints beyond the behavioral explanation.",
    ],
    "General Finance, Surveys, and Trading Systems": [
        "Surveys and systems can dominate citations while empirical evidence remains mixed across markets.",
        "Broad claims need careful mapping to data, costs, and investment horizons.",
        "Metadata-driven ranking cannot replace expert reading of full papers and backtests.",
    ],
}

KOREAN_CATEGORY_NAMES = {
    "Asset Pricing and Return Predictability": "자산가격결정과 수익률 예측",
    "Portfolio Optimization and Asset Allocation": "포트폴리오 최적화와 자산배분",
    "Machine Learning for Stock Prediction": "머신러닝 기반 주식 예측",
    "Deep Learning and Financial Time Series": "딥러닝과 금융 시계열",
    "Reinforcement Learning and Algorithmic Trading": "강화학습과 알고리즘 트레이딩",
    "Sentiment, News, and Alternative Data": "감성, 뉴스, 대체데이터",
    "Market Microstructure and High-Frequency Trading": "시장 미시구조와 고빈도 거래",
    "Risk, Volatility, and Forecast Evaluation": "위험, 변동성, 예측 평가",
    "Behavioral Finance and Investor Decision-Making": "행동재무와 투자자 의사결정",
    "General Finance, Surveys, and Trading Systems": "일반 금융, 서베이, 거래 시스템",
}

CATEGORY_COLORS = {
    "Asset Pricing and Return Predictability": ("#4f46e5", "#c4b5fd"),
    "Portfolio Optimization and Asset Allocation": ("#7c2d12", "#fed7aa"),
    "Machine Learning for Stock Prediction": ("#475569", "#cbd5e1"),
    "Deep Learning and Financial Time Series": ("#2563eb", "#93c5fd"),
    "Reinforcement Learning and Algorithmic Trading": ("#dc2626", "#fca5a5"),
    "Sentiment, News, and Alternative Data": ("#f59e0b", "#fde68a"),
    "Market Microstructure and High-Frequency Trading": ("#0891b2", "#67e8f9"),
    "Risk, Volatility, and Forecast Evaluation": ("#be123c", "#f9a8d4"),
    "Behavioral Finance and Investor Decision-Making": ("#a855f7", "#f0abfc"),
    "General Finance, Surveys, and Trading Systems": ("#0f766e", "#99f6e4"),
}

KOREAN_CATEGORY_NAMES = {
    "Asset Pricing and Return Predictability": "자산가격결정과 수익률 예측",
    "Portfolio Optimization and Asset Allocation": "포트폴리오 최적화와 자산배분",
    "Machine Learning for Stock Prediction": "머신러닝 기반 주식 예측",
    "Deep Learning and Financial Time Series": "딥러닝과 금융 시계열",
    "Reinforcement Learning and Algorithmic Trading": "강화학습과 알고리즘 트레이딩",
    "Sentiment, News, and Alternative Data": "감성, 뉴스, 대체데이터",
    "Market Microstructure and High-Frequency Trading": "시장 미시구조와 고빈도 거래",
    "Risk, Volatility, and Forecast Evaluation": "위험, 변동성, 예측 평가",
    "Behavioral Finance and Investor Decision-Making": "행동재무와 투자자 의사결정",
    "General Finance, Surveys, and Trading Systems": "일반 금융, 서베이, 거래 시스템",
}


def is_relevant(paper):
    title = norm_text(paper.get("title"))
    if not title or title_is_bad(title):
        return False
    if off_topic_match(paper):
        return False
    publication_types = publication_type_values(paper)
    if any(kind in publication_types for kind in ("editorial", "news", "lettersandcomments")):
        return False
    title_hits, text_hits = relevance_counts(paper)
    if title_hits:
        return True
    if finance_venue_match(paper) and text_hits >= 1:
        return True
    if finance_field_match(paper) and text_hits >= 2:
        return True
    return False


def importance_score(paper):
    title = paper.get("title") or ""
    abstract = paper.get("abstract") or ""
    venue = venue_name(paper)
    text = f"{title} {abstract} {venue}".lower()
    citations = int(paper.get("citationCount") or 0)
    influential = int(paper.get("influentialCitationCount") or 0)
    title_hits, text_hits = relevance_counts(paper)
    query_hits = len(paper.get("sourceQueries") or [])
    score = math.log1p(citations) * 22.0 + math.log1p(influential) * 14.0
    reasons = [f"citations={citations}", f"influential={influential}"]
    if any(v in venue.lower() for v in IMPORTANT_VENUES):
        score += 10
        reasons.append("recognized finance/AI venue")
    if re.search(r"\b(review|survey|systematic review|benchmark|dataset)\b", text):
        score += 5
        reasons.append("survey/benchmark signal")
    if re.search(r"\b(out-of-sample|backtest|transaction cost|trading strategy|portfolio|asset pricing|volatility|risk)\b", text):
        score += 5
        reasons.append("investment validation signal")
    if re.search(r"\b(machine learning|deep learning|reinforcement learning|neural|transformer|support vector|random forest|xgboost)\b", text):
        score += 4
        reasons.append("AI/ML method signal")
    if paper.get("openAccessPdf"):
        score += 1
        reasons.append("open PDF metadata")
    if title_hits:
        score += min(8, title_hits * 2)
        reasons.append(f"title trade matches={title_hits}")
    if text_hits:
        score += min(8, text_hits)
        reasons.append(f"trade term matches={text_hits}")
    if query_hits > 1:
        score += min(6, query_hits)
        reasons.append(f"query hits={query_hits}")
    return round(score, 3), "; ".join(reasons)


def category_for(paper):
    text = text_for(paper)
    priority = [
        ("Sentiment, News, and Alternative Data", r"\b(sentiment|news|twitter|social media|text mining|natural language|nlp|language model|llm|alternative data|earnings call)\b"),
        ("Market Microstructure and High-Frequency Trading", r"\b(microstructure|high[- ]frequency|limit order book|order flow|liquidity|bid[- ]ask|price impact|transaction cost)\b"),
        ("Reinforcement Learning and Algorithmic Trading", r"\b(reinforcement|q-learning|policy gradient|actor critic|algorithmic trading|trading strategy|market making|execution|backtest)\b"),
        ("Portfolio Optimization and Asset Allocation", r"\b(portfolio|asset allocation|mean[- ]variance|efficient frontier|risk parity|rebalanc|portfolio selection)\b"),
        ("Risk, Volatility, and Forecast Evaluation", r"\b(volatility|garch|value at risk|var\b|tail risk|drawdown|forecast evaluation|out-of-sample|stress test)\b"),
        ("Deep Learning and Financial Time Series", r"\b(deep learning|neural network|lstm|gru|cnn|rnn|transformer|attention|representation learning)\b"),
        ("Machine Learning for Stock Prediction", r"\b(machine learning|support vector|svm|random forest|xgboost|boosting|data mining|classification|kernel)\b"),
        ("Asset Pricing and Return Predictability", r"\b(asset pricing|return predictability|stock return|expected return|factor|cross-section|anomal|momentum|equity premium)\b"),
        ("Behavioral Finance and Investor Decision-Making", r"\b(behavioral|investor sentiment|investor attention|herding|overreaction|underreaction|disposition effect|market efficiency)\b"),
    ]
    for category, pattern in priority:
        if re.search(pattern, text, re.I):
            return category
    best = ("General Finance, Surveys, and Trading Systems", 0)
    for category, terms in CATEGORIES:
        score = sum(1 for term in terms if term in text)
        if score > best[1]:
            best = (category, score)
    return best[0]


def method_tags(paper):
    text = text_for(paper)
    rules = [
        ("survey/review", r"\b(review|survey|overview)\b"),
        ("forecasting", r"\b(forecast|prediction|predictability|out-of-sample)\b"),
        ("portfolio", r"\b(portfolio|asset allocation|rebalanc|mean[- ]variance)\b"),
        ("asset-pricing", r"\b(asset pricing|factor|anomaly|cross-section|return predictability)\b"),
        ("machine-learning", r"\b(machine learning|support vector|svm|random forest|xgboost|boosting|kernel)\b"),
        ("deep-learning", r"\b(deep learning|neural|lstm|gru|cnn|rnn|transformer|attention)\b"),
        ("reinforcement-learning", r"\b(reinforcement|q-learning|policy|actor critic)\b"),
        ("sentiment/text", r"\b(sentiment|news|text|nlp|language model|llm|social media)\b"),
        ("microstructure", r"\b(microstructure|high[- ]frequency|order book|order flow|liquidity)\b"),
        ("risk/volatility", r"\b(risk|volatility|garch|value at risk|drawdown|tail)\b"),
    ]
    tags = [name for name, pattern in rules if re.search(pattern, text, re.I)]
    return tags[:6] or ["metadata-ranked"]


def keyword_tags(paper, category=None):
    text = text_for(paper)
    raw_text = f"{paper.get('title', '')} {paper.get('abstract', '')} {paper.get('venue', '')}".lower()
    tags = set()
    if re.search(r"\b(stock price|stock return|stock market|equity market|forecast|prediction|predictability)\b", text, re.I):
        tags.add("stock-prediction")
    if re.search(r"\b(machine learning|support vector|svm|random forest|xgboost|boosting|kernel|data mining)\b", text, re.I):
        tags.add("machine-learning")
    if re.search(r"\b(deep learning|neural network|lstm|gru|cnn|rnn|transformer|attention)\b", text, re.I):
        tags.add("deep-learning")
    if re.search(r"\b(portfolio|asset allocation|mean[- ]variance|risk parity|rebalanc)\b", text, re.I):
        tags.add("portfolio")
    if re.search(r"\b(reinforcement|algorithmic trading|trading strategy|execution|market making|backtest)\b", text, re.I):
        tags.add("reinforcement-trading")
    if re.search(r"\b(sentiment|news|twitter|social media|text mining|nlp|language model|llm|alternative data|earnings call)\b", text, re.I):
        tags.add("sentiment-altdata")
    if re.search(r"\b(high[- ]frequency|microstructure|limit order book|order flow|liquidity|price impact)\b", text, re.I):
        tags.add("high-frequency")
    if re.search(r"\b(risk|volatility|garch|value at risk|tail risk|drawdown|stress test|forecast evaluation)\b", text, re.I):
        tags.add("risk-volatility")
    if re.search(r"\b(asset pricing|factor|anomaly|momentum|cross-section|equity premium|expected return)\b", text, re.I):
        tags.add("asset-pricing")
    if re.search(r"\b(behavioral|investor sentiment|attention|herding|overreaction|underreaction|market efficiency)\b", text, re.I):
        tags.add("behavioral-finance")
    ai_signal = re.search(
        r"\b(artificial intelligence|ai|machine learning|deep learning|neural network|lstm|gru|cnn|rnn|transformer|support vector|svm|random forest|xgboost|boosting|kernel|data mining|reinforcement|q-learning|policy gradient|actor critic|natural language processing|nlp|language model|large language model|llm|text mining)\b",
        raw_text,
        re.I,
    )
    trade_signal = re.search(
        r"\b(stock|equity|portfolio|asset allocation|investment|securities|financial market|capital market|stock market|asset pricing|return predictability|stock return|market return|stock price|price prediction|financial forecast|market forecast|stock forecast|financial time series|algorithmic trading|trading strategy|trading strategies|trading simulation|market making|execution|backtest|order book|high[- ]frequency|microstructure)\b",
        raw_text,
        re.I,
    )
    if ai_signal and trade_signal:
        tags.add("ai-based-trade")
    category = category or category_for(paper)
    defaults = {
        "Asset Pricing and Return Predictability": "asset-pricing",
        "Portfolio Optimization and Asset Allocation": "portfolio",
        "Machine Learning for Stock Prediction": "machine-learning",
        "Deep Learning and Financial Time Series": "deep-learning",
        "Reinforcement Learning and Algorithmic Trading": "reinforcement-trading",
        "Sentiment, News, and Alternative Data": "sentiment-altdata",
        "Market Microstructure and High-Frequency Trading": "high-frequency",
        "Risk, Volatility, and Forecast Evaluation": "risk-volatility",
        "Behavioral Finance and Investor Decision-Making": "behavioral-finance",
        "General Finance, Surveys, and Trading Systems": "stock-prediction",
    }
    if not tags and category in defaults:
        tags.add(defaults[category])
    ordered = [keyword for keyword, _, _ in KEYWORD_CONVENTION if keyword in tags]
    return ordered or ["stock-prediction"]


def research_limitations(paper, method_tag_list, keyword_list):
    category = paper.get("category") or category_for(paper)
    items = list(TAXONOMY_LIMITATIONS.get(category, TAXONOMY_LIMITATIONS["General Finance, Surveys, and Trading Systems"]))
    tags = set(method_tag_list + keyword_list)
    if "forecasting" in tags or "stock-prediction" in tags:
        items.append("Forecasting papers require careful checks for temporal leakage and transaction-cost sensitivity.")
    if "deep-learning" in tags:
        items.append("Deep-learning gains can be fragile under regime shifts and small changes in validation protocol.")
    if "reinforcement-trading" in tags:
        items.append("Trading-policy evidence should include realistic slippage, market impact, and risk constraints.")
    if "sentiment-altdata" in tags:
        items.append("Text and alternative-data signals need precise timestamping to avoid look-ahead bias.")
    return items[:3]


def localized_paper_fields(paper, language):
    return {
        "keyIdea": paper["keyIdea"],
        "strengths": paper["strengths"],
        "limitations": paper["limitations"],
    }


def enrich_paper(paper):
    base = dict(paper)
    category = category_for(base)
    methods = method_tags(base)
    keywords = keyword_tags(base, category)
    citations = int(base.get("citationCount") or 0)
    influential = int(base.get("influentialCitationCount") or 0)
    strengths = []
    if citations >= 100:
        strengths.append(f"high citation signal ({citations:,})")
    if influential >= 10:
        strengths.append(f"influential citation signal ({influential:,})")
    if "recognized finance/AI venue" in base.get("importanceReasons", ""):
        strengths.append("recognized finance/AI venue")
    if base.get("openAccessPdf"):
        strengths.append("open-access PDF metadata")
    if not strengths:
        strengths.append("selected by citation count from the audited stock-investment candidate pool")
    base["category"] = category
    base["methodTags"] = "; ".join(methods)
    base["keywordTags"] = "; ".join(keywords)
    base["keyIdea"] = first_sentence(
        base.get("abstract", ""),
        f"Positions {base.get('title', 'this paper')} within {category}.",
    )
    base["strengths"] = "; ".join(strengths[:4])
    base["limitations"] = "; ".join(research_limitations(base, methods, keywords))
    return base


def write_github_links(rows):
    links = {}
    for row in rows:
        github_url = official_github_url_from_paper_text(row)
        if github_url:
            links[title_key(row.get("title"))] = {
                "githubUrl": github_url,
                "githubOfficial": True,
                "mentionedInPaper": True,
                "source": "abstract metadata context",
            }
    payload = {
        "generated": date.today().isoformat(),
        "method": "official-looking GitHub URLs extracted only when paper abstract metadata mentions code, repository, software, model, or data context",
        "links": links,
    }
    DATA_DIR.mkdir(exist_ok=True)
    (DATA_DIR / GITHUB_LINKS_JSON).write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def write_link_audit(selected, candidates):
    def classify(value):
        value = norm_text(value)
        if not value:
            return "missing"
        if re.match(r"^https?://[^\s<>]+$", value):
            return "valid_url_syntax"
        return "invalid_url_syntax"

    fields = ["url", "semanticScholarUrl", "openAccessPdf", "githubUrl"]
    counts = {field: Counter() for field in fields}
    for row in selected:
        for field in fields:
            counts[field][classify(row.get(field))] += 1
    payload = {
        "generated": date.today().isoformat(),
        "scope": "selected papers",
        "audit_type": "URL syntax and metadata-presence audit; live HTTP status checks were not run to avoid unnecessary load on scholarly hosts",
        "selected_papers": len(selected),
        "candidate_records": len(candidates),
        "fields": {field: dict(counter) for field, counter in counts.items()},
    }
    (DATA_DIR / LINK_AUDIT_JSON).write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def trade_metadata(selected_count=None):
    return {
        "topic": PROJECT_TOPIC,
        "source": "Semantic Scholar Academic Graph bulk search",
        "generated": date.today().isoformat(),
        "years": YEARS,
        "candidate_pool_per_year": CANDIDATES_PER_YEAR,
        "selected_per_year": TARGET_PER_YEAR,
        "selected_total": selected_count if selected_count is not None else TARGET_TOTAL,
        "selection_scope": "top 100 papers per publication year by citation count from stock-investment and AI-trading candidate pools",
        "ranking": "citationCount desc, influentialCitationCount desc",
        "queries": QUERIES,
    }


def rewrite_data_metadata(selected_count):
    metadata = trade_metadata(selected_count)
    json_files = [DATA_DIR / PAPERS_JSON, DATA_DIR / CANDIDATES_JSON]
    json_files.extend(DATA_DIR / f"candidates_top1000_{year}.json" for year in YEARS)
    for path in json_files:
        if not path.exists():
            continue
        payload = json.loads(path.read_text(encoding="utf-8"))
        if isinstance(payload, dict) and "metadata" in payload:
            payload["metadata"] = metadata
            path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def render_category_chart(rows, start, end):
    data = category_stats(rows).most_common()
    width = 1000
    row_h = 44
    height = max(260, 96 + row_h * max(1, len(data)))
    max_value = max([value for _, value in data] or [1])
    bars = []
    if not data:
        bars.append('<text x="500" y="150" text-anchor="middle" fill="#64748b" font-size="24">No selected papers</text>')
    for i, (category, value) in enumerate(data):
        y = 72 + i * row_h
        bar_w = int((value / max_value) * 520)
        bars.append(f'<text x="32" y="{y + 24}" fill="#172033" font-size="15">{svg_text(category[:52])}</text>')
        bars.append(f'<rect x="420" y="{y + 4}" width="{bar_w}" height="26" rx="5" fill="{chart_color(i)}"/>')
        bars.append(f'<text x="{430 + bar_w}" y="{y + 24}" fill="#172033" font-size="15" font-weight="700">{value}</text>')
    return f"""<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" viewBox="0 0 {width} {height}" role="img" aria-label="Category distribution {start}-{end}">
<rect width="100%" height="100%" fill="#ffffff"/>
<text x="32" y="38" fill="#0f172a" font-size="24" font-weight="800">Trade Paper Taxonomy, {period_label(start, end)}</text>
<text x="32" y="60" fill="#64748b" font-size="14">Selected papers by category</text>
{''.join(bars)}
</svg>"""


def taxonomy_icon_svg(category):
    accent, soft = CATEGORY_COLORS.get(category, ("#0f766e", "#99f6e4"))
    label = category.split(" and ")[0][:30]
    return f"""<svg xmlns="http://www.w3.org/2000/svg" width="320" height="200" viewBox="0 0 320 200" role="img" aria-label="{svg_text(category)}">
<rect x="0" y="0" width="320" height="200" rx="18" fill="#f8fafc"/>
<path d="M42 142 L94 104 L137 120 L184 72 L230 94 L278 50" fill="none" stroke="{accent}" stroke-width="9" stroke-linecap="round" stroke-linejoin="round"/>
<circle cx="94" cy="104" r="13" fill="{soft}" stroke="{accent}" stroke-width="5"/>
<circle cx="184" cy="72" r="13" fill="{soft}" stroke="{accent}" stroke-width="5"/>
<circle cx="278" cy="50" r="13" fill="{soft}" stroke="{accent}" stroke-width="5"/>
<rect x="42" y="154" width="236" height="3" fill="#cbd5e1"/>
<text x="160" y="182" text-anchor="middle" fill="#172033" font-size="17" font-weight="800">{svg_text(label)}</text>
</svg>"""


def write_taxonomy_icons(selected):
    from PIL import Image, ImageDraw, ImageFont

    icon_dir = DOCS_DIR / "assets" / "taxonomy"
    icon_dir.mkdir(parents=True, exist_ok=True)
    try:
        font = ImageFont.truetype("arial.ttf", 24)
        small_font = ImageFont.truetype("arial.ttf", 18)
    except OSError:
        font = ImageFont.load_default()
        small_font = ImageFont.load_default()
    for category in category_stats(selected):
        slug = safe_slug(category)
        accent, soft = CATEGORY_COLORS.get(category, ("#0f766e", "#99f6e4"))
        write_svg(icon_dir / f"{slug}.svg", taxonomy_icon_svg(category))
        image = Image.new("RGB", (640, 400), "#f8fafc")
        draw = ImageDraw.Draw(image)
        draw.rounded_rectangle((8, 8, 632, 392), radius=32, fill="#f8fafc", outline="#d9dee8", width=2)
        points = [(72, 284), (188, 210), (270, 238), (374, 136), (470, 184), (568, 92)]
        draw.line(points, fill=accent, width=16, joint="curve")
        for x, y in points[1::2]:
            draw.ellipse((x - 22, y - 22, x + 22, y + 22), fill=soft, outline=accent, width=8)
        draw.line((72, 314, 568, 314), fill="#cbd5e1", width=5)
        title = category.replace(" and ", " & ")
        if len(title) > 42:
            title = title[:39] + "..."
        bbox = draw.textbbox((0, 0), title, font=font)
        draw.text(((640 - (bbox[2] - bbox[0])) / 2, 338), title, fill="#172033", font=font)
        draw.text((34, 34), "Awesome Trade", fill=accent, font=small_font)
        image.save(icon_dir / f"{slug}.png")


def research_overview_html():
    return """
    <section class="research-brief" id="researchBrief" aria-labelledby="research-timeline-title">
      <h2 id="research-timeline-title">Research Timeline</h2>
      <div class="timeline-copy">
        <p>The 2000-2026 trade corpus follows a shift from asset-pricing, technical-analysis, volatility, and portfolio research toward machine learning, deep learning, sentiment-aware investing, reinforcement learning, and high-frequency market microstructure. The selected papers should be read as an influence map of stock-investment research rather than as trading advice.</p>
        <p>Across the period, the core question remains stable: which signals survive realistic costs, changing regimes, and out-of-sample evaluation? AI-based stock investment papers add richer models and data sources, but the strongest work still makes finance constraints explicit.</p>
      </div>
      <h2>Research Insights</h2>
      <div class="research-insights">
        <article class="insight-box"><div class="insight-label">Prediction</div><h3>Forecast accuracy is only the first filter</h3><p>Stock prediction papers increasingly use ML and deep learning, but economic value depends on turnover, costs, capacity, and risk.</p><p class="insight-implication">Implication: compare predictive metrics with tradable portfolio outcomes.</p></article>
        <article class="insight-box"><div class="insight-label">Portfolio</div><h3>Allocation turns signals into decisions</h3><p>Portfolio papers show how estimation error, constraints, and risk budgets can dominate raw alpha signals.</p><p class="insight-implication">Implication: robust allocation is often as important as model accuracy.</p></article>
        <article class="insight-box"><div class="insight-label">Alternative Data</div><h3>Text and sentiment require timestamp discipline</h3><p>News, filings, social media, and LLM-derived signals are useful only when information timing is auditable.</p><p class="insight-implication">Implication: look-ahead bias is the first failure mode to rule out.</p></article>
        <article class="insight-box"><div class="insight-label">Execution</div><h3>Microstructure shapes realized returns</h3><p>Order flow, liquidity, market impact, and fees decide whether a signal survives implementation.</p><p class="insight-implication">Implication: trading research should model execution, not just prices.</p></article>
      </div>
    </section>
"""


def research_copy():
    return {"en": research_overview_html()}


def overall_research_templates():
    return {
        "en": {
            "timelineTitle": "Research Timeline",
            "summary": [
                "For {range}, the trade corpus contains {papers} selected papers across {activeYears} active years, with {citations} citations. The strongest taxonomy signals are {topCategories}, and the most active year is {peakYear} ({peakYearCount} papers).",
                "The leading citation-ranked paper is \"{topPaper}\" ({topPaperYear}, {topPaperCitations} citations) in {topPaperCategory}. Keywords such as {topKeywords} show how this period connects asset pricing, forecasting, portfolio construction, risk, market microstructure, sentiment, and AI-based trading.",
            ],
            "insightsTitle": "Research Insights",
            "insights": [
                {"label": "Signal Quality", "title": "Prediction must survive market frictions", "body": "{topCategories} dominate {range}, showing where citation-ranked influence concentrates.", "implication": "Implication: evaluate signals after costs, turnover, capacity, and risk constraints."},
                {"label": "Citation Mass", "title": "Citation peaks mark reusable finance infrastructure", "body": "The selected range carries {citations} citations, with citation mass peaking around {peakCitationYear}.", "implication": "Implication: high-impact papers often define data, factors, methods, or validation conventions."},
                {"label": "AI Trading", "title": "AI expands the signal surface", "body": "Frequent tags such as {topKeywords} indicate whether the period centers on forecasting, deep learning, sentiment, portfolio construction, or execution.", "implication": "Implication: richer models still need realistic backtests and market-aware validation."},
                {"label": "Open Gaps", "title": "Recent papers need recency correction", "body": "Newer AI trading papers can be underweighted before citation signals mature.", "implication": "Implication: use this map as a stable citation view, then layer recent expert review on top."},
            ],
        }
    }


def write_markdown_html(markdown_path, title=None, lang="en"):
    text = markdown_path.read_text(encoding="utf-8")
    title = title or markdown_path.stem
    body = [f"<h1>{html.escape(title)}</h1>"]
    in_list = False
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            if in_list:
                body.append("</ul>")
                in_list = False
            continue
        if stripped.startswith("# "):
            if in_list:
                body.append("</ul>")
                in_list = False
            body.append(f"<h1>{html.escape(stripped[2:])}</h1>")
        elif stripped.startswith("## "):
            if in_list:
                body.append("</ul>")
                in_list = False
            body.append(f"<h2>{html.escape(stripped[3:])}</h2>")
        elif stripped.startswith("### "):
            if in_list:
                body.append("</ul>")
                in_list = False
            body.append(f"<h3>{html.escape(stripped[4:])}</h3>")
        elif stripped.startswith("- "):
            if not in_list:
                body.append("<ul>")
                in_list = True
            body.append(f"<li>{html.escape(stripped[2:])}</li>")
        elif stripped.startswith("<"):
            if in_list:
                body.append("</ul>")
                in_list = False
            body.append(line)
        else:
            if in_list:
                body.append("</ul>")
                in_list = False
            body.append(f"<p>{html.escape(stripped)}</p>")
    if in_list:
        body.append("</ul>")
    html_doc = f"""<!doctype html>
<html lang="{html.escape(lang, quote=True)}">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{html.escape(title)}</title>
  <style>body{{font-family:Inter,Segoe UI,Arial,sans-serif;max-width:980px;margin:40px auto;padding:0 22px;line-height:1.65;color:#172033}}a{{color:#0f5f97}}code{{background:#f1f5f9;padding:2px 5px;border-radius:4px}}table{{width:100%;border-collapse:collapse}}td,th{{border:1px solid #d9dee8;padding:6px;vertical-align:top}}</style>
</head>
<body>{''.join(body)}</body>
</html>
"""
    markdown_path.with_suffix(".html").write_text(html_doc, encoding="utf-8")


README_LANGUAGE_FILES = {
    "en": "README.md",
    "de": "README.de.md",
    "es": "README.es.md",
    "fr": "README.fr.md",
    "ja": "README.ja.md",
    "ko": "README.ko.md",
    "pt": "README.pt.md",
    "ru": "README.ru.md",
    "zh": "README.zh.md",
}

README_LANGUAGE_LABELS = {
    "en": "English",
    "de": "Deutsch",
    "es": "Español",
    "fr": "français",
    "ja": "日本語",
    "ko": "한국어",
    "pt": "Português",
    "ru": "Русский",
    "zh": "中文",
}

README_LOCALIZED_COPY = {
    "de": {
        "description": "Eine taxonomieorientierte, nach Zitierungen gerankte Karte der Forschung zu Aktienanlagen und KI-gestütztem Trading von 2000 bis 2026.",
        "note": "Hinweis: Diese lokalisierte README wird aus dem aktuellen Datensatz der ausgewählten Arbeiten neu erzeugt. Paper-Titel, Autorennamen, URLs, Badges, Dateipfade und Identifikatoren bleiben aus Gründen der Reproduzierbarkeit unverändert.",
        "generated": "Erzeugt am {generated_date} aus frei zugänglichen Semantic-Scholar-Metadaten. Diese Ausgabe untersucht für {year_range} je Publikationsjahr bis zu {candidates_per_year} Kandidatenarbeiten zu Aktienanlagen und KI-Trading, hält einen geprüften Kandidatenpool von {candidate_count} Datensätzen, wählt pro Jahr die {target_per_year} meistzitierten Arbeiten aus ({selected_count} ausgewählte Arbeiten) und ordnet sie nach Forschungstaxonomie neu.",
        "curation": "Die Sammlung nutzt die Bulk-Suche des Semantic Scholar Academic Graph. Die Abfragen decken Aktienprognosen, Aktienrenditen, Asset Pricing, Portfoliooptimierung, finanzielle Zeitreihen, algorithmisches Trading, Reinforcement Learning, Sentiment/Nachrichten/alternative Daten, Markt-Mikrostruktur, Hochfrequenzhandel, Volatilität und Risikothemen ab. Für jedes Jahr von {start_year} bis {end_year} werden Ergebnisse auf das Publikationsjahr gefiltert, mit expliziten Relevanzausdrücken zu Aktienanlagen und KI-Trading in Titel-, Abstract- und Venue-Metadaten geprüft, nach DOI, arXiv, PubMed, CorpusId, paperId und anschließend Titel dedupliziert und auf höchstens {candidates_per_year} Kandidaten nach Zitierzahl reduziert. Die finale Awesome-Liste wählt innerhalb jedes Publikationsjahres die {target_per_year} meistzitierten Arbeiten aus; einflussreiche Zitierungen und ein deterministischer Metadaten-Score bleiben als Tie-Breaker und Audit-Signale erhalten.",
        "no_paid": "Taxonomie, Kernideen, Stärken, Einschränkungen, Methodentags und Keyword-Tags werden deterministisch aus öffentlichen Metadaten und regelbasierten Domänenkonventionen erzeugt. Es wurden keine bezahlten APIs, bezahlten LLMs, bezahlten Übersetzungen oder bezahlten Compute-Dienste verwendet.",
        "show": "Repräsentative Arbeiten anzeigen: ",
        "paper_word": "Arbeiten",
    },
    "es": {
        "description": "Un mapa de investigación sobre inversión en acciones y trading impulsado por IA, organizado por taxonomía y ordenado por citas, de 2000 a 2026.",
        "note": "Nota: este README localizado se regenera desde el conjunto actual de artículos seleccionados. Los títulos de los artículos, autores, URL, insignias, rutas de archivo e identificadores se conservan para mantener la reproducibilidad.",
        "generated": "Generado el {generated_date} a partir de metadatos públicos y gratuitos de Semantic Scholar. Esta edición investiga hasta {candidates_per_year} artículos candidatos sobre inversión en acciones y trading con IA por año de publicación para {year_range}, mantiene un conjunto auditado de {candidate_count} registros, selecciona los {target_per_year} artículos más citados de cada año ({selected_count} artículos seleccionados) y los reorganiza por taxonomía de investigación.",
        "curation": "La colección utiliza la búsqueda masiva de Semantic Scholar Academic Graph. Las consultas cubren predicción bursátil, rentabilidades de acciones, valoración de activos, optimización de carteras, series temporales financieras, trading algorítmico, aprendizaje por refuerzo, sentimiento/noticias/datos alternativos, microestructura de mercado, trading de alta frecuencia, volatilidad y riesgo. Para cada año desde {start_year} hasta {end_year}, los resultados se filtran por año de publicación, se revisan con expresiones explícitas de relevancia para inversión en acciones y trading con IA en metadatos de título, resumen y venue, se deduplican por DOI, arXiv, PubMed, CorpusId, paperId y luego título, y se reducen a un máximo de {candidates_per_year} candidatos por número de citas. La lista awesome final selecciona los {target_per_year} artículos más citados dentro de cada año; las citas influyentes y una puntuación determinística de importancia de metadatos se conservan como desempates y señales de auditoría.",
        "no_paid": "La taxonomía, ideas clave, fortalezas, limitaciones, etiquetas de método y etiquetas de palabra clave se generan de forma determinística a partir de metadatos públicos y convenciones de dominio basadas en reglas. No se usaron APIs de pago, LLMs de pago, traducción de pago ni cómputo de pago.",
        "show": "Mostrar artículos representativos de ",
        "paper_word": "artículos",
    },
    "fr": {
        "description": "Une carte de recherche sur l'investissement actions et le trading piloté par l'IA, structurée par taxonomie et classée par citations, de 2000 à 2026.",
        "note": "Note : ce README localisé est régénéré à partir du jeu actuel d'articles sélectionnés. Les titres d'articles, noms d'auteurs, URL, badges, chemins de fichiers et identifiants sont conservés pour la reproductibilité.",
        "generated": "Généré le {generated_date} à partir de métadonnées publiques gratuites de Semantic Scholar. Cette édition examine jusqu'à {candidates_per_year} articles candidats sur l'investissement actions et le trading par IA par année de publication pour {year_range}, conserve un vivier audité de {candidate_count} enregistrements, sélectionne les {target_per_year} articles les plus cités de chaque année ({selected_count} articles sélectionnés) et les réorganise par taxonomie de recherche.",
        "curation": "La collection utilise la recherche en masse du Semantic Scholar Academic Graph. Les requêtes couvrent la prédiction boursière, les rendements actions, l'asset pricing, l'optimisation de portefeuille, les séries temporelles financières, le trading algorithmique, l'apprentissage par renforcement, le sentiment/les nouvelles/les données alternatives, la microstructure de marché, le trading haute fréquence, la volatilité et les thèmes de risque. Pour chaque année de {start_year} à {end_year}, les résultats sont filtrés par année de publication, examinés avec des expressions explicites de pertinence pour l'investissement actions et le trading IA dans les métadonnées de titre, résumé et venue, dédupliqués par DOI, arXiv, PubMed, CorpusId, paperId puis titre, et réduits à au plus {candidates_per_year} candidats par nombre de citations. La liste awesome finale sélectionne les {target_per_year} articles les plus cités dans chaque année; les citations influentes et un score déterministe d'importance des métadonnées sont conservés comme signaux de départage et d'audit.",
        "no_paid": "La taxonomie, les idées clés, les forces, les limites, les tags de méthode et les tags de mots-clés sont générés de manière déterministe à partir de métadonnées publiques et de conventions de domaine fondées sur des règles. Aucune API payante, aucun LLM payant, aucune traduction payante ni aucun calcul payant n'a été utilisé.",
        "show": "Afficher les articles représentatifs pour ",
        "paper_word": "articles",
    },
    "ja": {
        "description": "2000年から2026年までの株式投資およびAI主導トレーディング研究を、分類体系を軸に引用数順で整理したマップです。",
        "note": "注: このローカライズ版 README は、現在の選定論文データセットから再生成されています。論文タイトル、著者名、URL、バッジ、ファイルパス、識別子は再現性のため原文のまま保持しています。",
        "generated": "無料で公開されている Semantic Scholar メタデータに基づき {generated_date} に生成しました。この版では、{year_range} の各出版年について株式投資およびAIトレーディング関連の候補論文を最大 {candidates_per_year} 件調査し、監査済み候補プール {candidate_count} 件を保持し、各年の引用数上位 {target_per_year} 件（選定論文 {selected_count} 件）を選び、研究分類体系に沿って再編成しています。",
        "curation": "このコレクションは Semantic Scholar Academic Graph の bulk search を使用しています。検索クエリは、株価予測、株式リターン、資産価格評価、ポートフォリオ最適化、金融時系列、アルゴリズム取引、強化学習、センチメント/ニュース/代替データ、市場ミクロ構造、高頻度取引、ボラティリティ、リスクを対象にしています。{start_year} 年から {end_year} 年までの各年について、結果を出版年で絞り込み、タイトル・要旨・venue メタデータに含まれる株式投資およびAIトレーディング関連表現でスクリーニングし、DOI、arXiv、PubMed、CorpusId、paperId、最後にタイトルで重複排除し、引用数順で最大 {candidates_per_year} 件に絞ります。最終的な awesome リストでは、各出版年内で引用数上位 {target_per_year} 件を選定し、影響力引用数と決定的なメタデータ重要度スコアをタイブレークおよび監査シグナルとして保持します。",
        "no_paid": "分類体系、主要アイデア、強み、限界、手法タグ、キーワードタグは、公開メタデータとルールベースのドメイン規約から決定的に生成しています。有料API、有料LLM、有料翻訳、有料計算資源は使用していません。",
        "show": "代表論文を表示: ",
        "paper_word": "本",
    },
    "ko": {
        "description": "2000년부터 2026년까지의 주식 투자 및 AI 기반 트레이딩 연구를 분류 우선 방식으로 정리하고 인용수 기준으로 순위화한 지도입니다.",
        "note": "참고: 이 현지화 README는 현재 선정 논문 데이터셋에서 재생성되었습니다. 논문 제목, 저자명, URL, 배지, 파일 경로, 식별자는 재현성을 위해 원문 그대로 유지했습니다.",
        "generated": "무료 공개 Semantic Scholar 메타데이터를 기준으로 {generated_date}에 생성했습니다. 이 판은 {year_range} 기간의 각 출판연도마다 주식 투자 및 AI 트레이딩 후보 논문을 최대 {candidates_per_year}편 조사하고, 감사된 후보 풀 {candidate_count}건을 유지하며, 매년 인용수 상위 {target_per_year}편({selected_count}편 선정)을 골라 연구 분류 체계별로 재구성합니다.",
        "curation": "이 컬렉션은 Semantic Scholar Academic Graph bulk search를 사용합니다. 쿼리는 주가 예측, 주식 수익률, 자산가격결정, 포트폴리오 최적화, 금융 시계열, 알고리즘 트레이딩, 강화학습, 감성/뉴스/대체데이터, 시장 미시구조, 고빈도 거래, 변동성, 리스크 주제를 포괄합니다. {start_year}년부터 {end_year}년까지 각 연도 결과를 출판연도로 필터링하고, 제목/초록/venue 메타데이터에서 주식 투자 및 AI 트레이딩 관련 명시적 표현으로 선별한 뒤, DOI, arXiv, PubMed, CorpusId, paperId, 제목 순으로 중복을 제거하고 인용수 기준 최대 {candidates_per_year}편의 후보로 축소합니다. 최종 awesome 목록은 각 출판연도 안에서 인용수 상위 {target_per_year}편을 선정하며, 영향력 인용수와 결정적 메타데이터 중요도 점수는 동률 처리 및 감사 신호로 보존합니다.",
        "no_paid": "분류 체계, 핵심 아이디어, 강점, 한계, 방법 태그, 키워드 태그는 공개 메타데이터와 규칙 기반 도메인 규약에서 결정적으로 생성됩니다. 유료 API, 유료 LLM, 유료 번역, 유료 컴퓨팅은 사용하지 않았습니다.",
        "show": "대표 논문 보기: ",
        "paper_word": "편",
    },
    "pt": {
        "description": "Um mapa de pesquisa sobre investimento em ações e trading orientado por IA, organizado por taxonomia e classificado por citações, de 2000 a 2026.",
        "note": "Nota: este README localizado é regenerado a partir do conjunto atual de artigos selecionados. Títulos dos artigos, autores, URLs, badges, caminhos de arquivo e identificadores são preservados para reprodutibilidade.",
        "generated": "Gerado em {generated_date} a partir de metadados públicos gratuitos do Semantic Scholar. Esta edição investiga até {candidates_per_year} artigos candidatos de investimento em ações e trading com IA por ano de publicação para {year_range}, mantém um conjunto auditado de {candidate_count} registros, seleciona os {target_per_year} artigos mais citados de cada ano ({selected_count} artigos selecionados) e os reorganiza por taxonomia de pesquisa.",
        "curation": "A coleção usa a busca em massa do Semantic Scholar Academic Graph. As consultas cobrem previsão de ações, retornos de ações, precificação de ativos, otimização de portfólio, séries temporais financeiras, trading algorítmico, aprendizado por reforço, sentimento/notícias/dados alternativos, microestrutura de mercado, trading de alta frequência, volatilidade e risco. Para cada ano de {start_year} a {end_year}, os resultados são filtrados pelo ano de publicação, triados com expressões explícitas de relevância para investimento em ações e trading com IA nos metadados de título, resumo e venue, deduplicados por DOI, arXiv, PubMed, CorpusId, paperId e depois título, e reduzidos a no máximo {candidates_per_year} candidatos por número de citações. A lista awesome final seleciona os {target_per_year} artigos mais citados dentro de cada ano; citações influentes e uma pontuação determinística de importância de metadados são mantidas como desempates e sinais de auditoria.",
        "no_paid": "A taxonomia, ideias-chave, pontos fortes, limitações, tags de método e tags de palavra-chave são gerados deterministicamente a partir de metadados públicos e convenções de domínio baseadas em regras. Nenhuma API paga, LLM pago, tradução paga ou computação paga foi usada.",
        "show": "Mostrar artigos representativos de ",
        "paper_word": "artigos",
    },
    "ru": {
        "description": "Таксономическая и ранжированная по цитируемости карта исследований по инвестициям в акции и AI-трейдингу за 2000-2026 годы.",
        "note": "Примечание: этот локализованный README заново создается из текущего набора выбранных работ. Названия статей, авторы, URL, бейджи, пути к файлам и идентификаторы сохранены для воспроизводимости.",
        "generated": "Сгенерировано {generated_date} на основе бесплатных публичных метаданных Semantic Scholar. Эта версия проверяет до {candidates_per_year} кандидатных работ по инвестициям в акции и AI-трейдингу на каждый год публикации за {year_range}, хранит аудированный пул из {candidate_count} записей, выбирает {target_per_year} наиболее цитируемых работ за каждый год ({selected_count} выбранных работ) и реорганизует их по исследовательской таксономии.",
        "curation": "Коллекция использует bulk search Semantic Scholar Academic Graph. Запросы охватывают прогнозирование акций, доходности акций, asset pricing, оптимизацию портфеля, финансовые временные ряды, алгоритмический трейдинг, обучение с подкреплением, sentiment/новости/альтернативные данные, микроструктуру рынка, высокочастотную торговлю, волатильность и риск. Для каждого года с {start_year} по {end_year} результаты фильтруются по году публикации, проверяются явными выражениями релевантности к инвестициям в акции и AI-трейдингу в метаданных title/abstract/venue, дедуплицируются по DOI, arXiv, PubMed, CorpusId, paperId, затем названию, и сокращаются максимум до {candidates_per_year} кандидатов по числу цитирований. Финальный awesome-список выбирает {target_per_year} наиболее цитируемых работ в каждом году; влиятельные цитирования и детерминированный metadata-importance score сохраняются как сигналы для тай-брейков и аудита.",
        "no_paid": "Таксономия, ключевые идеи, сильные стороны, ограничения, method tags и keyword tags генерируются детерминированно из публичных метаданных и правил предметной области. Платные API, платные LLM, платный перевод и платные вычисления не использовались.",
        "show": "Показать репрезентативные работы: ",
        "paper_word": "работ",
    },
    "zh": {
        "description": "这是一个按研究分类组织、按引用数排序的股票投资与 AI 驱动交易研究地图，覆盖 2000 年至 2026 年。",
        "note": "注：本地化 README 已从当前入选论文数据集重新生成。论文题名、作者、URL、徽章、文件路径和标识符为保持可复现性而保留原文。",
        "generated": "基于免费的公开 Semantic Scholar 元数据，于 {generated_date} 生成。本版针对 {year_range} 的每个发表年份，调查最多 {candidates_per_year} 篇股票投资与 AI 交易候选论文，保留 {candidate_count} 条已审计候选记录，按引用数选出每年排名前 {target_per_year} 的论文（共 {selected_count} 篇入选），并按研究分类重新组织。",
        "curation": "本集合使用 Semantic Scholar Academic Graph 批量搜索。查询覆盖股票预测、股票收益、资产定价、投资组合优化、金融时间序列、算法交易、强化学习、情绪/新闻/另类数据、市场微观结构、高频交易、波动率和风险主题。对于 {start_year} 至 {end_year} 的每一年，结果按发表年份过滤，并使用标题、摘要和 venue 元数据中的股票投资与 AI 交易相关表达进行筛选，再按 DOI、arXiv、PubMed、CorpusId、paperId 和标题去重，最后按引用数保留最多 {candidates_per_year} 篇候选论文。最终 awesome 列表在每个发表年份内选出引用数最高的 {target_per_year} 篇论文；影响力引用数和确定性的元数据重要性分数会保留为平局处理和审计信号。",
        "no_paid": "分类体系、关键思想、优势、局限、方法标签和关键词标签均由公开元数据和基于规则的领域约定确定性生成。未使用付费 API、付费 LLM、付费翻译或付费计算。",
        "show": "显示代表性论文：",
        "paper_word": "篇",
    },
}

README_STATIC_REPLACEMENTS = {
    "de": [
        ("## Project Links", "## Projektlinks"),
        ("- Open Interactive Website:", "- Interaktive Website:"),
        ("- Selected dataset:", "- Ausgewählter Datensatz:"),
        ("- Taxonomy dataset with paper-level ideas, strengths, and limitations:", "- Taxonomiedatensatz mit Ideen, Stärken und Einschränkungen auf Paper-Ebene:"),
        ("- Precomputed period analysis:", "- Vorberechnete Periodenanalyse:"),
        ("- Candidate Pool:", "- Kandidatenpool:"),
        ("- English review draft:", "- Englischer Review-Entwurf:"),
        ("- Korean review draft:", "- Koreanischer Review-Entwurf:"),
        ("## Keywords Convention", "## Keyword-Konvention"),
        ("These badges define the stock-investment and AI-trading keyword tags used to read and extend this collection.", "Diese Badges definieren die Keyword-Tags für Aktienanlage und KI-Trading, mit denen diese Sammlung gelesen und erweitert wird."),
        ("## Taxonomy Overview", "## Taxonomieüberblick"),
        ("**Total selected papers**", "**Insgesamt ausgewählte Arbeiten**"),
        ("**Candidate pool audited**", "**Geprüfter Kandidatenpool**"),
        ("**Citation count in selected set**", "**Zitierungen im ausgewählten Set**"),
        ("## Taxonomy Collections", "## Taxonomiesammlungen"),
        ("- Papers selected:", "- Ausgewählte Arbeiten:"),
        ("- Years covered:", "- Abgedeckte Jahre:"),
        ("- Citation count in selected set:", "- Zitierungen im ausgewählten Set:"),
        ("- Category Overview:", "- Kategorieüberblick:"),
        ("- Limitations:", "- Einschränkungen:"),
        ("## Yearly Selection Summary", "## Jährliche Auswahlübersicht"),
        ("| Year | Candidate records | Selected papers | Selected citations | Top selected paper |", "| Jahr | Kandidatenrecords | Ausgewählte Arbeiten | Zitierungen der Auswahl | Top-Paper der Auswahl |"),
        ("## Curation Method", "## Kurationsmethode"),
        ("## Limitations", "## Einschränkungen"),
        ("- This is a metadata-driven citation map, not a full systematic review of every PDF.", "- Dies ist eine metadatengetriebene Zitierkarte, keine vollständige systematische Durchsicht jedes PDFs."),
        ("- Citation count is an influence signal, not investment advice or proof of live trading profitability.", "- Zitierzahl ist ein Einflussindikator, keine Anlageberatung und kein Nachweis realer Trading-Profitabilität."),
        ("- Semantic Scholar metadata can omit venues, abstracts, PDFs, or influential citation counts for some records.", "- Semantic-Scholar-Metadaten können bei einzelnen Records Venues, Abstracts, PDFs oder einflussreiche Zitierungen auslassen."),
        ("## Acknowledgements", "## Danksagung"),
        ("## License", "## Lizenz"),
    ],
    "es": [
        ("## Project Links", "## Enlaces del proyecto"),
        ("- Open Interactive Website:", "- Sitio web interactivo:"),
        ("- Selected dataset:", "- Conjunto seleccionado:"),
        ("- Taxonomy dataset with paper-level ideas, strengths, and limitations:", "- Conjunto de taxonomía con ideas, fortalezas y limitaciones por artículo:"),
        ("- Precomputed period analysis:", "- Análisis de periodos precalculado:"),
        ("- Candidate Pool:", "- Conjunto de candidatos:"),
        ("- English review draft:", "- Borrador de revisión en inglés:"),
        ("- Korean review draft:", "- Borrador de revisión en coreano:"),
        ("## Keywords Convention", "## Convención de palabras clave"),
        ("These badges define the stock-investment and AI-trading keyword tags used to read and extend this collection.", "Estas insignias definen las etiquetas de palabras clave de inversión en acciones y trading con IA usadas para leer y ampliar esta colección."),
        ("## Taxonomy Overview", "## Resumen de taxonomía"),
        ("**Total selected papers**", "**Total de artículos seleccionados**"),
        ("**Candidate pool audited**", "**Conjunto de candidatos auditado**"),
        ("**Citation count in selected set**", "**Citas en el conjunto seleccionado**"),
        ("## Taxonomy Collections", "## Colecciones por taxonomía"),
        ("- Papers selected:", "- Artículos seleccionados:"),
        ("- Years covered:", "- Años cubiertos:"),
        ("- Citation count in selected set:", "- Citas en el conjunto seleccionado:"),
        ("- Category Overview:", "- Resumen de categoría:"),
        ("- Limitations:", "- Limitaciones:"),
        ("## Yearly Selection Summary", "## Resumen anual de selección"),
        ("| Year | Candidate records | Selected papers | Selected citations | Top selected paper |", "| Año | Registros candidatos | Artículos seleccionados | Citas seleccionadas | Artículo principal |"),
        ("## Curation Method", "## Método de curación"),
        ("## Limitations", "## Limitaciones"),
        ("- This is a metadata-driven citation map, not a full systematic review of every PDF.", "- Este es un mapa de citas basado en metadatos, no una revisión sistemática completa de cada PDF."),
        ("- Citation count is an influence signal, not investment advice or proof of live trading profitability.", "- El número de citas es una señal de influencia, no asesoría de inversión ni prueba de rentabilidad real."),
        ("- Semantic Scholar metadata can omit venues, abstracts, PDFs, or influential citation counts for some records.", "- Los metadatos de Semantic Scholar pueden omitir venues, resúmenes, PDFs o citas influyentes en algunos registros."),
        ("## Acknowledgements", "## Agradecimientos"),
        ("## License", "## Licencia"),
    ],
    "fr": [
        ("## Project Links", "## Liens du projet"),
        ("- Open Interactive Website:", "- Site web interactif :"),
        ("- Selected dataset:", "- Jeu de données sélectionné :"),
        ("- Taxonomy dataset with paper-level ideas, strengths, and limitations:", "- Jeu de taxonomie avec idées, forces et limites au niveau des articles :"),
        ("- Precomputed period analysis:", "- Analyse de période précalculée :"),
        ("- Candidate Pool:", "- Vivier de candidats :"),
        ("- English review draft:", "- Brouillon de revue en anglais :"),
        ("- Korean review draft:", "- Brouillon de revue en coréen :"),
        ("## Keywords Convention", "## Convention des mots-clés"),
        ("These badges define the stock-investment and AI-trading keyword tags used to read and extend this collection.", "Ces badges définissent les tags de mots-clés liés à l'investissement actions et au trading IA utilisés pour lire et étendre cette collection."),
        ("## Taxonomy Overview", "## Aperçu de la taxonomie"),
        ("**Total selected papers**", "**Total d'articles sélectionnés**"),
        ("**Candidate pool audited**", "**Vivier de candidats audité**"),
        ("**Citation count in selected set**", "**Citations dans l'ensemble sélectionné**"),
        ("## Taxonomy Collections", "## Collections taxonomiques"),
        ("- Papers selected:", "- Articles sélectionnés :"),
        ("- Years covered:", "- Années couvertes :"),
        ("- Citation count in selected set:", "- Citations dans l'ensemble sélectionné :"),
        ("- Category Overview:", "- Aperçu de la catégorie :"),
        ("- Limitations:", "- Limites :"),
        ("## Yearly Selection Summary", "## Résumé annuel de sélection"),
        ("| Year | Candidate records | Selected papers | Selected citations | Top selected paper |", "| Année | Candidats | Articles sélectionnés | Citations sélectionnées | Article principal |"),
        ("## Curation Method", "## Méthode de curation"),
        ("## Limitations", "## Limites"),
        ("- This is a metadata-driven citation map, not a full systematic review of every PDF.", "- Il s'agit d'une carte de citations pilotée par les métadonnées, et non d'une revue systématique complète de chaque PDF."),
        ("- Citation count is an influence signal, not investment advice or proof of live trading profitability.", "- Le nombre de citations est un signal d'influence, pas un conseil d'investissement ni une preuve de rentabilité en trading réel."),
        ("- Semantic Scholar metadata can omit venues, abstracts, PDFs, or influential citation counts for some records.", "- Les métadonnées Semantic Scholar peuvent omettre venues, résumés, PDFs ou citations influentes pour certains enregistrements."),
        ("## Acknowledgements", "## Remerciements"),
        ("## License", "## Licence"),
    ],
    "ja": [
        ("## Project Links", "## プロジェクトリンク"),
        ("- Open Interactive Website:", "- インタラクティブサイト:"),
        ("- Selected dataset:", "- 選定データセット:"),
        ("- Taxonomy dataset with paper-level ideas, strengths, and limitations:", "- 論文単位のアイデア・強み・限界を含む分類データセット:"),
        ("- Precomputed period analysis:", "- 事前計算済み期間分析:"),
        ("- Candidate Pool:", "- 候補プール:"),
        ("- English review draft:", "- 英語レビュー草稿:"),
        ("- Korean review draft:", "- 韓国語レビュー草稿:"),
        ("## Keywords Convention", "## キーワード規約"),
        ("These badges define the stock-investment and AI-trading keyword tags used to read and extend this collection.", "これらのバッジは、このコレクションを読む・拡張するための株式投資およびAIトレーディング関連キーワードタグを定義します。"),
        ("## Taxonomy Overview", "## 分類体系の概要"),
        ("**Total selected papers**", "**選定論文総数**"),
        ("**Candidate pool audited**", "**監査済み候補プール**"),
        ("**Citation count in selected set**", "**選定セット内の引用数**"),
        ("## Taxonomy Collections", "## 分類別コレクション"),
        ("- Papers selected:", "- 選定論文:"),
        ("- Years covered:", "- 対象年:"),
        ("- Citation count in selected set:", "- 選定セット内の引用数:"),
        ("- Category Overview:", "- カテゴリ概要:"),
        ("- Limitations:", "- 限界:"),
        ("## Yearly Selection Summary", "## 年別選定サマリー"),
        ("| Year | Candidate records | Selected papers | Selected citations | Top selected paper |", "| 年 | 候補レコード | 選定論文 | 選定論文の引用数 | 最上位論文 |"),
        ("## Curation Method", "## キュレーション方法"),
        ("## Limitations", "## 限界"),
        ("- This is a metadata-driven citation map, not a full systematic review of every PDF.", "- これはメタデータ駆動の引用マップであり、すべてのPDFを精査した完全な系統的レビューではありません。"),
        ("- Citation count is an influence signal, not investment advice or proof of live trading profitability.", "- 引用数は影響度のシグナルであり、投資助言や実運用収益性の証明ではありません。"),
        ("- Semantic Scholar metadata can omit venues, abstracts, PDFs, or influential citation counts for some records.", "- Semantic Scholar メタデータでは、一部のレコードで venue、要旨、PDF、影響力引用数が欠落する場合があります。"),
        ("## Acknowledgements", "## 謝辞"),
        ("## License", "## ライセンス"),
    ],
    "ko": [
        ("## Project Links", "## 프로젝트 링크"),
        ("- Open Interactive Website:", "- 인터랙티브 웹사이트:"),
        ("- Selected dataset:", "- 선정 데이터셋:"),
        ("- Taxonomy dataset with paper-level ideas, strengths, and limitations:", "- 논문별 아이디어, 강점, 한계가 포함된 분류 데이터셋:"),
        ("- Precomputed period analysis:", "- 사전 계산된 기간 분석:"),
        ("- Candidate Pool:", "- 후보 풀:"),
        ("- English review draft:", "- 영어 리뷰 초안:"),
        ("- Korean review draft:", "- 한국어 리뷰 초안:"),
        ("## Keywords Convention", "## 키워드 규칙"),
        ("These badges define the stock-investment and AI-trading keyword tags used to read and extend this collection.", "이 배지는 컬렉션을 읽고 확장할 때 사용하는 주식 투자 및 AI 트레이딩 키워드 태그를 정의합니다."),
        ("## Taxonomy Overview", "## 분류 체계 개요"),
        ("**Total selected papers**", "**전체 선정 논문**"),
        ("**Candidate pool audited**", "**감사된 후보 풀**"),
        ("**Citation count in selected set**", "**선정 세트의 총 인용수**"),
        ("## Taxonomy Collections", "## 분류별 컬렉션"),
        ("- Papers selected:", "- 선정 논문:"),
        ("- Years covered:", "- 포함 연도:"),
        ("- Citation count in selected set:", "- 선정 세트의 인용수:"),
        ("- Category Overview:", "- 분류 개요:"),
        ("- Limitations:", "- 한계:"),
        ("## Yearly Selection Summary", "## 연도별 선정 요약"),
        ("| Year | Candidate records | Selected papers | Selected citations | Top selected paper |", "| 연도 | 후보 기록 | 선정 논문 | 선정 논문 인용수 | 최상위 선정 논문 |"),
        ("## Curation Method", "## 큐레이션 방법"),
        ("## Limitations", "## 한계"),
        ("- This is a metadata-driven citation map, not a full systematic review of every PDF.", "- 이 자료는 메타데이터 기반 인용 지도이며, 모든 PDF를 읽은 완전한 체계적 문헌고찰은 아닙니다."),
        ("- Citation count is an influence signal, not investment advice or proof of live trading profitability.", "- 인용수는 영향력 신호일 뿐이며 투자 조언이나 실제 거래 수익성의 증거가 아닙니다."),
        ("- Semantic Scholar metadata can omit venues, abstracts, PDFs, or influential citation counts for some records.", "- Semantic Scholar 메타데이터에는 일부 레코드의 venue, 초록, PDF, 영향력 인용수가 누락될 수 있습니다."),
        ("## Acknowledgements", "## 감사의 말"),
        ("## License", "## 라이선스"),
    ],
    "pt": [
        ("## Project Links", "## Links do projeto"),
        ("- Open Interactive Website:", "- Site interativo:"),
        ("- Selected dataset:", "- Conjunto selecionado:"),
        ("- Taxonomy dataset with paper-level ideas, strengths, and limitations:", "- Conjunto de taxonomia com ideias, pontos fortes e limitações por artigo:"),
        ("- Precomputed period analysis:", "- Análise de períodos pré-computada:"),
        ("- Candidate Pool:", "- Conjunto de candidatos:"),
        ("- English review draft:", "- Rascunho de revisão em inglês:"),
        ("- Korean review draft:", "- Rascunho de revisão em coreano:"),
        ("## Keywords Convention", "## Convenção de palavras-chave"),
        ("These badges define the stock-investment and AI-trading keyword tags used to read and extend this collection.", "Estes badges definem as tags de palavras-chave de investimento em ações e trading com IA usadas para ler e ampliar esta coleção."),
        ("## Taxonomy Overview", "## Visão geral da taxonomia"),
        ("**Total selected papers**", "**Total de artigos selecionados**"),
        ("**Candidate pool audited**", "**Conjunto de candidatos auditado**"),
        ("**Citation count in selected set**", "**Citações no conjunto selecionado**"),
        ("## Taxonomy Collections", "## Coleções por taxonomia"),
        ("- Papers selected:", "- Artigos selecionados:"),
        ("- Years covered:", "- Anos cobertos:"),
        ("- Citation count in selected set:", "- Citações no conjunto selecionado:"),
        ("- Category Overview:", "- Visão geral da categoria:"),
        ("- Limitations:", "- Limitações:"),
        ("## Yearly Selection Summary", "## Resumo anual de seleção"),
        ("| Year | Candidate records | Selected papers | Selected citations | Top selected paper |", "| Ano | Registros candidatos | Artigos selecionados | Citações selecionadas | Artigo principal |"),
        ("## Curation Method", "## Método de curadoria"),
        ("## Limitations", "## Limitações"),
        ("- This is a metadata-driven citation map, not a full systematic review of every PDF.", "- Este é um mapa de citações orientado por metadados, não uma revisão sistemática completa de cada PDF."),
        ("- Citation count is an influence signal, not investment advice or proof of live trading profitability.", "- O número de citações é um sinal de influência, não aconselhamento de investimento nem prova de lucratividade real."),
        ("- Semantic Scholar metadata can omit venues, abstracts, PDFs, or influential citation counts for some records.", "- Os metadados do Semantic Scholar podem omitir venues, resumos, PDFs ou citações influentes em alguns registros."),
        ("## Acknowledgements", "## Agradecimentos"),
        ("## License", "## Licença"),
    ],
    "ru": [
        ("## Project Links", "## Ссылки проекта"),
        ("- Open Interactive Website:", "- Интерактивный сайт:"),
        ("- Selected dataset:", "- Выбранный набор данных:"),
        ("- Taxonomy dataset with paper-level ideas, strengths, and limitations:", "- Таксономический набор с идеями, сильными сторонами и ограничениями по каждой работе:"),
        ("- Precomputed period analysis:", "- Предварительно рассчитанный анализ периодов:"),
        ("- Candidate Pool:", "- Пул кандидатов:"),
        ("- English review draft:", "- Черновик обзора на английском:"),
        ("- Korean review draft:", "- Черновик обзора на корейском:"),
        ("## Keywords Convention", "## Соглашение по ключевым словам"),
        ("These badges define the stock-investment and AI-trading keyword tags used to read and extend this collection.", "Эти бейджи задают keyword tags для инвестиций в акции и AI-трейдинга, используемые для чтения и расширения коллекции."),
        ("## Taxonomy Overview", "## Обзор таксономии"),
        ("**Total selected papers**", "**Всего выбранных работ**"),
        ("**Candidate pool audited**", "**Аудированный пул кандидатов**"),
        ("**Citation count in selected set**", "**Цитирования в выбранном наборе**"),
        ("## Taxonomy Collections", "## Коллекции по таксономии"),
        ("- Papers selected:", "- Выбранные работы:"),
        ("- Years covered:", "- Охваченные годы:"),
        ("- Citation count in selected set:", "- Цитирования в выбранном наборе:"),
        ("- Category Overview:", "- Обзор категории:"),
        ("- Limitations:", "- Ограничения:"),
        ("## Yearly Selection Summary", "## Годовая сводка выбора"),
        ("| Year | Candidate records | Selected papers | Selected citations | Top selected paper |", "| Год | Записи-кандидаты | Выбранные работы | Цитирования выбранных | Лидирующая работа |"),
        ("## Curation Method", "## Метод курирования"),
        ("## Limitations", "## Ограничения"),
        ("- This is a metadata-driven citation map, not a full systematic review of every PDF.", "- Это карта цитирований на основе метаданных, а не полный систематический обзор каждого PDF."),
        ("- Citation count is an influence signal, not investment advice or proof of live trading profitability.", "- Число цитирований является сигналом влияния, а не инвестиционным советом или доказательством прибыльности реальной торговли."),
        ("- Semantic Scholar metadata can omit venues, abstracts, PDFs, or influential citation counts for some records.", "- Метаданные Semantic Scholar могут не содержать venue, abstract, PDF или influential citations для некоторых записей."),
        ("## Acknowledgements", "## Благодарности"),
        ("## License", "## Лицензия"),
    ],
    "zh": [
        ("## Project Links", "## 项目链接"),
        ("- Open Interactive Website:", "- 交互式网站："),
        ("- Selected dataset:", "- 入选数据集："),
        ("- Taxonomy dataset with paper-level ideas, strengths, and limitations:", "- 包含论文级思想、优势和局限的分类数据集："),
        ("- Precomputed period analysis:", "- 预计算期间分析："),
        ("- Candidate Pool:", "- 候选池："),
        ("- English review draft:", "- 英文综述草稿："),
        ("- Korean review draft:", "- 韩文综述草稿："),
        ("## Keywords Convention", "## 关键词约定"),
        ("These badges define the stock-investment and AI-trading keyword tags used to read and extend this collection.", "这些徽章定义了阅读和扩展本集合时使用的股票投资与 AI 交易关键词标签。"),
        ("## Taxonomy Overview", "## 分类概览"),
        ("**Total selected papers**", "**入选论文总数**"),
        ("**Candidate pool audited**", "**已审计候选池**"),
        ("**Citation count in selected set**", "**入选集合引用数**"),
        ("## Taxonomy Collections", "## 分类集合"),
        ("- Papers selected:", "- 入选论文："),
        ("- Years covered:", "- 覆盖年份："),
        ("- Citation count in selected set:", "- 入选集合引用数："),
        ("- Category Overview:", "- 分类概览："),
        ("- Limitations:", "- 局限："),
        ("## Yearly Selection Summary", "## 年度选择摘要"),
        ("| Year | Candidate records | Selected papers | Selected citations | Top selected paper |", "| 年份 | 候选记录 | 入选论文 | 入选论文引用数 | 最高引用入选论文 |"),
        ("## Curation Method", "## 策展方法"),
        ("## Limitations", "## 局限"),
        ("- This is a metadata-driven citation map, not a full systematic review of every PDF.", "- 这是基于元数据的引用地图，并非对每篇 PDF 的完整系统综述。"),
        ("- Citation count is an influence signal, not investment advice or proof of live trading profitability.", "- 引用数是影响力信号，不是投资建议，也不是实盘交易盈利能力证明。"),
        ("- Semantic Scholar metadata can omit venues, abstracts, PDFs, or influential citation counts for some records.", "- Semantic Scholar 元数据可能缺少部分记录的 venue、摘要、PDF 或影响力引用数。"),
        ("## Acknowledgements", "## 致谢"),
        ("## License", "## 许可证"),
    ],
}


def readme_language_selector_lines(owner, repo, active="en"):
    del owner, repo
    lines = ["<div align=\"center\">"]
    last_index = len(README_LANGUAGE_FILES) - 1
    for index, (code, filename) in enumerate(README_LANGUAGE_FILES.items()):
        label = README_LANGUAGE_LABELS[code]
        rendered_label = f"<strong>{label}</strong>" if code == active else label
        suffix = " |" if index != last_index else ""
        lines.append(f'  <a href="{filename}">{rendered_label}</a>{suffix}')
    lines.extend(["</div>", ""])
    return [
        *lines,
    ]


def readme_generation_context(selected, candidates):
    return {
        "generated_date": date.today().isoformat(),
        "candidates_per_year": f"{CANDIDATES_PER_YEAR:,}",
        "target_per_year": f"{TARGET_PER_YEAR:,}",
        "year_range": YEAR_RANGE_TEXT,
        "candidate_count": f"{len(candidates):,}",
        "selected_count": f"{len(selected):,}",
        "start_year": START_YEAR,
        "end_year": END_YEAR,
    }


def readme_english_curation_method():
    return f"The collection uses Semantic Scholar Academic Graph bulk search. Queries cover stock prediction, equity returns, asset pricing, portfolio optimization, financial time series, algorithmic trading, reinforcement learning, sentiment/news/alternative data, market microstructure, high-frequency trading, volatility, and risk themes. For each year from {START_YEAR} through {END_YEAR}, results are filtered to the publication year, screened with explicit stock-investment and AI-trading relevance expressions in title/abstract/venue metadata, deduplicated by DOI, arXiv, PubMed, CorpusId, paperId, then title, and reduced to at most {CANDIDATES_PER_YEAR:,} candidates by citation count. The final awesome list selects the top {TARGET_PER_YEAR:,} papers within each publication year by citation count; influential citation count and a deterministic metadata importance score are retained as tie-breakers and audit signals."


def replace_readme_selector(text, language):
    start = text.find('<div align="center">\n  <a href="README.md"')
    if start == -1:
        return text
    end = text.find("</div>", start)
    if end == -1:
        return text
    end += len("</div>")
    selector = "\n".join(readme_language_selector_lines(PROJECT_OWNER, PROJECT_REPO, active=language)).rstrip()
    note = README_LOCALIZED_COPY[language]["note"]
    return f"{text[:start]}{selector}\n\n{note}{text[end:]}"


def localize_readme_keywords(text, language):
    descriptions = globals().get("KEYWORD_DESCRIPTION_LANG", {}).get(language)
    if not descriptions:
        return text
    for keyword, english_description, _color in KEYWORD_CONVENTION:
        localized = descriptions.get(keyword)
        if localized:
            text = text.replace(f"**{keyword}**: {english_description}", f"**{keyword}**: {localized}")
    return text


def localized_readme_text(english_text, language, selected, candidates):
    copy = README_LOCALIZED_COPY[language]
    context = readme_generation_context(selected, candidates)
    text = replace_readme_selector(english_text, language)
    text = text.replace(PROJECT_DESCRIPTION, copy["description"], 1)
    text = re.sub(
        r"Generated on \d{4}-\d{2}-\d{2} from free public Semantic Scholar metadata\. This edition investigates up to [\s\S]*?research taxonomy\.",
        copy["generated"].format(**context),
        text,
        count=1,
    )
    text = text.replace(readme_english_curation_method(), copy["curation"].format(**context), 1)
    text = text.replace(
        "The taxonomy, key ideas, strengths, limitations, method tags, and keyword tags are generated deterministically from public metadata and rule-based domain conventions. No paid API, paid LLM, paid translation, or paid compute was used.",
        copy["no_paid"],
        1,
    )
    partial_year_warning = f"- Papers from {END_YEAR} are structurally citation-disadvantaged because the year is partial and citation accumulation is still immature as of {date.today().isoformat()}."
    text = text.replace(
        partial_year_warning,
        {
            "de": f"- Arbeiten aus {END_YEAR} sind strukturell benachteiligt, weil das Jahr nur teilweise abgeschlossen ist und die Zitierakkumulation am {date.today().isoformat()} noch unreif ist.",
            "es": f"- Los artículos de {END_YEAR} tienen desventaja estructural porque el año es parcial y la acumulación de citas aún es inmadura al {date.today().isoformat()}.",
            "fr": f"- Les articles de {END_YEAR} sont structurellement désavantagés car l'année est partielle et l'accumulation de citations reste immature au {date.today().isoformat()}.",
            "ja": f"- {END_YEAR} 年の論文は、年が途中であり {date.today().isoformat()} 時点で引用蓄積がまだ成熟していないため、構造的に不利です。",
            "ko": f"- {END_YEAR}년 논문은 연도가 아직 진행 중이고 {date.today().isoformat()} 기준 인용 축적이 충분히 성숙하지 않아 구조적으로 불리합니다.",
            "pt": f"- Artigos de {END_YEAR} têm desvantagem estrutural porque o ano é parcial e a acumulação de citações ainda está imatura em {date.today().isoformat()}.",
            "ru": f"- Работы {END_YEAR} года структурно находятся в невыгодном положении, поскольку год неполный, а накопление цитирований на {date.today().isoformat()} еще незрелое.",
            "zh": f"- {END_YEAR} 年论文在结构上处于引用劣势，因为该年份尚未完整，且截至 {date.today().isoformat()} 引用积累仍不成熟。",
        }[language],
        1,
    )
    text = text.replace(
        "- Stock-investment research is especially sensitive to transaction costs, survivorship bias, look-ahead bias, market impact, and regime change; those require full-paper and data-level review.",
        {
            "de": "- Aktienanlageforschung ist besonders empfindlich gegenüber Transaktionskosten, Survivorship Bias, Look-ahead Bias, Markteinfluss und Regimewechseln; diese Punkte erfordern eine Prüfung auf Paper- und Datenebene.",
            "es": "- La investigación de inversión en acciones es especialmente sensible a costes de transacción, sesgo de supervivencia, look-ahead bias, impacto de mercado y cambio de régimen; todo ello requiere revisión a nivel de artículo y datos.",
            "fr": "- La recherche sur l'investissement actions est particulièrement sensible aux coûts de transaction, au biais de survie, au look-ahead bias, à l'impact de marché et aux changements de régime; ces points exigent une revue au niveau des articles et des données.",
            "ja": "- 株式投資研究は、取引コスト、サバイバーシップバイアス、ルックアヘッドバイアス、マーケットインパクト、レジーム変化に特に敏感であり、論文本文とデータレベルの確認が必要です。",
            "ko": "- 주식 투자 연구는 거래비용, 생존편향, look-ahead bias, 시장충격, regime change에 특히 민감하므로 논문 전문과 데이터 수준 검토가 필요합니다.",
            "pt": "- A pesquisa de investimento em ações é especialmente sensível a custos de transação, viés de sobrevivência, look-ahead bias, impacto de mercado e mudança de regime; isso exige revisão no nível do artigo e dos dados.",
            "ru": "- Исследования инвестиций в акции особенно чувствительны к транзакционным издержкам, survivorship bias, look-ahead bias, market impact и смене режимов; это требует проверки на уровне полной статьи и данных.",
            "zh": "- 股票投资研究尤其容易受到交易成本、生存者偏差、前视偏差、市场冲击和 regime change 的影响，因此需要论文全文和数据层面的审查。",
        }[language],
        1,
    )
    text = text.replace(
        "This repository and interactive site were created with appreciation for [jehyunlee/paper-curation](https://github.com/jehyunlee/paper-curation). Its paper-curation workflow and repository organization informed the approach used here for a taxonomy-first, citation-ranked research map.",
        {
            "de": "Dieses Repository und die interaktive Website wurden mit Dank an [jehyunlee/paper-curation](https://github.com/jehyunlee/paper-curation) erstellt. Dessen Workflow und Repository-Struktur haben den hier verwendeten Ansatz für eine taxonomieorientierte, zitiergerankte Forschungskarte geprägt.",
            "es": "Este repositorio y el sitio interactivo se crearon con agradecimiento a [jehyunlee/paper-curation](https://github.com/jehyunlee/paper-curation). Su flujo de curación de papers y organización del repositorio inspiraron el enfoque de este mapa de investigación taxonómico y ordenado por citas.",
            "fr": "Ce dépôt et le site interactif ont été créés avec reconnaissance envers [jehyunlee/paper-curation](https://github.com/jehyunlee/paper-curation). Son workflow de curation d'articles et son organisation de dépôt ont inspiré l'approche utilisée ici.",
            "ja": "このリポジトリとインタラクティブサイトは [jehyunlee/paper-curation](https://github.com/jehyunlee/paper-curation) への謝意とともに作成しました。その paper-curation ワークフローとリポジトリ構成が、本プロジェクトの分類体系優先・引用順位型研究マップの設計に影響を与えています。",
            "ko": "이 저장소와 인터랙티브 사이트는 [jehyunlee/paper-curation](https://github.com/jehyunlee/paper-curation)에 대한 감사와 함께 만들었습니다. 해당 paper-curation 워크플로와 저장소 구성이 이곳의 분류 우선, 인용 순위 기반 연구 지도 접근에 참고가 되었습니다.",
            "pt": "Este repositório e o site interativo foram criados com agradecimento a [jehyunlee/paper-curation](https://github.com/jehyunlee/paper-curation). Seu fluxo de curadoria de papers e organização do repositório inspiraram a abordagem usada aqui.",
            "ru": "Этот репозиторий и интерактивный сайт созданы с благодарностью к [jehyunlee/paper-curation](https://github.com/jehyunlee/paper-curation). Его workflow и структура репозитория повлияли на подход к этой таксономической карте исследований, ранжированной по цитированиям.",
            "zh": "本仓库和交互式网站在致谢 [jehyunlee/paper-curation](https://github.com/jehyunlee/paper-curation) 的基础上创建。其论文策展流程和仓库组织方式启发了这里按分类优先、按引用排序的研究地图方法。",
        }[language],
        1,
    )
    text = text.replace(
        "CC-BY-4.0 for text and metadata curation. Upstream paper metadata belongs to the original sources.",
        {
            "de": "Text und Metadatenkuratierung stehen unter CC-BY-4.0. Die ursprünglichen Paper-Metadaten gehören den jeweiligen Quellen.",
            "es": "CC-BY-4.0 para texto y curación de metadatos. Los metadatos originales pertenecen a sus fuentes.",
            "fr": "CC-BY-4.0 pour le texte et la curation des métadonnées. Les métadonnées amont appartiennent aux sources originales.",
            "ja": "テキストおよびメタデータのキュレーションは CC-BY-4.0 です。元の論文メタデータは各ソースに帰属します。",
            "ko": "텍스트와 메타데이터 큐레이션은 CC-BY-4.0을 따릅니다. 원천 논문 메타데이터는 각 원본 출처에 속합니다.",
            "pt": "CC-BY-4.0 para texto e curadoria de metadados. Os metadados originais pertencem às fontes originais.",
            "ru": "Текст и курирование метаданных распространяются по CC-BY-4.0. Исходные метаданные статей принадлежат оригинальным источникам.",
            "zh": "文本和元数据策展采用 CC-BY-4.0。上游论文元数据归原始来源所有。",
        }[language],
        1,
    )
    for old, new in README_STATIC_REPLACEMENTS[language]:
        text = text.replace(old, new)
    text = text.replace("Show representative papers for ", copy["show"])
    text = re.sub(r"(\d[\d,]*) papers\b", lambda match: f"{match.group(1)} {copy['paper_word']}", text)
    text = localize_readme_keywords(text, language)
    return text


def write_multilingual_readmes(readme_path, selected, candidates):
    english_text = readme_path.read_text(encoding="utf-8")
    for language, filename in README_LANGUAGE_FILES.items():
        if language == "en":
            continue
        localized_path = ROOT / filename
        localized_path.write_text(localized_readme_text(english_text, language, selected, candidates), encoding="utf-8")
        write_markdown_html(localized_path, f"{PROJECT_TITLE} ({README_LANGUAGE_LABELS[language]})", lang=language)


def write_readme(selected, candidates):
    total_cites = sum(p["citationCount"] for p in selected)
    candidate_counts = Counter(p["year"] for p in candidates)
    stats = year_stats(selected)
    lines = [
        f"# {PROJECT_TITLE}",
        "",
        "[![Awesome](https://awesome.re/badge-flat.svg)](https://awesome.re)",
        "",
        PROJECT_DESCRIPTION,
        "",
        '<p align="center">',
        f'  <a href="{PROJECT_PAGES_URL}">',
        f'    <img src="https://img.shields.io/badge/Open_Interactive_Website-{PROJECT_OWNER}.github.io%2F{PROJECT_REPO.replace("-", "--")}-0f766e?style=for-the-badge" alt="Open Interactive Website">',
        "  </a>",
        "</p>",
        "",
        *readme_language_selector_lines(PROJECT_OWNER, PROJECT_REPO),
        f"Generated on {date.today().isoformat()} from free public Semantic Scholar metadata. This edition investigates up to {CANDIDATES_PER_YEAR:,} stock-investment and AI-trading candidate papers per publication year for {YEAR_RANGE_TEXT}, keeps an audited candidate pool of {len(candidates):,} records, selects the top {TARGET_PER_YEAR:,} papers from each year by citation count ({len(selected):,} papers selected), and reorganizes them by research taxonomy.",
        "",
        "## Project Links",
        "",
        f"- Open Interactive Website: {PROJECT_PAGES_URL}",
        f"- Selected dataset: `data/{PAPERS_CSV}`",
        f"- Taxonomy dataset with paper-level ideas, strengths, and limitations: `data/{TAXONOMY_CSV}`",
        f"- Precomputed period analysis: `data/{PERIOD_ANALYSIS_JSON}`",
        f"- Candidate Pool: `data/{CANDIDATES_CSV}`",
        "- English review draft: `paper/review_en.html`, `paper/review_en.docx`",
        "- Korean review draft: `paper/review_ko.html`",
        "",
        "## Keywords Convention",
        "",
        "These badges define the stock-investment and AI-trading keyword tags used to read and extend this collection.",
        "",
        *readme_keyword_convention_lines(),
        "",
        "## Taxonomy Overview",
        "",
        f"- **Total selected papers**: {len(selected):,} papers",
        f"- **Candidate pool audited**: {len(candidates):,} papers ({', '.join(f'{year}: {candidate_counts[year]:,}' for year in YEARS)})",
        f"- **Citation count in selected set**: {total_cites:,}",
    ]
    for category, count in category_stats(selected).most_common():
        lines.append(f"- **{category}**: {count:,} papers")
    lines.extend(["", "## Taxonomy Collections", ""])
    groups = category_groups(selected)
    for category, count in category_stats(selected).most_common():
        rows = groups[category]
        citations = sum(p["citationCount"] for p in rows)
        years = sorted({p["year"] for p in rows})
        lines.extend(
            [
                f"### {category}",
                "",
                f"- Papers selected: **{len(rows):,}**",
                f"- Years covered: **{years[0]}-{years[-1]}**",
                f"- Citation count in selected set: **{citations:,}**",
                "- Category Overview:",
                *[f"  - {item}" for item in TAXONOMY_TRENDS.get(category, [])],
                "- Limitations:",
                *[f"  - {item}" for item in TAXONOMY_LIMITATIONS.get(category, [])],
                "",
                f"<details>",
                f"<summary><strong>Show representative papers for {category}</strong></summary>",
                "",
                readme_taxonomy_table(rows),
                "",
                "</details>",
                "",
            ]
        )
    lines.extend(["## Yearly Selection Summary", "", "| Year | Candidate records | Selected papers | Selected citations | Top selected paper |", "|---:|---:|---:|---:|---|"])
    for year in YEARS:
        stat = stats.get(year)
        top = md_link(stat["top"]["title"], stat["top"]["url"]) if stat else "n/a"
        lines.append(f"| {year} | {candidate_counts[year]:,} | {stat['count'] if stat else 0:,} | {stat['citations'] if stat else 0:,} | {top} |")
    lines.extend(
        [
            "",
            "## Curation Method",
            "",
            f"The collection uses Semantic Scholar Academic Graph bulk search. Queries cover stock prediction, equity returns, asset pricing, portfolio optimization, financial time series, algorithmic trading, reinforcement learning, sentiment/news/alternative data, market microstructure, high-frequency trading, volatility, and risk themes. For each year from {START_YEAR} through {END_YEAR}, results are filtered to the publication year, screened with explicit stock-investment and AI-trading relevance expressions in title/abstract/venue metadata, deduplicated by DOI, arXiv, PubMed, CorpusId, paperId, then title, and reduced to at most {CANDIDATES_PER_YEAR:,} candidates by citation count. The final awesome list selects the top {TARGET_PER_YEAR:,} papers within each publication year by citation count; influential citation count and a deterministic metadata importance score are retained as tie-breakers and audit signals.",
            "",
            "The taxonomy, key ideas, strengths, limitations, method tags, and keyword tags are generated deterministically from public metadata and rule-based domain conventions. No paid API, paid LLM, paid translation, or paid compute was used.",
            "",
            "## Limitations",
            "",
            "- This is a metadata-driven citation map, not a full systematic review of every PDF.",
            "- Citation count is an influence signal, not investment advice or proof of live trading profitability.",
            "- Semantic Scholar metadata can omit venues, abstracts, PDFs, or influential citation counts for some records.",
            f"- Papers from {END_YEAR} are structurally citation-disadvantaged because the year is partial and citation accumulation is still immature as of {date.today().isoformat()}.",
            "- Stock-investment research is especially sensitive to transaction costs, survivorship bias, look-ahead bias, market impact, and regime change; those require full-paper and data-level review.",
            "",
            "## Acknowledgements",
            "",
            "This repository and interactive site were created with appreciation for [jehyunlee/paper-curation](https://github.com/jehyunlee/paper-curation). Its paper-curation workflow and repository organization informed the approach used here for a taxonomy-first, citation-ranked research map.",
            "",
            "## License",
            "",
            "CC-BY-4.0 for text and metadata curation. Upstream paper metadata belongs to the original sources.",
        ]
    )
    readme_path = ROOT / "README.md"
    readme_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    write_markdown_html(readme_path, PROJECT_TITLE)
    write_multilingual_readmes(readme_path, selected, candidates)


def review_sections(selected, korean=False):
    stats = year_stats(selected)
    cats = category_stats(selected)
    total_cites = sum(p["citationCount"] for p in selected)
    top_cited = sorted(selected, key=lambda p: p["citationCount"], reverse=True)[:12]
    top_scored = sorted(selected, key=lambda p: p["importanceScore"], reverse=True)[:12]
    peak_year = max(stats, key=lambda y: stats[y]["citations"])
    leading_cat, leading_count = cats.most_common(1)[0]
    if korean:
        title = f"{YEAR_RANGE_TEXT} 주식 투자 및 AI 트레이딩 연구 동향: 공개 메타데이터 기반 citation-ranked 리뷰"
        abstract = (
            f"이 리뷰 초안은 {START_YEAR}년부터 {END_YEAR}년까지 주식 투자, 주가/수익률 예측, 포트폴리오, "
            f"알고리즘 트레이딩, AI 기반 투자 연구를 연도별 최대 {CANDIDATES_PER_YEAR:,}편의 후보 논문으로 조사하고, "
            f"각 연도에서 citation이 높은 {TARGET_PER_YEAR:,}편을 선정해 taxonomy-first 방식으로 정리한다."
        )
        methods = (
            "Semantic Scholar 공개 메타데이터에 stock prediction, asset pricing, portfolio optimization, "
            "machine learning, deep learning, reinforcement learning trading, sentiment/news, high-frequency trading, "
            "volatility/risk 관련 질의를 보내고, 명시적 관련성 필터와 중복 제거를 거쳐 인용수 기준으로 정렬했다."
        )
        findings = [
            f"선정 논문 {len(selected):,}편은 총 {total_cites:,}회의 인용을 포함하며, citation mass가 가장 큰 연도는 {peak_year}년이다.",
            f"가장 큰 분류는 {KOREAN_CATEGORY_NAMES.get(leading_cat, leading_cat)}({leading_count:,}편)이다.",
            "AI 기반 주식 투자 연구는 예측 정확도뿐 아니라 거래비용, turnover, market impact, regime shift 검증이 함께 필요하다.",
            f"{END_YEAR}년 논문은 아직 인용 축적 시간이 짧으므로 최신성은 별도 전문가 검토가 필요하다.",
        ]
        caveat = "이 문서는 PDF 전문 기반 systematic review가 아니라 공개 메타데이터 기반 citation map이다. 인용수는 영향력 신호일 뿐 투자 조언이나 실거래 수익성을 보장하지 않는다."
        conclusion = "주식 투자 연구는 전통적 자산가격결정과 포트폴리오 이론 위에 머신러닝, 딥러닝, 텍스트/대체데이터, 강화학습 기반 의사결정이 결합되는 방향으로 확장되고 있다."
    else:
        title = f"Stock Investment and AI Trading Research from {START_YEAR} to {END_YEAR}: A Metadata-Driven Citation Map"
        abstract = (
            f"This draft review maps stock investment and AI-driven trading research from {START_YEAR} through {END_YEAR}, "
            f"investigating up to {CANDIDATES_PER_YEAR:,} candidate papers per year from free public Semantic Scholar metadata "
            f"and selecting the top {TARGET_PER_YEAR:,} papers from each year by citation count ({len(selected):,} papers selected)."
        )
        methods = (
            "Queries covered stock prediction, asset pricing, portfolio optimization, financial time series, machine learning, "
            "deep learning, reinforcement learning trading, sentiment/news, market microstructure, high-frequency trading, volatility, and risk. "
            "Records were filtered by publication year, screened for explicit stock-investment or AI-trading relevance, deduplicated, and ranked by citation count."
        )
        findings = [
            f"The {len(selected):,} selected papers account for {total_cites:,} citations in the selected set, with the largest citation mass in {peak_year}.",
            f"The largest category is {leading_cat} ({leading_count:,} papers).",
            "AI-based stock-investment work broadens the signal surface, but economic value still depends on costs, turnover, capacity, and risk.",
            f"Papers from {END_YEAR} are structurally citation-disadvantaged because citation accumulation is still immature.",
        ]
        caveat = "This is a metadata-driven citation map rather than a full systematic review of every PDF. Citation count is an influence signal, not investment advice or evidence of live trading profitability."
        conclusion = "Stock-investment research increasingly combines asset-pricing and portfolio foundations with machine learning, deep learning, alternative data, and sequential trading decisions."
    category_lines = [f"{cat}: {count}" for cat, count in cats.most_common()]
    year_lines = [
        f"{year}: {stats[year]['count']} selected papers, {stats[year]['citations']:,} citations, top selected paper: {stats[year]['top']['title']}"
        for year in YEARS
        if year in stats
    ]
    return {
        "title": title,
        "abstract": abstract,
        "methods": methods,
        "findings": findings,
        "category_lines": category_lines,
        "year_lines": year_lines,
        "top_cited": top_cited,
        "top_scored": top_scored,
        "caveat": caveat,
        "conclusion": conclusion,
    }


def review_sections(selected, korean=False):
    stats = year_stats(selected)
    cats = category_stats(selected)
    total_cites = sum(p["citationCount"] for p in selected)
    top_cited = sorted(selected, key=lambda p: p["citationCount"], reverse=True)[:12]
    top_scored = sorted(selected, key=lambda p: p["importanceScore"], reverse=True)[:12]
    peak_year = max(stats, key=lambda y: stats[y]["citations"])
    leading_cat, leading_count = cats.most_common(1)[0]
    if korean:
        title = f"{YEAR_RANGE_TEXT} 주식 투자 및 AI 트레이딩 연구 동향: 공개 메타데이터 기반 citation-ranked 리뷰"
        abstract = (
            f"이 리뷰 초안은 {START_YEAR}년부터 {END_YEAR}년까지 주식 투자, 주가 및 수익률 예측, 포트폴리오 최적화, "
            f"알고리즘 트레이딩, AI 기반 투자 연구를 연도별 최대 {CANDIDATES_PER_YEAR:,}편의 후보 논문으로 조사하고, "
            f"각 연도에서 인용수가 높은 상위 {TARGET_PER_YEAR:,}편을 선정해 taxonomy-first 방식으로 정리한다."
        )
        methods = (
            "Semantic Scholar 공개 메타데이터에 stock prediction, asset pricing, portfolio optimization, "
            "machine learning, deep learning, reinforcement learning trading, sentiment/news, high-frequency trading, "
            "volatility/risk 관련 질의를 적용했다. 기록은 출판연도별로 필터링하고 명시적 관련성 조건과 중복 제거를 거친 뒤 인용수 기준으로 정렬했다."
        )
        findings = [
            f"선정 논문 {len(selected):,}편은 총 {total_cites:,}회의 인용을 포함하며, 선정 세트의 citation mass가 가장 큰 연도는 {peak_year}년이다.",
            f"가장 큰 분류는 {KOREAN_CATEGORY_NAMES.get(leading_cat, leading_cat)}({leading_count:,}편)이다.",
            "AI 기반 주식 투자 연구는 신호 공간을 넓히지만, 실제 경제적 가치는 거래비용, turnover, market impact, capacity, regime shift 검증에 크게 좌우된다.",
            f"{END_YEAR}년 논문은 아직 인용 축적 기간이 짧으므로 최신성은 별도 전문가 검토로 보완해야 한다.",
        ]
        caveat = (
            "이 문서는 PDF 전문 기반 systematic review가 아니라 공개 메타데이터 기반 citation map이다. "
            "인용수는 영향력 신호일 뿐 투자 조언이나 실거래 수익성을 보장하지 않는다."
        )
        conclusion = (
            "주식 투자 연구는 전통적 자산가격결정과 포트폴리오 이론 위에 머신러닝, 딥러닝, 텍스트 및 대체데이터, "
            "강화학습 기반 순차 의사결정을 결합하는 방향으로 확장되고 있다."
        )
    else:
        title = f"Stock Investment and AI Trading Research from {START_YEAR} to {END_YEAR}: A Metadata-Driven Citation Map"
        abstract = (
            f"This draft review maps stock investment and AI-driven trading research from {START_YEAR} through {END_YEAR}, "
            f"investigating up to {CANDIDATES_PER_YEAR:,} candidate papers per year from free public Semantic Scholar metadata "
            f"and selecting the top {TARGET_PER_YEAR:,} papers from each year by citation count ({len(selected):,} papers selected)."
        )
        methods = (
            "Queries covered stock prediction, asset pricing, portfolio optimization, financial time series, machine learning, "
            "deep learning, reinforcement learning trading, sentiment/news, market microstructure, high-frequency trading, volatility, and risk. "
            "Records were filtered by publication year, screened for explicit stock-investment or AI-trading relevance, deduplicated, and ranked by citation count."
        )
        findings = [
            f"The {len(selected):,} selected papers account for {total_cites:,} citations in the selected set, with the largest citation mass in {peak_year}.",
            f"The largest category is {leading_cat} ({leading_count:,} papers).",
            "AI-based stock-investment work broadens the signal surface, but economic value still depends on costs, turnover, capacity, and risk.",
            f"Papers from {END_YEAR} are structurally citation-disadvantaged because citation accumulation is still immature.",
        ]
        caveat = "This is a metadata-driven citation map rather than a full systematic review of every PDF. Citation count is an influence signal, not investment advice or evidence of live trading profitability."
        conclusion = "Stock-investment research increasingly combines asset-pricing and portfolio foundations with machine learning, deep learning, alternative data, and sequential trading decisions."
    category_lines = [
        f"{KOREAN_CATEGORY_NAMES.get(cat, cat) if korean else cat}: {count}"
        for cat, count in cats.most_common()
    ]
    year_lines = [
        f"{year}: {stats[year]['count']} selected papers, {stats[year]['citations']:,} citations, top selected paper: {stats[year]['top']['title']}"
        for year in YEARS
        if year in stats
    ]
    return {
        "title": title,
        "abstract": abstract,
        "methods": methods,
        "findings": findings,
        "category_lines": category_lines,
        "year_lines": year_lines,
        "top_cited": top_cited,
        "top_scored": top_scored,
        "caveat": caveat,
        "conclusion": conclusion,
    }


def write_project_files(selected, candidates):
    citation = f"""cff-version: 1.2.0
title: "Awesome Trade: A Metadata-Driven Citation Map of Stock Investment and AI Trading Research, {YEAR_RANGE_TEXT}"
message: "If you use this curated dataset or review draft, please cite this repository."
type: dataset
authors:
  - name: "Honggi"
repository-code: "{PROJECT_GITHUB_URL}"
date-released: "{date.today().isoformat()}"
license: "CC-BY-4.0"
keywords:
  - "stock investment"
  - "AI trading"
  - "stock prediction"
  - "portfolio optimization"
  - "bibliometrics"
"""
    (ROOT / "CITATION.cff").write_text(citation, encoding="utf-8")
    (ROOT / "LICENSE").write_text("CC-BY-4.0 for text and metadata curation; upstream paper metadata belongs to original sources.\n", encoding="utf-8")
    (ROOT / ".gitignore").write_text("__pycache__/\n*.pyc\n.tools/\ndata/cache/\n.playwright-cli/\noutput/playwright/\nbuild.log\nbuild.err.log\n", encoding="utf-8")
    publish = f"""@echo off
setlocal
cd /d "%~dp0"

set "GH_EXE=%~dp0.tools\\gh\\bin\\gh.exe"
if not exist "%GH_EXE%" if exist "%~dp0..\\awesome-BCI\\.tools\\gh\\bin\\gh.exe" set "GH_EXE=%~dp0..\\awesome-BCI\\.tools\\gh\\bin\\gh.exe"
if not exist "%GH_EXE%" set "GH_EXE=gh"

"%GH_EXE%" auth status
if errorlevel 1 (
  echo.
  echo GitHub login is required. Run:
  echo   "%GH_EXE%" auth login --hostname github.com --web --scopes repo
  exit /b 1
)

"%GH_EXE%" repo view {PROJECT_OWNER}/{PROJECT_REPO} >nul 2>nul
if errorlevel 1 (
  "%GH_EXE%" repo create {PROJECT_OWNER}/{PROJECT_REPO} --public --description "Awesome Trade: metadata-driven stock investment and AI trading paper curation, {YEAR_RANGE_TEXT}" --source . --remote origin --push
) else (
  git remote set-url origin {PROJECT_GITHUB_URL}.git
  git push -u origin main
)
if errorlevel 1 exit /b %errorlevel%

"%GH_EXE%" api repos/{PROJECT_OWNER}/{PROJECT_REPO}/pages -X POST -f "source[branch]=main" -f "source[path]=/docs" >nul 2>nul
if errorlevel 1 (
  "%GH_EXE%" api repos/{PROJECT_OWNER}/{PROJECT_REPO}/pages -X PUT -f "source[branch]=main" -f "source[path]=/docs" >nul 2>nul
)

echo.
echo Done: {PROJECT_GITHUB_URL}
echo Pages: {PROJECT_PAGES_URL}
"""
    (ROOT / "publish_to_github.bat").write_text(publish, encoding="utf-8")
    method = f"""# Curation Method

## Scope

- Topic: stock investment and AI-driven trading research across stock prediction, asset pricing, portfolio optimization, financial time series, market microstructure, risk, sentiment/news, and algorithmic trading.
- Period: {YEAR_RANGE_TEXT}.
- Candidate target: up to {CANDIDATES_PER_YEAR:,} papers per year.
- Final selection: top {TARGET_PER_YEAR:,} papers per year by citation count from the audited yearly candidate pools.

## Data Source

Metadata comes from the free public Semantic Scholar Academic Graph bulk search endpoint. No paid API, paid LLM, paid translation, or paid compute was used.

## Ranking

Records are filtered to the requested publication year, screened for explicit stock-investment and AI-trading relevance in title/abstract/venue metadata, deduplicated by DOI, arXiv, PubMed, CorpusId, paperId, and normalized title, then ranked by citation count. Influential citation count and a deterministic metadata importance score are retained as audit fields.

## Enrichment

Taxonomy, key ideas, strengths, limitations, method tags, and keyword convention tags are generated with deterministic rules from public metadata. This repository is a metadata-driven citation map, not a full-text PDF systematic review.

## GitHub-Awesome Skill2 and Paper-Curation Provenance

This regeneration follows `github-awesome-skill2` in metadata-adapter mode for a large citation-ranked awesome repository. The workflow inspected the local `jehyunlee/paper-curation` checkout and is configured for Zotero-free folder-source PDF staging under `{(ROOT.parent / "paper-curation" / "paper" / PROJECT_REPO).resolve()}`. Full PDF LLM review stages from paper-curation were not run because they require separate explicit approval for paid or metered APIs.
"""
    PAPER_DIR.mkdir(exist_ok=True)
    (PAPER_DIR / "curation_method.md").write_text(method, encoding="utf-8")
    write_curation_method_html(method)


def write_skill2_provenance(selected, candidates):
    payload = {
        "skill": "github-awesome-skill2",
        "mode": "metadata-adapter",
        "paper_curation_source": str((ROOT.parent / "paper-curation").resolve()),
        "zotero_used": False,
        "paid_or_metered_api_used": False,
        "folder_source_pdf_dir": str((ROOT.parent / "paper-curation" / "paper" / PROJECT_REPO).resolve()),
        "selected_dataset": f"data/{PAPERS_CSV}",
        "candidate_dataset": f"data/{CANDIDATES_CSV}",
        "selected_papers": len(selected),
        "candidate_records": len(candidates),
        "period": YEAR_RANGE_TEXT,
        "candidate_target_per_year": CANDIDATES_PER_YEAR,
        "selection_target_per_year": TARGET_PER_YEAR,
        "ranking": "citationCount descending with influentialCitationCount and metadata importance score retained as audit signals",
        "note": "The repository/site outputs are deterministic metadata curation artifacts; full PDF LLM reviews require separate explicit approval.",
    }
    (DATA_DIR / SKILL2_PROVENANCE_JSON).write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def postprocess_generated_files():
    replacements = {
        "Awesome Test": PROJECT_TITLE,
        "awesome-test": PROJECT_REPO,
        "AI research": "stock investment and AI-driven trading research",
        "AI-specific convention": "trade-specific convention",
        "AI Paper Taxonomy": "Trade Paper Taxonomy",
        "https://github.com/honggi82/awesome-test": PROJECT_GITHUB_URL,
    }
    targets = [
        DOCS_DIR / "index.html",
        DOCS_DIR / "paper" / "review_en.html",
        DOCS_DIR / "paper" / "review_ko.html",
        DOCS_DIR / "paper" / "curation_method.html",
    ]
    for path in targets:
        if not path.exists():
            continue
        text = path.read_text(encoding="utf-8")
        for old, new in replacements.items():
            text = text.replace(old, new)
        path.write_text(text, encoding="utf-8")


LANGUAGES = {
    "en": "English",
    "ko": "한국어",
    "zh": "中文",
    "ja": "日本語",
}

UI_LABELS = {
    "en": {
        "papers": "papers",
        "categories": "categories",
        "overview": "Category Overview",
        "limitations": "Limitations",
        "analysis": "Selected-period analysis",
        "totalSelected": "Total selected papers",
        "categoryCount": "Categories",
        "keyIdea": "Key idea",
        "strengths": "Strengths",
        "paperLimitations": "Limitations",
        "citations": "citations",
        "selectedKeyword": "Selected keyword",
        "matchingPapers": "Matching papers",
        "allKeywords": "all",
        "categoryChart": "Category distribution",
        "citationChart": "Yearly citation mass",
        "topPaper": "Top paper",
        "allTaxonomies": "All Taxonomies",
    },
    "ko": {
        "papers": "편",
        "categories": "분류",
        "overview": "분류 개요",
        "limitations": "한계",
        "analysis": "선택 기간 분석",
        "totalSelected": "선정 논문",
        "categoryCount": "분류",
        "keyIdea": "핵심 아이디어",
        "strengths": "강점",
        "paperLimitations": "한계",
        "citations": "회 인용",
        "selectedKeyword": "선택 키워드",
        "matchingPapers": "일치 논문",
        "allKeywords": "전체",
        "categoryChart": "분류별 분포",
        "citationChart": "연도별 인용 규모",
        "topPaper": "대표 논문",
        "allTaxonomies": "전체 분류",
    },
    "zh": {
        "papers": "篇论文",
        "categories": "个分类",
        "overview": "分类概览",
        "limitations": "局限",
        "analysis": "所选期间分析",
        "totalSelected": "入选论文",
        "categoryCount": "分类",
        "keyIdea": "核心观点",
        "strengths": "优势",
        "paperLimitations": "局限",
        "citations": "次引用",
        "selectedKeyword": "选中关键词",
        "matchingPapers": "匹配论文",
        "allKeywords": "全部",
        "categoryChart": "分类分布",
        "citationChart": "年度引用规模",
        "topPaper": "代表论文",
        "allTaxonomies": "全部分类",
    },
    "ja": {
        "papers": "本",
        "categories": "分類",
        "overview": "分類概要",
        "limitations": "限界",
        "analysis": "選択期間の分析",
        "totalSelected": "選定論文",
        "categoryCount": "分類",
        "keyIdea": "主要アイデア",
        "strengths": "強み",
        "paperLimitations": "限界",
        "citations": "件の引用",
        "selectedKeyword": "選択キーワード",
        "matchingPapers": "該当論文",
        "allKeywords": "すべて",
        "categoryChart": "分類別分布",
        "citationChart": "年別引用規模",
        "topPaper": "代表論文",
        "allTaxonomies": "全分類",
    },
}

CATEGORY_NAMES = {
    "ko": {
        "Asset Pricing and Return Predictability": "자산가격결정과 수익률 예측",
        "Portfolio Optimization and Asset Allocation": "포트폴리오 최적화와 자산배분",
        "Machine Learning for Stock Prediction": "머신러닝 기반 주식 예측",
        "Deep Learning and Financial Time Series": "딥러닝과 금융 시계열",
        "Reinforcement Learning and Algorithmic Trading": "강화학습과 알고리즘 트레이딩",
        "Sentiment, News, and Alternative Data": "감성, 뉴스, 대체데이터",
        "Market Microstructure and High-Frequency Trading": "시장 미시구조와 고빈도 거래",
        "Risk, Volatility, and Forecast Evaluation": "위험, 변동성, 예측 평가",
        "Behavioral Finance and Investor Decision-Making": "행동재무와 투자자 의사결정",
        "General Finance, Surveys, and Trading Systems": "일반 금융, 서베이, 거래 시스템",
    },
    "zh": {
        "Asset Pricing and Return Predictability": "资产定价与收益预测",
        "Portfolio Optimization and Asset Allocation": "投资组合优化与资产配置",
        "Machine Learning for Stock Prediction": "用于股票预测的机器学习",
        "Deep Learning and Financial Time Series": "深度学习与金融时间序列",
        "Reinforcement Learning and Algorithmic Trading": "强化学习与算法交易",
        "Sentiment, News, and Alternative Data": "情绪、新闻与另类数据",
        "Market Microstructure and High-Frequency Trading": "市场微观结构与高频交易",
        "Risk, Volatility, and Forecast Evaluation": "风险、波动率与预测评估",
        "Behavioral Finance and Investor Decision-Making": "行为金融与投资者决策",
        "General Finance, Surveys, and Trading Systems": "通用金融、综述与交易系统",
    },
    "ja": {
        "Asset Pricing and Return Predictability": "資産価格評価とリターン予測",
        "Portfolio Optimization and Asset Allocation": "ポートフォリオ最適化と資産配分",
        "Machine Learning for Stock Prediction": "株式予測のための機械学習",
        "Deep Learning and Financial Time Series": "深層学習と金融時系列",
        "Reinforcement Learning and Algorithmic Trading": "強化学習とアルゴリズム取引",
        "Sentiment, News, and Alternative Data": "センチメント、ニュース、代替データ",
        "Market Microstructure and High-Frequency Trading": "市場マイクロストラクチャーと高頻度取引",
        "Risk, Volatility, and Forecast Evaluation": "リスク、ボラティリティ、予測評価",
        "Behavioral Finance and Investor Decision-Making": "行動ファイナンスと投資家意思決定",
        "General Finance, Surveys, and Trading Systems": "一般金融、サーベイ、取引システム",
    },
}

KOREAN_CATEGORY_NAMES = CATEGORY_NAMES["ko"]

KEYWORD_DESCRIPTION_LANG = {
    "ko": {
        "stock-prediction": "주가, 수익률, 방향성, 추세, 시장지수 예측과 forecasting.",
        "machine-learning": "SVM, 트리, 부스팅, 랜덤 포레스트, 커널, 데이터마이닝 등 전통적 ML 기반 시장 분석.",
        "deep-learning": "LSTM/GRU/CNN, attention, transformer, 표현학습 기반 금융 시계열 모델.",
        "ai-based-trade": "AI/ML 예측 모델, NLP/LLM, 자동화 전략을 활용한 주식 거래와 투자 신호.",
        "portfolio": "포트폴리오 최적화, 자산배분, 포트폴리오 선택, 리스크 패리티, 리밸런싱.",
        "reinforcement-trading": "강화학습, 알고리즘 트레이딩, 실행, market making, 백테스트된 거래 전략.",
        "sentiment-altdata": "뉴스, 소셜미디어, NLP, LLM, 실적발표, 대체데이터 기반 투자 신호.",
        "high-frequency": "시장 미시구조, limit order book, order flow, 유동성, 고빈도 거래.",
        "risk-volatility": "변동성 예측, VaR, drawdown, tail risk, stress test, 예측 평가.",
        "asset-pricing": "자산가격결정, 주식 팩터, anomaly, 수익률 예측성, 횡단면 주식 수익률.",
        "behavioral-finance": "투자자 심리, attention, 행동 편향, 시장 효율성, 의사결정.",
    },
    "zh": {
        "stock-prediction": "股票价格、收益、方向、趋势或市场指数预测。",
        "machine-learning": "SVM、树模型、提升方法、随机森林、核方法和数据挖掘在市场中的应用。",
        "deep-learning": "神经网络、LSTM/GRU/CNN、注意力、Transformer 与金融时间序列表征学习。",
        "ai-based-trade": "使用 AI/ML 预测模型、NLP/LLM 和自动化策略形成股票交易与投资信号。",
        "portfolio": "投资组合优化、资产配置、组合选择、风险平价与再平衡。",
        "reinforcement-trading": "强化学习、算法交易、执行、做市与回测交易策略。",
        "sentiment-altdata": "新闻、社交媒体、NLP、LLM、财报电话会与另类数据投资信号。",
        "high-frequency": "市场微观结构、限价订单簿、订单流、流动性与高频交易。",
        "risk-volatility": "波动率预测、VaR、回撤、尾部风险、压力测试与预测评估。",
        "asset-pricing": "资产定价、股票因子、异象、收益可预测性与横截面股票收益。",
        "behavioral-finance": "投资者情绪、注意力、行为偏差、市场效率与决策。",
    },
    "ja": {
        "stock-prediction": "株価、リターン、方向、トレンド、市場指数の予測。",
        "machine-learning": "SVM、木モデル、ブースティング、ランダムフォレスト、カーネル、データマイニングの市場応用。",
        "deep-learning": "ニューラルネット、LSTM/GRU/CNN、attention、Transformer、金融時系列の表現学習。",
        "ai-based-trade": "AI/ML 予測モデル、NLP/LLM、自動化戦略を用いた株式取引と投資シグナル。",
        "portfolio": "ポートフォリオ最適化、資産配分、ポートフォリオ選択、リスクパリティ、リバランス。",
        "reinforcement-trading": "強化学習、アルゴリズム取引、執行、マーケットメイク、バックテスト済み取引戦略。",
        "sentiment-altdata": "ニュース、SNS、NLP、LLM、決算説明会、代替データによる投資シグナル。",
        "high-frequency": "市場マイクロストラクチャー、板情報、注文フロー、流動性、高頻度取引。",
        "risk-volatility": "ボラティリティ予測、VaR、ドローダウン、テールリスク、ストレステスト、予測評価。",
        "asset-pricing": "資産価格評価、株式ファクター、アノマリー、リターン予測可能性、クロスセクション株式リターン。",
        "behavioral-finance": "投資家心理、注目度、行動バイアス、市場効率性、意思決定。",
    },
}

CATEGORY_LIMITATIONS_LANG = {
    "ko": {
        "Asset Pricing and Return Predictability": ["수익률 예측성은 거래비용, crowded trade, 다중검정 보정 이후 약해질 수 있습니다.", "인용 영향력은 잘 알려진 anomaly를 선호할 수 있으며 실제 성과 약화를 반영하지 못할 수 있습니다."],
        "Portfolio Optimization and Asset Allocation": ["배분 성과는 추정오차, turnover, 공매도 제약, 벤치마크 선택에 크게 좌우될 수 있습니다.", "메타데이터만으로 비용, 레버리지, 리밸런싱 규칙을 확인할 수 없습니다."],
        "Machine Learning for Stock Prediction": ["예측 정확도가 비용 차감 후 위험조정수익으로 이어지지 않을 수 있습니다.", "feature leakage, 생존편향, 비정상성은 지속적인 위험입니다."],
        "Deep Learning and Financial Time Series": ["딥러닝 모델은 잡음이 크고 비정상적인 금융 데이터에 과적합될 수 있습니다.", "검증 프로토콜이 조금만 달라져도 보고된 성과가 흔들릴 수 있습니다."],
        "Reinforcement Learning and Algorithmic Trading": ["시뮬레이션 보상은 slippage, latency, market impact가 있는 실제 거래와 다를 수 있습니다.", "정책은 regime shift와 stress period에서 취약할 수 있습니다."],
        "Sentiment, News, and Alternative Data": ["텍스트 및 대체데이터 연구에서는 timestamp 정렬과 look-ahead bias 통제가 핵심입니다.", "정보가 crowded해지면 sentiment signal은 빠르게 약화될 수 있습니다."],
        "Market Microstructure and High-Frequency Trading": ["고빈도 결과는 시장 접근, latency, 수수료, order-book 재구성 품질에 민감합니다.", "공개 메타데이터는 실행 가정을 충분히 보여주지 않습니다."],
        "Risk, Volatility, and Forecast Evaluation": ["위험 예측은 regime break와 극단 사건에서 실패할 수 있습니다.", "loss function 선택이 모델 순위를 크게 바꿀 수 있습니다."],
        "Behavioral Finance and Investor Decision-Making": ["행동 메커니즘은 시장별·기간별로 달라질 수 있습니다.", "sentiment proxy는 여러 혼재된 경로를 동시에 포착할 수 있습니다."],
        "General Finance, Surveys, and Trading Systems": ["서베이와 시스템 논문은 인용을 많이 받을 수 있지만 실증 근거는 시장별로 다를 수 있습니다.", "메타데이터 기반 순위는 전문적인 원문 독해와 백테스트 검증을 대체하지 않습니다."],
    },
    "zh": {
        "Asset Pricing and Return Predictability": ["收益可预测性在交易成本、拥挤交易和多重检验修正后可能消失。", "引用影响力可能偏向知名异象，即使其实盘表现后来减弱。"],
        "Portfolio Optimization and Asset Allocation": ["配置收益可能被估计误差、换手率、卖空约束和基准选择所主导。", "仅靠元数据无法确认成本、杠杆或再平衡规则。"],
        "Machine Learning for Stock Prediction": ["预测准确率不一定转化为扣除成本后的风险调整收益。", "特征泄漏、生存者偏差和非平稳性仍是主要风险。"],
        "Deep Learning and Financial Time Series": ["深度模型可能过拟合噪声大、非平稳且信噪比低的金融数据。", "验证协议的细微变化也可能影响报告结果。"],
        "Reinforcement Learning and Algorithmic Trading": ["模拟奖励可能与含滑点、延迟和市场冲击的可执行交易不同。", "策略在 regime shift 或压力时期可能脆弱。"],
        "Sentiment, News, and Alternative Data": ["文本和另类数据研究必须严格处理时间戳对齐与前视偏差。", "当信息变得拥挤时，情绪信号可能快速衰减。"],
        "Market Microstructure and High-Frequency Trading": ["高频结果高度依赖市场接入、延迟、费用和订单簿重建质量。", "公开元数据很少揭示真实执行假设。"],
        "Risk, Volatility, and Forecast Evaluation": ["风险预测在 regime break 和极端事件下可能失效。", "损失函数选择会显著改变模型排名。"],
        "Behavioral Finance and Investor Decision-Making": ["行为机制可能具有市场和时期特异性。", "情绪代理变量可能混合多个渠道。"],
        "General Finance, Surveys, and Trading Systems": ["综述和系统论文可能占据高引用，但实证证据在不同市场中并不一致。", "元数据排序不能替代专家阅读全文和回测验证。"],
    },
    "ja": {
        "Asset Pricing and Return Predictability": ["リターン予測可能性は、取引コスト、混雑、複数検定補正後に消えることがあります。", "引用影響度は有名なアノマリーを過大評価する場合があります。"],
        "Portfolio Optimization and Asset Allocation": ["配分の改善は、推定誤差、回転率、空売り制約、ベンチマーク選択に左右されます。", "メタデータだけではコスト、レバレッジ、リバランス規則を確認できません。"],
        "Machine Learning for Stock Prediction": ["予測精度は、コスト控除後のリスク調整リターンに直結しない場合があります。", "特徴量リーク、生存者バイアス、非定常性は継続的な懸念です。"],
        "Deep Learning and Financial Time Series": ["深層モデルはノイズが大きく非定常な金融データに過学習しやすいです。", "検証手順の小さな違いで報告性能が変わることがあります。"],
        "Reinforcement Learning and Algorithmic Trading": ["シミュレーション報酬は、スリッページ、遅延、市場インパクトを含む実行可能取引と異なる場合があります。", "方策は regime shift やストレス期に脆弱になり得ます。"],
        "Sentiment, News, and Alternative Data": ["テキスト・代替データではタイムスタンプ整合性と look-ahead bias の管理が重要です。", "情報が混雑すると sentiment signal は急速に弱まる可能性があります。"],
        "Market Microstructure and High-Frequency Trading": ["高頻度結果は市場アクセス、遅延、手数料、注文簿再構成品質に強く依存します。", "公開メタデータは実行仮定を十分に示しません。"],
        "Risk, Volatility, and Forecast Evaluation": ["リスク予測は regime break や極端事象で失敗することがあります。", "損失関数の選択がモデル順位を大きく左右します。"],
        "Behavioral Finance and Investor Decision-Making": ["行動メカニズムは市場や時期によって変化します。", "センチメント代理変数は複数の経路を混在して捉える場合があります。"],
        "General Finance, Surveys, and Trading Systems": ["サーベイやシステム論文は高引用になりやすい一方、実証証拠は市場ごとに異なります。", "メタデータ順位は専門家による全文読解やバックテスト検証を代替しません。"],
    },
}


def localized_category_name(category, language):
    return CATEGORY_NAMES.get(language, {}).get(category, category)


def category_display_name(category, language):
    return localized_category_name(category, language)


def localized_limitations_for(category, language):
    if language == "en":
        return TAXONOMY_LIMITATIONS.get(category, TAXONOMY_LIMITATIONS["General Finance, Surveys, and Trading Systems"])[:2]
    return CATEGORY_LIMITATIONS_LANG.get(language, {}).get(
        category,
        CATEGORY_LIMITATIONS_LANG[language]["General Finance, Surveys, and Trading Systems"],
    )


def localized_strengths_for(paper, language):
    citations = int(paper.get("citationCount") or 0)
    influential = int(paper.get("influentialCitationCount") or 0)
    pieces = []
    if language == "en":
        return paper["strengths"]
    if language == "ko":
        if citations >= 100:
            pieces.append(f"높은 인용 신호({citations:,}회)")
        if influential >= 10:
            pieces.append(f"영향력 인용 신호({influential:,}회)")
        if paper.get("openAccessPdf"):
            pieces.append("오픈액세스 PDF 메타데이터 존재")
        if "recognized finance/AI venue" in str(paper.get("importanceReasons") or ""):
            pieces.append("금융/AI 주요 venue 신호")
        return "; ".join(pieces or ["감사된 후보군에서 인용수 기준으로 선정"])
    if language == "zh":
        if citations >= 100:
            pieces.append(f"高引用信号（{citations:,}次）")
        if influential >= 10:
            pieces.append(f"影响力引用信号（{influential:,}次）")
        if paper.get("openAccessPdf"):
            pieces.append("包含开放获取 PDF 元数据")
        if "recognized finance/AI venue" in str(paper.get("importanceReasons") or ""):
            pieces.append("具有金融/AI 重要期刊或会议信号")
        return "; ".join(pieces or ["从审计后的候选池中按引用数入选"])
    if citations >= 100:
        pieces.append(f"高い引用シグナル（{citations:,}件）")
    if influential >= 10:
        pieces.append(f"影響力引用シグナル（{influential:,}件）")
    if paper.get("openAccessPdf"):
        pieces.append("オープンアクセス PDF メタデータあり")
    if "recognized finance/AI venue" in str(paper.get("importanceReasons") or ""):
        pieces.append("金融/AI の主要 venue シグナル")
    return "; ".join(pieces or ["監査済み候補群から引用数により選定"])


def localized_paper_fields(paper, language):
    if language == "en":
        return {
            "keyIdea": paper["keyIdea"],
            "strengths": paper["strengths"],
            "limitations": paper["limitations"],
        }
    title = paper.get("title") or "this paper"
    category = localized_category_name(paper.get("category") or category_for(paper), language)
    if language == "ko":
        key_idea = f"이 논문은 \"{title}\"을(를) 중심으로 {category} 관점에서 주식 투자와 AI 트레이딩 연구의 핵심 쟁점을 다룹니다."
    elif language == "zh":
        key_idea = f"本文围绕《{title}》，从{category}视角梳理其对股票投资与 AI 交易研究的贡献。"
    else:
        key_idea = f"この論文は『{title}』を軸に、{category}の観点から株式投資と AI 取引研究への示唆を整理します。"
    return {
        "keyIdea": key_idea,
        "strengths": localized_strengths_for(paper, language),
        "limitations": "; ".join(localized_limitations_for(paper.get("category") or category_for(paper), language)),
    }


def language_period_analysis(language, category, rows, start, end):
    count = len(rows)
    citations = sum(p["citationCount"] for p in rows)
    top = max(rows, key=lambda p: (p["citationCount"], p["influentialCitationCount"], p["title"]))
    active_year, active_count = Counter(p["year"] for p in rows).most_common(1)[0]
    tags = ", ".join(top_metadata_values(rows, "methodTags")) or "metadata-ranked"
    venues = ", ".join(top_metadata_values(rows, "venue")) or "mixed venues"
    name = localized_category_name(category, language)
    if language == "ko":
        overview = [
            f"{start}-{end} 기간의 {name} 분류에는 선정 논문 {count:,}편과 인용 {citations:,}회가 포함됩니다. 가장 활발한 연도는 {active_year}년({active_count:,}편)이며, 인용 기준 대표 논문은 \"{top['title']}\"({top['citationCount']:,}회 인용)입니다.",
            f"자주 등장하는 방법 태그는 {tags}이고, 주요 venue는 {venues}입니다.",
        ]
        recency = "인용수 기반 기간 분석은 최신 논문에 구조적으로 불리하므로, 최신성은 별도 전문가 검토로 보완해야 합니다."
    elif language == "zh":
        overview = [
            f"在 {start}-{end} 期间，{name} 分类包含 {count:,} 篇入选论文和 {citations:,} 次引用。最活跃年份为 {active_year} 年（{active_count:,} 篇），引用排名最高的代表论文是《{top['title']}》（{top['citationCount']:,} 次引用）。",
            f"常见方法标签包括 {tags}，主要 venue 集中在 {venues}。",
        ]
        recency = "基于引用数的期间视图会系统性低估新近论文，因此需要额外的专家新近性审查。"
    elif language == "ja":
        overview = [
            f"{start}-{end} 年の {name} 分類には、選定論文 {count:,} 本と {citations:,} 件の引用が含まれます。最も活発な年は {active_year} 年（{active_count:,} 本）で、引用順位の代表論文は『{top['title']}』（{top['citationCount']:,} 件）です。",
            f"頻出する手法タグは {tags} で、主な venue は {venues} です。",
        ]
        recency = "引用数ベースの期間分析は新しい論文に不利なため、新規性は別途専門家レビューで補う必要があります。"
    else:
        overview = [
            f"In {start}-{end}, {name} contains {count:,} selected papers and {citations:,} citations. The most active year is {active_year} ({active_count:,} papers), and the leading citation-ranked paper is \"{top['title']}\" ({top['citationCount']:,} citations).",
            f"Frequent method tags include {tags}, with venue concentration around {venues}.",
        ]
        recency = "Citation-ranked period views structurally disadvantage very recent papers, so novelty needs separate expert review."
    return {
        "categoryName": name,
        "overview": overview,
        "limitations": [*localized_limitations_for(category, language), recency],
    }


def research_copy():
    return {
        "en": research_overview_html(),
        "ko": "<h2 id='research-timeline-title'>연구 타임라인</h2><div class='timeline-copy'><p>2000-2026년 trade corpus는 자산가격결정, 기술적 분석, 변동성, 포트폴리오 연구에서 머신러닝, 딥러닝, 감성 기반 투자, 강화학습, 고빈도 시장 미시구조로 확장되는 흐름을 보여줍니다.</p><p>이 지도는 투자 조언이 아니라 주식 투자 연구의 인용 영향력 지도입니다.</p></div>",
        "zh": "<h2 id='research-timeline-title'>研究时间线</h2><div class='timeline-copy'><p>2000-2026 年的 trade corpus 展示了从资产定价、技术分析、波动率和组合研究，向机器学习、深度学习、情绪投资、强化学习和高频市场微观结构扩展的过程。</p><p>该视图是股票投资研究的引用影响力地图，并非投资建议。</p></div>",
        "ja": "<h2 id='research-timeline-title'>研究タイムライン</h2><div class='timeline-copy'><p>2000-2026 年の trade corpus は、資産価格評価、テクニカル分析、ボラティリティ、ポートフォリオ研究から、機械学習、深層学習、センチメント投資、強化学習、高頻度市場マイクロストラクチャーへ広がる流れを示します。</p><p>これは株式投資研究の引用影響マップであり、投資助言ではありません。</p></div>",
    }


def overall_research_templates():
    return {
        "en": {
            "timelineTitle": "Research Timeline",
            "summary": [
                "For {range}, the trade corpus contains {papers} selected papers across {activeYears} active years, with {citations} citations. The strongest taxonomy signals are {topCategories}, and the most active year is {peakYear} ({peakYearCount} papers).",
                "The leading citation-ranked paper is \"{topPaper}\" ({topPaperYear}, {topPaperCitations} citations) in {topPaperCategory}. Keywords such as {topKeywords} show how this period connects asset pricing, forecasting, portfolio construction, risk, market microstructure, sentiment, and AI-based trading.",
            ],
            "insightsTitle": "Research Insights",
            "insights": [
                {"label": "Signal Quality", "title": "Prediction must survive market frictions", "body": "{topCategories} dominate {range}, showing where citation-ranked influence concentrates.", "implication": "Implication: evaluate signals after costs, turnover, capacity, and risk constraints."},
                {"label": "Citation Mass", "title": "Citation peaks mark reusable finance infrastructure", "body": "The selected range carries {citations} citations, with citation mass peaking around {peakCitationYear}.", "implication": "Implication: high-impact papers often define data, factors, methods, or validation conventions."},
                {"label": "AI Trading", "title": "AI expands the signal surface", "body": "Frequent tags such as {topKeywords} indicate whether the period centers on forecasting, deep learning, sentiment, portfolio construction, or execution.", "implication": "Implication: richer models still need realistic backtests and market-aware validation."},
                {"label": "Open Gaps", "title": "Recent papers need recency correction", "body": "Newer AI trading papers can be underweighted before citation signals mature.", "implication": "Implication: use this map as a stable citation view, then layer recent expert review on top."},
            ],
        },
        "ko": {
            "timelineTitle": "연구 타임라인",
            "summary": [
                "{range} 기간의 trade corpus는 {activeYears}개 활성 연도에 걸쳐 선정 논문 {papers}편과 인용 {citations}회를 포함합니다. 가장 강한 taxonomy 신호는 {topCategories}이며, 논문 수가 가장 많은 연도는 {peakYear}년({peakYearCount}편)입니다.",
                "인용 기준 대표 논문은 {topPaperCategory} 분류의 \"{topPaper}\"({topPaperYear}, {topPaperCitations}회 인용)입니다. {topKeywords} 같은 키워드는 이 기간이 자산가격결정, 예측, 포트폴리오, 위험, 시장 미시구조, 감성, AI 기반 트레이딩을 어떻게 연결하는지 보여줍니다.",
            ],
            "insightsTitle": "연구 인사이트",
            "insights": [
                {"label": "신호 품질", "title": "예측은 시장 마찰을 통과해야 합니다", "body": "{range}에서는 {topCategories}가 인용 영향력을 주도합니다.", "implication": "시사점: 신호는 비용, turnover, capacity, 위험 제약 이후에 평가해야 합니다."},
                {"label": "인용 규모", "title": "인용 피크는 재사용 가능한 금융 인프라를 보여줍니다", "body": "선택 기간은 {citations}회의 인용을 가지며 citation mass는 {peakCitationYear}년 부근에서 가장 큽니다.", "implication": "시사점: 영향력 높은 논문은 데이터, 팩터, 방법론, 검증 관행을 정의하는 경우가 많습니다."},
                {"label": "AI 트레이딩", "title": "AI는 신호 공간을 확장합니다", "body": "{topKeywords} 같은 태그는 해당 기간이 forecasting, deep learning, sentiment, portfolio, execution 중 어디에 집중되는지 보여줍니다.", "implication": "시사점: 풍부한 모델도 현실적인 백테스트와 시장 인식 검증이 필요합니다."},
                {"label": "열린 과제", "title": "최신 논문에는 recency 보정이 필요합니다", "body": "새로운 AI 트레이딩 논문은 인용 신호가 성숙하기 전까지 과소평가될 수 있습니다.", "implication": "시사점: 이 지도는 안정적인 citation view로 사용하고, 최신 전문 리뷰를 덧붙여야 합니다."},
            ],
        },
        "zh": {
            "timelineTitle": "研究时间线",
            "summary": [
                "在 {range} 期间，trade corpus 覆盖 {activeYears} 个活跃年份，包含 {papers} 篇入选论文和 {citations} 次引用。最强的 taxonomy 信号是 {topCategories}，论文数最多的年份是 {peakYear}（{peakYearCount} 篇）。",
                "引用排名最高的代表论文是 {topPaperCategory} 分类中的《{topPaper}》（{topPaperYear}，{topPaperCitations} 次引用）。{topKeywords} 等关键词显示该期间如何连接资产定价、预测、组合构建、风险、市场微观结构、情绪和 AI 交易。",
            ],
            "insightsTitle": "研究洞察",
            "insights": [
                {"label": "信号质量", "title": "预测必须经受市场摩擦检验", "body": "{topCategories} 主导 {range} 的引用影响力。", "implication": "启示：信号应在成本、换手率、容量和风险约束之后评估。"},
                {"label": "引用规模", "title": "引用峰值标记可复用的金融基础设施", "body": "所选期间共有 {citations} 次引用，引用规模在 {peakCitationYear} 附近达到峰值。", "implication": "启示：高影响论文往往定义数据、因子、方法或验证惯例。"},
                {"label": "AI 交易", "title": "AI 扩展了信号空间", "body": "{topKeywords} 等高频标签显示该时期是否聚焦于预测、深度学习、情绪、组合构建或执行。", "implication": "启示：更复杂模型仍需要现实回测和市场感知验证。"},
                {"label": "开放问题", "title": "新近论文需要时间校正", "body": "新的 AI 交易论文在引用信号成熟前可能被低估。", "implication": "启示：将该地图作为稳定引用视图，再叠加最新专家审查。"},
            ],
        },
        "ja": {
            "timelineTitle": "研究タイムライン",
            "summary": [
                "{range} の trade corpus は {activeYears} 年にわたり、選定論文 {papers} 本と {citations} 件の引用を含みます。最も強い taxonomy シグナルは {topCategories} で、最も活発な年は {peakYear} 年（{peakYearCount} 本）です。",
                "引用順位の代表論文は {topPaperCategory} の『{topPaper}』（{topPaperYear}、{topPaperCitations} 件）です。{topKeywords} などのキーワードは、この期間が資産価格評価、予測、ポートフォリオ構築、リスク、市場マイクロストラクチャー、センチメント、AI 取引をどう結びつけるかを示します。",
            ],
            "insightsTitle": "研究インサイト",
            "insights": [
                {"label": "シグナル品質", "title": "予測は市場摩擦を越えて有効である必要があります", "body": "{range} では {topCategories} が引用影響力を主導しています。", "implication": "示唆：シグナルはコスト、回転率、容量、リスク制約の後で評価すべきです。"},
                {"label": "引用規模", "title": "引用ピークは再利用可能な金融インフラを示します", "body": "選択期間は {citations} 件の引用を持ち、引用規模は {peakCitationYear} 年付近でピークになります。", "implication": "示唆：影響力の高い論文はデータ、ファクター、方法、検証慣行を定義することが多いです。"},
                {"label": "AI 取引", "title": "AI はシグナル空間を広げます", "body": "{topKeywords} などのタグは、期間が予測、深層学習、センチメント、ポートフォリオ構築、執行のどこに焦点を置くかを示します。", "implication": "示唆：高度なモデルにも現実的なバックテストと市場を意識した検証が必要です。"},
                {"label": "未解決課題", "title": "新しい論文には recency 補正が必要です", "body": "新しい AI 取引論文は引用シグナルが成熟するまで過小評価される可能性があります。", "implication": "示唆：このマップを安定した引用ビューとして使い、最新の専門家レビューを重ねてください。"},
            ],
        },
    }


def site_static_copy():
    base_hero = f"A taxonomy-first, citation-ranked map of stock investment and AI-driven trading research from {START_YEAR} through {END_YEAR}. Each year investigates up to {CANDIDATES_PER_YEAR:,} candidate papers; the final collection selects the top {TARGET_PER_YEAR:,} papers from each year by citation count ({TARGET_TOTAL:,} papers total)."
    return {
        "en": {
            "hero": base_hero,
            "navReadme": "README",
            "navDataset": "CSV Dataset",
            "navTaxonomyCsv": "Taxonomy CSV",
            "navPeriodJson": "Period Analysis JSON",
            "navKeywords": "Keywords Convention",
            "navCandidatePool": "Candidate Pool",
            "navReview": "Review Paper",
            "navKoreanReview": "Korean Review",
            "statSelected": "selected papers",
            "statYears": "years represented",
            "statCitations": "citation count total",
            "statCategories": "topic categories",
            "filterPeriod": "Period",
            "filterLanguage": "Language",
            "filterStart": "Start year",
            "filterEnd": "End year",
            "filterReset": "Reset",
            "keywordsTitle": "Keywords Convention",
            "keywordsIntro": "These clickable keyword tags define the trade-specific convention used to scan, filter, and extend this collection.",
            "taxonomyTitle": "Taxonomy",
            "taxonomyIntro": "Each taxonomy section lists papers with publication year, venue, citation count, influential citations, score, keywords, key idea, strengths, research-focused limitations, and paper links. Sections are collapsed by default.",
            "categoryNames": {},
        },
        "ko": {
            "hero": f"{START_YEAR}-{END_YEAR}년 주식 투자 및 AI 기반 트레이딩 연구를 taxonomy-first, citation-ranked 방식으로 정리한 지도입니다. 각 연도별 최대 {CANDIDATES_PER_YEAR:,}편의 후보 논문을 조사하고, 인용수 기준 상위 {TARGET_PER_YEAR:,}편씩 총 {TARGET_TOTAL:,}편을 선정했습니다.",
            "navReadme": "README",
            "navDataset": "CSV 데이터셋",
            "navTaxonomyCsv": "Taxonomy CSV",
            "navPeriodJson": "기간 분석 JSON",
            "navKeywords": "키워드 규칙",
            "navCandidatePool": "후보 논문 풀",
            "navReview": "영문 리뷰",
            "navKoreanReview": "한국어 리뷰",
            "statSelected": "선정 논문",
            "statYears": "포함 연도",
            "statCitations": "총 인용수",
            "statCategories": "주제 분류",
            "filterPeriod": "기간",
            "filterLanguage": "언어",
            "filterStart": "시작 연도",
            "filterEnd": "종료 연도",
            "filterReset": "초기화",
            "keywordsTitle": "키워드 규칙",
            "keywordsIntro": "아래 키워드 태그는 이 컬렉션을 검색, 필터링, 확장할 때 사용하는 거래 연구용 규칙입니다.",
            "taxonomyTitle": "분류 체계",
            "taxonomyIntro": "각 분류 섹션은 출판연도, venue, 인용수, 영향력 인용, score, 키워드, 핵심 아이디어, 강점, 연구상 한계, 논문 링크를 보여줍니다. 섹션은 기본적으로 접혀 있습니다.",
            "categoryNames": CATEGORY_NAMES["ko"],
        },
        "zh": {
            "hero": f"这是 {START_YEAR}-{END_YEAR} 年股票投资与 AI 交易研究的 taxonomy-first、citation-ranked 地图。每年最多调查 {CANDIDATES_PER_YEAR:,} 篇候选论文，并按引用数选出前 {TARGET_PER_YEAR:,} 篇，共 {TARGET_TOTAL:,} 篇。",
            "navReadme": "README",
            "navDataset": "CSV 数据集",
            "navTaxonomyCsv": "Taxonomy CSV",
            "navPeriodJson": "期间分析 JSON",
            "navKeywords": "关键词规则",
            "navCandidatePool": "候选论文池",
            "navReview": "英文综述",
            "navKoreanReview": "韩文综述",
            "statSelected": "入选论文",
            "statYears": "覆盖年份",
            "statCitations": "总引用数",
            "statCategories": "主题分类",
            "filterPeriod": "期间",
            "filterLanguage": "语言",
            "filterStart": "开始年份",
            "filterEnd": "结束年份",
            "filterReset": "重置",
            "keywordsTitle": "关键词规则",
            "keywordsIntro": "这些可点击关键词定义了本集合用于检索、过滤和扩展的交易研究规则。",
            "taxonomyTitle": "分类体系",
            "taxonomyIntro": "每个分类部分列出论文年份、venue、引用数、影响力引用、score、关键词、核心观点、优势、研究局限和论文链接。各部分默认折叠。",
            "categoryNames": CATEGORY_NAMES["zh"],
        },
        "ja": {
            "hero": f"{START_YEAR}-{END_YEAR} 年の株式投資と AI 取引研究を taxonomy-first、citation-ranked 方式で整理したマップです。各年最大 {CANDIDATES_PER_YEAR:,} 本の候補論文を調査し、引用数上位 {TARGET_PER_YEAR:,} 本ずつ、合計 {TARGET_TOTAL:,} 本を選定しました。",
            "navReadme": "README",
            "navDataset": "CSV データセット",
            "navTaxonomyCsv": "Taxonomy CSV",
            "navPeriodJson": "期間分析 JSON",
            "navKeywords": "キーワード規則",
            "navCandidatePool": "候補論文プール",
            "navReview": "英語レビュー",
            "navKoreanReview": "韓国語レビュー",
            "statSelected": "選定論文",
            "statYears": "対象年",
            "statCitations": "総引用数",
            "statCategories": "主題分類",
            "filterPeriod": "期間",
            "filterLanguage": "言語",
            "filterStart": "開始年",
            "filterEnd": "終了年",
            "filterReset": "リセット",
            "keywordsTitle": "キーワード規則",
            "keywordsIntro": "これらのクリック可能なキーワードタグは、このコレクションを検索、フィルタ、拡張するための取引研究用ルールです。",
            "taxonomyTitle": "分類体系",
            "taxonomyIntro": "各分類セクションには、出版年、venue、引用数、影響力引用、score、キーワード、主要アイデア、強み、研究上の限界、論文リンクが表示されます。セクションは既定で折りたたまれています。",
            "categoryNames": CATEGORY_NAMES["ja"],
        },
    }


def site_keyword_convention_html():
    rows = []
    for keyword, description, color in KEYWORD_CONVENTION:
        attrs = {
            "data-description-en": description,
            "data-description-ko": KEYWORD_DESCRIPTION_LANG["ko"].get(keyword, description),
            "data-description-zh": KEYWORD_DESCRIPTION_LANG["zh"].get(keyword, description),
            "data-description-ja": KEYWORD_DESCRIPTION_LANG["ja"].get(keyword, description),
        }
        attr_text = " ".join(f"{key}='{html.escape(value, quote=True)}'" for key, value in attrs.items())
        rows.append(
            f"<button class='keyword-item' type='button' data-keyword='{html.escape(keyword)}' aria-pressed='false'>"
            f"<span class='keyword-chip' style='--chip-color:#{color}'>{html.escape(keyword)}</span>"
            f"<span class='keyword-description' {attr_text}>{html.escape(description)}</span></button>"
        )
    return "\n".join(rows)


def main():
    for path in (DATA_DIR, DOCS_DIR, PAPER_DIR):
        path.mkdir(exist_ok=True)
    if "--reuse-candidates" in sys.argv:
        selected, selected_by_year, candidates_by_year = reuse_existing_candidates()
    else:
        selected, selected_by_year, candidates_by_year = collect_papers()
    all_candidate_rows = [p for rows in candidates_by_year.values() for p in rows]
    write_github_links(selected + all_candidate_rows)
    candidates = write_data(selected, selected_by_year, candidates_by_year)
    rewrite_data_metadata(len(selected))
    write_link_audit(selected, candidates)
    write_taxonomy_dataset(selected)
    write_period_analysis(selected)
    write_readme(selected, candidates)
    write_taxonomy_icons(selected)
    write_charts(selected)
    write_site(selected)
    write_review_html(selected, korean=False)
    write_review_html(selected, korean=True)
    write_review_docx(selected)
    write_project_files(selected, candidates)
    write_skill2_provenance(selected, candidates)
    copy_public_assets()
    postprocess_generated_files()
    print(f"[done] generated {len(selected):,} selected papers from {len(candidates):,} candidates", flush=True)


if __name__ == "__main__":
    main()
